#1) python imports packages include here
import os
import json
from datetime import datetime, timedelta
from random import randint
from zipfile import ZipFile 
import operator
import requests
import subprocess as sp
from docx import Document
# import docx2txt
from urllib.request import urlopen
from bs4 import BeautifulSoup
from zipfile import ZipFile
from io import BytesIO
import shutil
import urllib, mimetypes
import re
import pandas as pd
from PyPDF2 import PdfFileReader, PdfFileWriter
import glob
from PIL import Image
# from pptx import Presentation

try:
	from functools import reduce
except ImportError:  # Python < 3
	pass
from decimal import Decimal
#2) all django imports
from django.conf import settings
from django.shortcuts import render, redirect
from django.http import HttpResponse, Http404, JsonResponse, StreamingHttpResponse
from django.core.mail import send_mail
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout
from django.utils import timezone 
from django.shortcuts import get_list_or_404, get_object_or_404
from django.utils.crypto import get_random_string
from django.core.files.storage import FileSystemStorage
from django.db.models.signals import post_delete,post_save
from wsgiref.util import FileWrapper
from django.http import FileResponse
from django.db.models import Max, F, Min, Q, Sum
from django.core import serializers as sez
from users_and_permission.models import DataroomGroupFolderSpecificPermissions, DataroomGroups
from users_and_permission.serializers import DataroomGroupFolderSpecificPermissionsSerializer
from django.core.mail import EmailMessage, EmailMultiAlternatives
from django.template import Context

from django.template.loader import render_to_string, get_template




#3) all rest framework packages include here
from rest_framework.generics import (
	UpdateAPIView,
	ListAPIView,
	ListCreateAPIView)
from rest_framework.authtoken.models import Token
from rest_framework.authentication import TokenAuthentication, SessionAuthentication, BasicAuthentication
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework import status
from rest_framework.parsers import JSONParser, FileUploadParser
from rest_framework.renderers import JSONRenderer, TemplateHTMLRenderer
from rest_framework import permissions, routers, serializers, viewsets
from rest_framework.decorators import api_view, renderer_classes, permission_classes

#4) all common app imports
from userauth.serializers import UserSerializer, AccessHistorySerializer
from userauth import constants, utils
from userauth.models import Profile, AccessHistory, User

#5) include all apps packages here

# import all required modules here 
from .models import DataroomFolder, IndexDownload,FolderView,FolderPrint,FolderDownload, Categories, ManageDataroomCategories
# import all required serializers here 
from .serializers import DataroomFolderSerializer, IndexDownloadSerializer, FolderViewSerializer,FolderPrintSerializer,FolderDownloadSerializer, CategoriesSerializer, RcentUpdateSerializer, RecentUpdateSerializer, FolderActivityPrintSerializer, FolderActivityDownloadSerializer, FolderActivityViewSerializer
from dataroom.models import Dataroom, DataroomOverview, DataroomView, Watermarking
from dataroom.serializers import DataroomViewSerializer,DataroomSerializer
from users_and_permission.models import DataroomMembers, RcentUpdate
from users_and_permission.serializers import *
# add debugger here
import pdb
from itertools import chain
from qna.models import *
from qna.serializers import *
import time
from rest_framework import pagination, serializers
from .utils import *

paginator = pagination.PageNumberPagination()

def change_all_indexes(delete_folder):
	#pdb.set_trace()
	if delete_folder.is_folder == True:
		if delete_folder.is_root_folder:
			dataroom_folder = DataroomFolder.objects.filter(is_root_folder=True, is_deleted=False, index__gte = delete_folder.index, is_folder=True).update(index = F('index')-1 )
			dataroom_folder = DataroomFolder.objects.filter(is_root_folder=True, is_deleted=False, is_folder=True).order_by('index')
			return dataroom_folder
		else:
			dataroom_folder = DataroomFolder.objects.filter(parent_folder_id=delete_folder.parent_folder_id, is_deleted=False, index__gte = delete_folder.index, is_folder=True, is_root_folder=False).update(index = F('index')-1)  
			dataroom_folder = DataroomFolder.objects.filter(parent_folder_id=delete_folder.parent_folder_id, is_deleted=False,  is_root_folder=False).order_by('index')
			return dataroom_folder
	else:
		dataroom_folder = DataroomFolder.objects.filter(parent_folder_id=delete_folder.parent_folder_id, is_deleted=False, index__gte = delete_folder.index, is_folder=False, is_root_folder=False).update(index = F('index')-1)  
		dataroom_folder = DataroomFolder.objects.filter(parent_folder_id=delete_folder.parent_folder_id, is_deleted=False,  is_root_folder=False).order_by('index')
		return dataroom_folder
	   
class DeleteDataroomFolder(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def delete(self, request, uuid, format=None):
		user = request.user
		current_date = timezone.now()

		try:
			#current_date = datetime.now()
			delete_folder = DataroomFolder.objects.get(dataroom_folder_uuid=uuid)
			delete_folder.is_deleted = True
			delete_folder.deleted_by = user
			delete_folder.deleted_by_date = current_date#add current date
			delete_folder.save()
			dataroom_delete_folders = change_all_indexes(delete_folder)
			# print ("dataroom delete folders", dataroom_delete_folders)
			dataroom_folder_serializer = DataroomFolderSerializer(dataroom_delete_folders, many=True)
			#data = {"msg":"Folder/File successfully deleted", 'data': dataroom_folder_serializer.data}
			# print ("all data is", dataroom_folder_serializer.data)
			return Response(dataroom_folder_serializer.data, status=status.HTTP_201_CREATED)
		except DataroomFolder.DoesNotExist:
			data = {"msg":"Folder/File does not exist"}
			return Response(data, status=status.HTTP_400_BAD_REQUEST)

class getAllSubfolderList(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		from_date = request.GET.get('from_date')
		to_date = request.GET.get('to_date')
		name = request.GET.get('name')
		file_type = request.GET.get('type')
		uploaded_by = request.GET.get('uploaded_by')
		print("name", name, "uploaded_by", uploaded_by, "from_date", from_date, "to_date", to_date)
		import datetime
		if to_date == '':
			todays_date = ''
		else:
			todays_date = datetime.datetime.strftime(datetime.datetime.strptime( to_date, '%Y-%m-%d'),'%Y-%m-%d 23:59:59+05:30')
		if from_date == '':
			first_date = ''
		else:
			first_date = datetime.datetime.strftime(datetime.datetime.strptime( from_date, '%Y-%m-%d'),'%Y-%m-%d 00:00:00+05:30')
		q_list = [Q(uploaded_by__last_name__icontains=uploaded_by), Q(uploaded_by__first_name__icontains=uploaded_by), Q(name__icontains=name), Q(name__icontains=file_type)]
		if from_date == '' and to_date == '':
			dataroom_folders = DataroomFolder.objects.filter(parent_folder_id = pk, is_root_folder=False, is_deleted=False).filter(reduce(operator.or_, q_list)).order_by('index')
		elif from_date == '' and to_date != '':
			dataroom_folders = DataroomFolder.objects.filter(parent_folder_id = pk, is_root_folder=False, is_deleted=False,  created_date__lte=todays_date).filter(reduce(operator.or_, q_list)).order_by('index')
		elif from_date != '' and to_date == '':
			dataroom_folders = DataroomFolder.objects.filter(parent_folder_id = pk, is_root_folder=False, is_deleted=False,  created_date__gte=first_date).filter(reduce(operator.or_, q_list)).order_by('index')
		else:
			dataroom_folders = DataroomFolder.objects.filter(parent_folder_id = pk, is_root_folder=False, is_deleted=False,  created_date__gte=first_date, created_date__lte=todays_date).filter(reduce(operator.or_, q_list)).order_by('index')
		dataroom_folder_serializer = DataroomFolderSerializer(dataroom_folders, many=True)
		data = dataroom_folder_serializer.data
		# #print("daaaaaaaaaaaaaaaaa", data)
		for da in data:
			da['perm'] = {}
			da['names'] = da['name']
			try:
				member = DataroomMembers.objects.filter(member_id=user.id, dataroom_id=da['dataroom'], is_deleted=False).first()
				#print("10_april_today====>",member.is_la_user, member.is_dataroom_admin)
				if member.is_la_user == True or member.is_dataroom_admin == True:
					da['perm']['is_view_only'] = True
					da['perm']['is_no_access'] = False
					da['perm']['is_access'] = True
					da['perm']['is_view_and_print'] = True
					da['perm']['is_view_and_print_and_download'] = True
					da['perm']['is_upload'] = True
					da['perm']['is_watermarking'] = True
					da['perm']['is_drm'] = True
					da['perm']['is_editor'] = True  
					da['perm']['is_shortcut'] = True  
				else:
					# #print("member--------", member.end_user_group.first().id, user.id, da['dataroom'])
					perm_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=da['id'],dataroom_id=da['dataroom'], dataroom_groups_id=member.end_user_group.first().id)
					print(da['id'],"ravishankar")
					grp_perm = DataroomGroupPermission.objects.filter(dataroom_groups_id=member.end_user_group.first().id, dataroom_id=da['dataroom']).first()
					# #print(grp_perm.id)
					if grp_perm.is_watermarking:
						if grp_perm.is_doc_as_pdf or grp_perm.is_excel_as_pdf :
							if perm_obj.watermarking_file:
								da['names'] = perm_obj.watermarking_file.url.split("/")[-1]
					da['perm']['is_view_only'] = perm_obj.is_view_only
					da['perm']['is_no_access'] = perm_obj.is_no_access
					da['perm']['is_view_and_print'] = perm_obj.is_view_and_print
					da['perm']['is_view_and_print_and_download'] = perm_obj.is_view_and_print_and_download
					da['perm']['is_upload'] = perm_obj.is_upload
					da['perm']['is_watermarking'] = perm_obj.is_watermarking
					da['perm']['is_drm'] = perm_obj.is_drm
					da['perm']['is_editor'] = perm_obj.is_editor
					da['perm']['is_access'] = perm_obj.is_access
					print(perm_obj.is_access,"perm_obj.is_access")
					da['perm']['is_shortcut'] = perm_obj.is_shortcut
			except:
				da['perm']['is_view_only'] = False
				da['perm']['is_no_access'] = True
				da['perm']['is_access'] = False
				da['perm']['is_view_and_print'] = False
				da['perm']['is_view_and_print_and_download'] = False
				da['perm']['is_upload'] = False
				da['perm']['is_watermarking'] = False
				da['perm']['is_drm'] = False
				da['perm']['is_editor'] = False
				da['perm']['is_shortcut'] = False
			# else:
			#     da['perm']['is_view_only'] = True
			#     da['perm']['is_no_access'] = False
			#     da['perm']['is_access'] = True
			#     da['perm']['is_view_and_print'] = True
			#     da['perm']['is_view_and_print_and_download'] = True
			#     da['perm']['is_upload'] = True
			#     da['perm']['is_watermarking'] = True
			#     da['perm']['is_drm'] = True
			#     da['perm']['is_editor'] = True
		return Response(sorted(data, key = lambda i: i['is_folder']) , status=status.HTTP_201_CREATED)

class GetAllRootFolders(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		# #print("userrrrrr", user.is_admin, user)
		from_date = request.GET.get('from_date')
		to_date = request.GET.get('to_date')
		import datetime
		dataroom_folders = DataroomFolder.objects.filter(dataroom_id=pk, is_root_folder=True, is_deleted=False).order_by('index')
		dataroom_folder_serializer = DataroomFolderSerializer(dataroom_folders, many=True)
		data = dataroom_folder_serializer.data
		for da in data:

			# #print("daaaaaaaaaa", da)
			try:
				member = DataroomMembers.objects.filter(member_id=user.id, dataroom_id=da['dataroom'], is_deleted=False).first()
				if member.is_la_user == True or member.is_dataroom_admin == True: 
					da['is_view_only'] = True
					da['is_no_access'] = False
					da['is_access'] = True
					da['is_view_and_print'] = True
					da['is_view_and_print_and_download'] = True
					da['is_upload'] = True
					da['is_watermarking'] = True
					da['is_drm'] = True
					da['is_editor'] = True
					da['is_shortcut'] = True
				else:
					# #print("member--------", member.end_user_group.first().id, user.id, da['dataroom'])
					perm_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=da['id'],dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id)
					da['is_no_access'] = perm_obj.is_no_access
					da['is_access'] = perm_obj.is_access
					da['is_view_only'] = perm_obj.is_view_only
					da['is_view_and_print'] = perm_obj.is_view_and_print
					da['is_view_and_print_and_download'] = perm_obj.is_view_and_print_and_download
					da['is_upload'] = perm_obj.is_upload
					da['is_watermarking'] = perm_obj.is_watermarking
					da['is_drm'] = perm_obj.is_drm
					da['is_editor'] = perm_obj.is_editor
					da['is_shortcut'] = perm_obj.is_shortcut
			except:
				da['is_view_only'] = False
				da['is_no_access'] = True
				da['is_access'] = False
				da['is_view_and_print'] = False
				da['is_view_and_print_and_download'] = False
				da['is_upload'] = False
				da['is_watermarking'] = False
				da['is_drm'] = False
				da['is_editor'] = False
				da['is_shortcut'] = False
		return Response(data , status=status.HTTP_201_CREATED)


class searchFoldersFiles(APIView):
	authentication_classes = (TokenAuthentication,)
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		test = request.GET.get('test','')
		from_date = request.GET.get('from_date', '')
		to_date = request.GET.get('to_date', '')
		name = request.GET.get('name', '')
		file_type = request.GET.get('type', '')
		uploaded_by = request.GET.get('uploaded_by', '')
		# #print(from_date, to_date, name, file_type, uploaded_by, test)
		import datetime
		if to_date == '':
			todays_date = ''
		else:
			todays_date = datetime.datetime.strftime(datetime.datetime.strptime( to_date, '%Y-%m-%d'),'%Y-%m-%d 23:59:59+05:30')
		if from_date == '':
			first_date = ''
		else:
			first_date = datetime.datetime.strftime(datetime.datetime.strptime( from_date, '%Y-%m-%d'),'%Y-%m-%d 00:00:00+05:30')
		if test != '':
			q_list = [Q(uploaded_by__last_name__icontains=uploaded_by), Q(uploaded_by__first_name__icontains=uploaded_by), Q(name__icontains=test)]
		elif name != '':
			q_list = [Q(uploaded_by__last_name__icontains=uploaded_by), Q(uploaded_by__first_name__icontains=uploaded_by), Q(name__icontains=name)]
		elif file_type != '':
			q_list = [Q(uploaded_by__last_name__icontains=uploaded_by), Q(uploaded_by__first_name__icontains=uploaded_by), Q(name__icontains=file_type)] 
		else:
			q_list = [Q(uploaded_by__last_name__icontains=uploaded_by), Q(uploaded_by__first_name__icontains=uploaded_by), Q(name__icontains=test), Q(name__icontains=name), Q(name__icontains=file_type)]
		if from_date == '' and to_date == '':
			dataroom_folders = DataroomFolder.objects.filter(dataroom_id = pk, is_folder=True, is_deleted=False).filter(reduce(operator.or_, q_list)).order_by('index')
		elif from_date == '' and to_date != '':
			dataroom_folders = DataroomFolder.objects.filter(dataroom_id = pk, is_folder=True, is_deleted=False,  created_date__lte=todays_date).filter(reduce(operator.or_, q_list)).order_by('index')
		elif from_date != '' and to_date == '':
			dataroom_folders = DataroomFolder.objects.filter(dataroom_id = pk, is_folder=True, is_deleted=False,  created_date__gte=first_date).filter(reduce(operator.or_, q_list)).order_by('index')
		else:
			dataroom_folders = DataroomFolder.objects.filter(dataroom_id = pk, is_folder=True, is_deleted=False,  created_date__gte=first_date, created_date__lte=todays_date).filter(reduce(operator.or_, q_list)).order_by('index')
		# dataroom_folders = DataroomFolder.objects.filter(dataroom_id=pk,name__icontains=test, is_folder=True, is_deleted=False).order_by('index')
		dataroom_folder_serializer = DataroomFolderSerializer(dataroom_folders, many=True)
		folders_data =  dataroom_folder_serializer.data
		from . import utils
		for da in folders_data:
			da['index'] = utils.getIndexes(da)
			da['names'] = da['name']
			parent_list = []
			folder_id =da['parent_folder']
			get_root_folder = da['is_root_folder']
			while (get_root_folder == False):
			   if da['is_root_folder'] == False:
				   # #print("folderrrrr", folder_id)
				   folder_obj = DataroomFolder.objects.get(id = int(folder_id))
				   folder_serializer = DataroomFolderSerializer(folder_obj, many=False)
				   parent_list.append(folder_serializer.data)
				   if folder_obj.parent_folder_id != None:
					   folder_id = folder_obj.parent_folder_id 
				   else:
					   get_root_folder = True
				   dataroom_folder = folder_obj
			   else:
				   get_root_folder = True
			parent_list.reverse()
			da['parent_folders'] = parent_list
			try:
				member = DataroomMembers.objects.filter(member_id=user.id, dataroom_id=da['dataroom'], is_deleted=False).first()
				if member.is_la_user == True or member.is_dataroom_admin == True: 
					da['is_view_only'] = True
					da['is_no_access'] = False
					da['is_access'] = True
					da['is_view_and_print'] = True
					da['is_view_and_print_and_download'] = True
					da['is_upload'] = True
					da['is_watermarking'] = True
					da['is_drm'] = True
					da['is_editor'] = True
					da['is_shortcut'] = True
				else:
					# #print("member--------", member.end_user_group.first().id, user.id, da['dataroom'])
					perm_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=da['id'],dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id)
					da['is_no_access'] = perm_obj.is_no_access
					da['is_access'] = perm_obj.is_access
					da['is_view_only'] = perm_obj.is_view_only
					da['is_view_and_print'] = perm_obj.is_view_and_print
					da['is_view_and_print_and_download'] = perm_obj.is_view_and_print_and_download
					da['is_upload'] = perm_obj.is_upload
					da['is_watermarking'] = perm_obj.is_watermarking
					da['is_drm'] = perm_obj.is_drm
					da['is_editor'] = perm_obj.is_editor
					da['is_shortcut'] = perm_obj.is_shortcut
			except:
				da['is_view_only'] = False
				da['is_no_access'] = True
				da['is_access'] = False
				da['is_view_and_print'] = False
				da['is_view_and_print_and_download'] = False
				da['is_upload'] = False
				da['is_watermarking'] = False
				da['is_drm'] = False
				da['is_editor'] = False
				da['is_shortcut'] = False
		if from_date == '' and to_date == '':
			dataroom_files = DataroomFolder.objects.filter(dataroom_id = pk, is_folder=False, is_deleted=False).filter(reduce(operator.or_, q_list)).order_by('index')
		elif from_date == '' and to_date != '':
			dataroom_files = DataroomFolder.objects.filter(dataroom_id = pk, is_folder=False, is_deleted=False,  created_date__lte=todays_date).filter(reduce(operator.or_, q_list)).order_by('index')
		elif from_date != '' and to_date == '':
			dataroom_files = DataroomFolder.objects.filter(dataroom_id = pk, is_folder=False, is_deleted=False,  created_date__gte=first_date).filter(reduce(operator.or_, q_list)).order_by('index')
		else:
			dataroom_files = DataroomFolder.objects.filter(dataroom_id = pk, is_folder=False, is_deleted=False,  created_date__gte=first_date, created_date__lte=todays_date).filter(reduce(operator.or_, q_list)).order_by('index')
		
		# dataroom_files = DataroomFolder.objects.filter(dataroom_id=pk,name__icontains=test, is_folder=False, is_deleted=False).order_by('index')
		dataroom_files_serializer = DataroomFolderSerializer(dataroom_files, many=True)
		files_data =  dataroom_files_serializer.data
		for da in files_data:
			da['index'] = utils.getIndexes(da)
			da['names'] = da['name']
			parent_list = []
			folder_id =da['parent_folder']
			get_root_folder = da['is_root_folder']
			# #print("folderrrrr", folder_id)
			while (get_root_folder == False):
			   if da['is_root_folder'] == False:
				   folder_obj = DataroomFolder.objects.get(id = folder_id)
				   folder_serializer = DataroomFolderSerializer(folder_obj, many=False)
				   parent_list.append(folder_serializer.data)
				   if folder_obj.parent_folder_id != None:
					   folder_id = folder_obj.parent_folder_id 
				   else:
					   get_root_folder = True
				   dataroom_folder = folder_obj
			   else:
				   get_root_folder = True
			parent_list.reverse()
			da['parent_folders'] = parent_list
			try:
				member = DataroomMembers.objects.filter(member_id=user.id, dataroom_id=da['dataroom'], is_deleted=False).first()
				if member.is_la_user == True or member.is_dataroom_admin == True: 
					da['is_view_only'] = True
					da['is_no_access'] = False
					da['is_access'] = True
					da['is_view_and_print'] = True
					da['is_view_and_print_and_download'] = True
					da['is_upload'] = True
					da['is_watermarking'] = True
					da['is_drm'] = True
					da['is_editor'] = True
					da['is_shortcut'] = True
				else:
					# #print("member--------", member.end_user_group.first().id, user.id, da['dataroom'])
					perm_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=da['id'],dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id)
					grp_perm = DataroomGroupPermission.objects.filter(dataroom_groups_id=member.end_user_group.first().id, dataroom_id=da['dataroom']).first()
					# #print(grp_perm.id)
					if grp_perm.is_watermarking:
						if grp_perm.is_doc_as_pdf or grp_perm.is_excel_as_pdf :
							if perm_obj.watermarking_file:
								da['names'] = perm_obj.watermarking_file.url.split("/")[-1]
					da['is_no_access'] = perm_obj.is_no_access
					da['is_access'] = perm_obj.is_access
					da['is_view_only'] = perm_obj.is_view_only
					da['is_view_and_print'] = perm_obj.is_view_and_print
					da['is_view_and_print_and_download'] = perm_obj.is_view_and_print_and_download
					da['is_upload'] = perm_obj.is_upload
					da['is_watermarking'] = perm_obj.is_watermarking
					da['is_drm'] = perm_obj.is_drm
					da['is_editor'] = perm_obj.is_editor
					da['is_shortcut'] = perm_obj.is_shortcut
			except:
				da['is_view_only'] = False
				da['is_no_access'] = True
				da['is_access'] = False
				da['is_view_and_print'] = False
				da['is_view_and_print_and_download'] = False
				da['is_upload'] = False
				da['is_watermarking'] = False
				da['is_drm'] = False
				da['is_editor'] = False
				da['is_shortcut'] = False
		# #print("filessssssssssssssss", files_data)
		return Response({'folders':folders_data, 'files':files_data}, status=status.HTTP_201_CREATED)

# get all root folders for users and permission

#GetAllRootFoldersForUsersAndPermission
class GetAllRootFoldersForUsersAndPermission(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		dataroom_folders = DataroomFolder.objects.filter(dataroom_id=pk, is_root_folder=True, is_deleted=False).order_by('index')
		dataroom_folder_serializer = DataroomFolderSerializer(dataroom_folders, many=True)
		dataroom_folders = dataroom_folder_serializer.data[:]
		# print ("dataroom folders data", dataroom_folders)
		for folder in dataroom_folders:
			folder['is_no_access'] = False
			folder['is_view_only'] = False
			folder['is_view_and_print'] = False
			folder['is_upload'] = False
			folder['is_watermarking'] = False
			folder['is_drm'] = False
			folder['is_editor'] = False
			folder['is_shortcut'] = False

		# print ("folder data", dataroom_folders)
		
		return Response(dataroom_folders , status=status.HTTP_201_CREATED)

class GetFileInfo(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def post(self, request, pk, format=None):
		user = request.user
		data = request.data
		# print ("data is ", data)
		dataroom_folders = DataroomFolder.objects.get(dataroom_id=pk, dataroom_folder_uuid=data.get('file_id'))
		dataroom_folder_serializer = DataroomFolderSerializer(dataroom_folders, many=False)
		data = dataroom_folder_serializer.data
		from . import utils
		data['index'] = utils.getIndexes(data)
		return Response(data, status=status.HTTP_201_CREATED)

class GetSingleFolderDetail(APIView):
	authentication_classes = ( TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user_obj = request.user
		dataroom_folder = DataroomFolder.objects.get(pk=pk)
		dataroom_folder_serializer = DataroomFolderSerializer(dataroom_folder, many=False)
		data = dataroom_folder_serializer.data
		user = User.objects.get(id=data['user'])
		serializer = UserSerializer(user, many=False)
		data['user'] = serializer.data
		from . import utils
		data['index'] = utils.getIndexes(data)
		data['viewed_count'] = FolderView.objects.filter(dataroom_id = dataroom_folder.dataroom_id,folder_id =dataroom_folder.id).count()
		data['download_count'] = FolderDownload.objects.filter(dataroom_id = dataroom_folder.dataroom_id,folder_id =dataroom_folder.id).count()
		data['print_count'] = FolderPrint.objects.filter(dataroom_id = dataroom_folder.dataroom_id,folder_id =dataroom_folder.id).count()
		parent_list = []
		data['parent_folders'] = []
		get_root_folder = False
		folder_id =dataroom_folder.parent_folder_id
		while (get_root_folder == False):
		   if dataroom_folder.is_root_folder == False:
			   folder_obj = DataroomFolder.objects.get(id = folder_id)
			   folder_serializer = DataroomFolderSerializer(folder_obj, many=False)
			   parent_list.append(folder_serializer.data)
			   folder_id = folder_obj.parent_folder_id
			   dataroom_folder = folder_obj
		   else:
			   get_root_folder = True
		data['parent_folders'] = parent_list

		try:
			member = DataroomMembers.objects.filter(member_id=user_obj.id, dataroom_id=dataroom_folder.dataroom_id, is_deleted=False).first()
			if member.is_la_user == True or member.is_dataroom_admin == True: 
				data['perm'] = {}
				data['perm']['is_view_only'] = True
				data['perm']['is_no_access'] = False
				data['perm']['is_access'] = True
				data['perm']['is_view_and_print'] = True
				data['perm']['is_view_and_print_and_download'] = True
				data['perm']['is_upload'] = True
				data['perm']['is_watermarking'] = True
				data['perm']['is_drm'] = True
				data['perm']['is_editor'] = True
				data['perm']['is_shortcut'] = True
			else:
				perm_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=pk,dataroom_id=dataroom_folder.dataroom_id, dataroom_groups_id=member.end_user_group.first().id)
				perm_serializer = DataroomGroupFolderSpecificPermissionsSerializer(perm_obj, many=False)
				data['perm'] = perm_serializer.data
		except:
		   data['perm'] = {}
		return Response(data, status=status.HTTP_201_CREATED)



class CreateNewSubfolder(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def assign_index(self, data , user):
		index = 0
		try :
			# print ("inside create new folder", data)
			max_index = DataroomFolder.objects.filter(dataroom_id=data.get('dataroom_id'), parent_folder_id=data.get('folder_id'), is_root_folder=False, is_folder=True, is_deleted=False).aggregate(Max('index'))
			index = max_index.get('index__max')
			# print ('max index now is', index)
			if index is None:
				index = 0
		except :
			pass
		# print ("index after +1 is", index+1)
		return index+1

	def get_single_folder(self, pk):
		dataroom_folder = DataroomFolder.objects.get(pk=pk)
		return dataroom_folder

	def post(self, request, format=None):
		data = request.data
		user = request.user
		folder_obj = DataroomFolder.objects.filter(name=data.get('name'),parent_folder_id=int(data.get('folder_id')),is_folder=True, is_deleted =False).first()
		if folder_obj == None:
			data["user"] = user.id
			data["dataroom"] = data.get('dataroom_id')
			data["index"] = self.assign_index(data, user)
			data["last_updated_user"] = user.id
			data["is_root_folder"] = False
			data["is_folder"] = True
			data["parent_folder"] = int(data.get('folder_id'))#self.get_single_folder(int(data.get('folder_id')))
			data["deleted_by"] = user.id
			dataroom_folder_serializer = DataroomFolderSerializer(data = data)
			# print ("datarooom folder serializer ", dataroom_folder_serializer.is_valid())
			# print ("datarooom folder serializer ", dataroom_folder_serializer.errors)
			if dataroom_folder_serializer.is_valid():
				dataroom_folder_serializer.save()
				return Response(dataroom_folder_serializer.data, status=status.HTTP_201_CREATED)
			return Response(dataroom_folder_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
		else:
			return Response({'msg':'Given Name folder Already Exist in this folder'}, status=status.HTTP_201_CREATED)
		


class CreateNewRootFolder(APIView):
	"""docstring for DataroomOverviewAll"""
	authentication_classes = (TokenAuthentication,)
	permission_classes = (IsAuthenticated,)
	
	def assign_index(self, data , user):
		index = 0
		try :
			# print ("create new root folder data", data)
			max_index = DataroomFolder.objects.filter(dataroom_id=data.get('dataroom_id'), is_root_folder=True, is_folder=True, is_deleted=False).aggregate(Max('index'))
			index = max_index.get('index__max')
			# print ("max index is", index)
			if index is None:
				index = 0
		except :
			pass
		# print ("index +1 is", index)
		return index+1

	def post(self, request, format=None):
		data = request.data
		user = request.user
		# #print("data", data)
		folder_obj = DataroomFolder.objects.filter(name=data.get('name'),is_root_folder=True, dataroom_id=int(data.get('dataroom_id')), is_deleted =False).first()
		if folder_obj == None:
			data["user"] = user.id
			data["dataroom"] = data.get('dataroom_id')
			data["index"] = self.assign_index(data, user)
			data["is_folder"] = True
			data["last_updated_user"] = user.id
			data["is_root_folder"] = True
			data["is_root_folder"] = True
			dataroom_folder_serializer = DataroomFolderSerializer(data = data)
			if dataroom_folder_serializer.is_valid():
				dataroom_folder_serializer.save()
				return Response(dataroom_folder_serializer.data, status=status.HTTP_201_CREATED)
			return Response(dataroom_folder_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
		else:
			return Response({'msg':'Folder Already Exist'}, status=status.HTTP_201_CREATED)


class UpdateFolderDetails(APIView):
	"""
	Retrieve, update or delete a snippet instance.
	"""
	authentication_classes = (TokenAuthentication,)
	permission_classes = (IsAuthenticated,)
	
	def get_object(self, pk, request):
		try:
			user = request.user
			return DataroomFolder.objects.get( pk=pk)
		except DataroomFolder.DoesNotExist:
			raise Http404

	def get(self, request, pk, format=None):
		folder = self.get_object(pk, request)
		serializer = DataroomFolderSerializer(folder)
		return Response(serializer.data)

	def put(self, request, pk, format=None):
		folder = self.get_object(pk, request)
		data = request.data
		serializer = DataroomFolderSerializer(folder, data=data)
		if serializer.is_valid():
			serializer.save()
			return Response(serializer.data, status=status.HTTP_201_CREATED)
		return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

	def delete(self, request, pk, format=None):
		folder = self.get_object(pk, request)
		folder.is_deleted = True
		folder.save()
		return Response(status=status.HTTP_204_NO_CONTENT)

class UpdateFileDetails(APIView):
	"""
	Retrieve, update or delete a snippet instance.
	"""
	authentication_classes = (TokenAuthentication,)
	permission_classes = (IsAuthenticated,)
 

	def get_object(self, pk, request):
		try:
			user = request.user
			return Dataroom.objects.get(pk=pk)
		except Dataroom.DoesNotExist:
			raise Http404

	def get(self, request, pk, format=None):
		dataroom = self.get_object(pk, request)
		serializer = DataroomFilesSerializer(dataroom)
		return Response(serializer.data)

	def put(self, request, pk, format=None):
		dataroom = self.get_object(pk, request)
		data = request.data
		serializer = DataroomFilesSerializer(dataroom, data=data)
		if serializer.is_valid():
			serializer.save()
			return Response(serializer.data)
		return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

	def delete(self, request, pk, format=None):
		dataroom = self.get_object(pk, request)
		dataroom.delete()
		return Response(status=status.HTTP_204_NO_CONTENT)

import sys
import subprocess
import re


def convert_to(folder, source):
	args = [libreoffice_exec(), '--headless', '--convert-to', 'pdf', '--outdir', folder, source]

	process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
	filename = re.search('-> (.*?) using filter', process.stdout.decode())

	return filename.group(1)


def libreoffice_exec():
	# TODO: Provide support for more platforms
	if sys.platform == 'darwin':
		return '/usr/bin/soffice'
	return 'libreoffice'


class uploadFileInsideFolder(APIView):

	authentication_classes = (TokenAuthentication,)
	permission_classes = (IsAuthenticated,)
	#parser_classes = (FileUploadParser,)

	def assign_index(self, data , user, pk):
		index = 0
		try :
			# print ("data under assign_index is:", data)
			# print ("priamry key is", pk)
			max_index = DataroomFolder.objects.filter(parent_folder_id=pk, dataroom_id=data, is_root_folder=False, is_folder=False, is_deleted=False).aggregate(Max('index'))
			# print ("max indx is", max_index)
			index = max_index.get('index__max')
			# print ("index is", index)
			if index is None:
				index = 0
		except :
			pass
		# print ("new index is", index+1)
		return index+1

	def assign_index_folder(self, data , user, pk):
		index = 0
		try :
			# print ("inside create new folder", data)
			if int(pk) == 0:
				max_index = DataroomFolder.objects.filter(dataroom_id=data, is_root_folder=True, is_folder=True, is_deleted=False).aggregate(Max('index'))
			else:
				max_index = DataroomFolder.objects.filter(dataroom_id=data, parent_folder_id=pk, is_root_folder=False, is_folder=True, is_deleted=False).aggregate(Max('index'))
			index = max_index.get('index__max')
			# print ('max index now is', index)
			if index is None:
				index = 0
		except :
			pass
		# print ("index after +1 is", index+1)
		return index+1

	def create_folder(self, name, pk, dataroomId, user):
		data = {}
		# #print("pllllllllll", type(pk))
		if int(pk) == 0:
			data['is_root_folder'] = True
			data["parent_folder"] = None
		else:
			data['is_root_folder'] = False
			data["parent_folder"] = pk
		data["user"] = user.id
		data["dataroom"] = dataroomId
		data["index"] =  self.assign_index_folder(dataroomId, user, pk)
		data["is_folder"] = True
		data["last_updated_user"] = user.id
		data["name"] = name

		dataroom_folder_serializer = DataroomFolderSerializer(data = data)
		# #print("dataroom_folder_serializer", dataroom_folder_serializer.is_valid())
		# #print("dataroom_folder_serializer", dataroom_folder_serializer.errors)
		if dataroom_folder_serializer.is_valid():
			dataroom_folder_serializer.save()
		# folder = DataroomFolder()
		# folder.name = name
		# folder.is_folder = True
		# folder.parent_folder_id = pk
		# folder.dataroom_id = dataroomId
		# folder.index = self.assign_index_folder(dataroomId, user, pk)
		# folder.last_updated_user_id = user.id
		# folder.user_id = user.id
		# folder.is_root_folder = False
		# folder.save()
		folder = DataroomFolder.objects.get(id=dataroom_folder_serializer.data.get("id"))
		return folder




	def post(self, request, pk, format=None):
		user = request.user
		print("in post method")
		data = {}
		datas = request.data
		is_notify = request.GET.get("notify")
		dataroom_filess = request.FILES.getlist('file')
		print(len(dataroom_filess))
		print(dataroom_filess,"data room files")
		dataroomId = datas.get('dataroomId')
		print("datas====>",datas)

		if len(dataroom_filess) > 0:
			print("in more than zero condition")
			for i, dataroom_files in enumerate(dataroom_filess):
				print(i ,"in dtaroom")
				# #print("-----", i, dataroom_files.__dict__)
				if datas.getlist('folder')[i] == 'Yes':
					path = datas.getlist('path')[i].split("/")
					#print("path ======>",path)
					del path[-1]
					parent_id = pk
					# #print("parent_id", parent_id)
					for pa in path:
						# #print("paaaaaaaa", pa, dataroomId, parent_id)
						if int(parent_id) == 0:
							da = DataroomFolder.objects.filter(name = pa, dataroom_id = dataroomId, is_deleted=False, is_root_folder=True)
						else:
							da = DataroomFolder.objects.filter(name = pa,dataroom_id = dataroomId, parent_folder_id=parent_id, is_deleted= False, is_folder=True)
						# #print("daaaaaaaa", da, len(da))
						if len(da) > 0:
							parent_id = da.first().id
						else:
							folder = self.create_folder(pa, parent_id, dataroomId, user)
							parent_id = folder.id
				else:
					print("m in else now")
					parent_id = pk
					print(parent_id,"parent_id")
				folder_obj = DataroomFolder.objects.filter(name = dataroom_files.name, parent_folder_id=parent_id).first()
				print(folder_obj,"folder_obj")
				if folder_obj == None:
					data['name'] = dataroom_files.name
					data['parent_path'] = dataroom_files.name
					data['is_root_folder'] = False
					data['is_folder'] = False
					data['is_infected'] = False
					data['file_content_type'] = dataroom_files.content_type
					data['file_size'] = dataroom_files._size
					#print("size ====>",dataroom_files._size)
					data['parent_folder'] = parent_id

					data['pages'] = 1
					data['index'] = self.assign_index(dataroomId, user, parent_id)
					data['version'] = 0
					data['path'] = dataroom_files
					data['user'] = user.id
					data['last_updated_user'] = user.id
					data['dataroom'] = dataroomId
					dataroom_folder_serializer = DataroomFolderSerializer(data = data)
					# print ("dataroom_folder_serializer", dataroom_folder_serializer.is_valid())
					# print ("dataroom_folder_errors", dataroom_folder_serializer.errors)
					dataroom_obj = Dataroom.objects.filter(id=dataroomId).first()
					all_folders_size = DataroomFolder.objects.filter(dataroom_id=dataroomId).aggregate(Sum('file_size')).get('file_size__sum')
					# #print("dataroom filesss Size", dataroom_files._size/1000, all_folders_size/1000, dataroom_obj.dataroom_storage_allocated*1000000)
					dataroom_storage = dataroom_obj.dataroom_storage_allocated*1000
					file_size = (dataroom_files._size) + (all_folders_size if all_folders_size !=None else 0)
					file_size = file_size/1000/1000
					# #print("updated file size ===>",file_size)
					# time.sleep(10)
					# #print("Test")
					# #print("Print", dataroom_files._size,all_folders_size)
					# file_size = 100
					# #print("dataroom filesss Size", dataroom_storage,Decimal(file_size))
					# #print(dataroom_storage)
					# #print(dataroom_files._size/1024)
					# #print(all_folders_size)
					# #print(file_size)
					percent  = (Decimal(file_size)/dataroom_storage)*100
					from emailer import utils as ut
					if percent >= Decimal(80) and percent < Decimal(100):
						ut.send_80_percent_mail(data, user, dataroom_obj)
					elif percent >= Decimal(100):
						ut.send_100_percent_mail(data, user, dataroom_obj)
					else:
						pass 
					# #print("Percent",percent)
					if dataroom_folder_serializer.is_valid():
						dataroom_folder_serializer.save()
						# dataroom files 
						print(dataroom_folder_serializer.data,"serializer data here")
						
						# save file 
						dataroom_file = DataroomFolder.objects.get(id=dataroom_folder_serializer.data.get('id'))

						pn=1
						if str(dataroom_folder_serializer.data.get('name')).split('.')[1]=='xlsx' or str(dataroom_folder_serializer.data.get('name')).split('.')[1]=='csv' or str(dataroom_folder_serializer.data.get('name')).split('.')[1]=='xls':
							print("in xlsx")
							import os
							filepath=dataroom_folder_serializer.data['path']
							print(filepath,"filepath")
							pathsplit=filepath.split('/')
							print(pathsplit,"pathsplit")
							from azure.storage.blob import BlockBlobService, PublicAccess,ContentSettings
							pathu=pathsplit[-2]+"/"+pathsplit[-1]
							block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
							container_name='docullycontainer'
							block_blob_service.get_blob_to_path(container_name, pathu, pathsplit[-1])
							dataroom_file.pages = int()
							dataroom_file.save()
							print(pathu,"pathu")
							with open(pathsplit[-1], 'rb') as ifh:
								read_data=ifh.read()
								with open("/home/cdms_backend/cdms2/DocFile_server/"+pathsplit[-1], 'wb') as fileee:
									fileee.write(read_data)
									print("yes done pptx")
							convert_to('/home/cdms_backend/cdms2/server_DocFile/',  "/home/cdms_backend/cdms2/DocFile_server/"+pathsplit[-1])
							fileext=pathsplit[-1].split('.')
							print(fileext,"fileext")
							blob_na=pathsplit[4]+"/"+fileext[0]+".pdf"
							print(blob_na,"blob_na")
							block_blob_service.create_blob_from_path(
							container_name,
							blob_na,
							"/home/cdms_backend/cdms2/server_DocFile/"+fileext[0]+".pdf",content_settings=ContentSettings(content_type='application/pdf'))
							print("file uploaded in server pptx")
							print("/home/cdms_backend/cdms2/server_DocFile/"+fileext[0]+".pdf")
							os.remove('/home/cdms_backend/cdms2/DocFile_server/'+pathsplit[-1])
							os.remove('/home/cdms_backend/cdms2/server_DocFile/'+fileext[0]+".pdf")
							print("file removed in folder pptx")
						#print(str(dataroom_folder_serializer.data.get('name')).split('.')[1])
						if str(dataroom_folder_serializer.data.get('name')).split('.')[1]=='pptx' or str(dataroom_folder_serializer.data.get('name')).split('.')[1]=='ppt':
							print("in pptx")
							import os
							filepath=dataroom_folder_serializer.data['path']
							print(filepath,"filepath")
							pathsplit=filepath.split('/')
							print(pathsplit,"pathsplit")
							from azure.storage.blob import BlockBlobService, PublicAccess,ContentSettings
							pathu=pathsplit[-2]+"/"+pathsplit[-1]
							block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
							container_name='docullycontainer'
							block_blob_service.get_blob_to_path(container_name, pathu, pathsplit[-1])
							dataroom_file.pages = int()
							dataroom_file.save()
							print(pathu,"pathu")
							with open(pathsplit[-1], 'rb') as ifh:
								read_data=ifh.read()
								with open("/home/cdms_backend/cdms2/DocFile_server/"+pathsplit[-1], 'wb') as fileee:
									fileee.write(read_data)
									print("yes done pptx")
							convert_to('/home/cdms_backend/cdms2/server_DocFile/',  "/home/cdms_backend/cdms2/DocFile_server/"+pathsplit[-1])
							fileext=pathsplit[-1].split('.')
							print(fileext,"fileext")
							blob_na=pathsplit[4]+"/"+fileext[0]+".pdf"
							print(blob_na,"blob_na")
							block_blob_service.create_blob_from_path(
							container_name,
							blob_na,
							"/home/cdms_backend/cdms2/server_DocFile/"+fileext[0]+".pdf",content_settings=ContentSettings(content_type='application/pdf'))
							print("file uploaded in server pptx")
							print("/home/cdms_backend/cdms2/server_DocFile/"+fileext[0]+".pdf")
							os.remove('/home/cdms_backend/cdms2/DocFile_server/'+pathsplit[-1])
							os.remove('/home/cdms_backend/cdms2/server_DocFile/'+fileext[0]+".pdf")
							print("file removed in folder pptx")
						if str(dataroom_folder_serializer.data.get('name')).split('.')[1]=='docx' or str(dataroom_folder_serializer.data.get('name')).split('.')[1]=='doc' or str(dataroom_folder_serializer.data.get('name')).split('.')[1]=='odt':
							#print(123)
							print("if1")
							from docx import Document
							import os
							from azure.storage.blob import BlockBlobService, PublicAccess
							from docx import Document
							
							fn='./'+str(dataroom_folder_serializer.data.get('path')).replace('%20',' ').replace(' ','\ ')
							#print(fn)
							filepath=dataroom_folder_serializer.data['path']
							print(filepath,"filepath")
							pathsplit=filepath.split('/')
							print(pathsplit,"pathsplit")
							from azure.storage.blob import BlockBlobService, PublicAccess,ContentSettings
							pathu=pathsplit[-2]+"/"+pathsplit[-1]
							block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
							container_name='docullycontainer'
							block_blob_service.get_blob_to_path(container_name, pathu, pathsplit[-1])
							dataroom_file.pages = int()
							dataroom_file.save()
							print(pathu,"pathu")
							with open(pathsplit[-1], 'rb') as ifh:
								read_data=ifh.read()
								with open("/home/cdms_backend/cdms2/DocFile_server/"+pathsplit[-1], 'wb') as fileee:
									fileee.write(read_data)
									print("yes done")
							convert_to('/home/cdms_backend/cdms2/server_DocFile/',  "/home/cdms_backend/cdms2/DocFile_server/"+pathsplit[-1])
							fileext=pathsplit[-1].split('.')
							print(fileext,"fileext")
							blob_na=pathsplit[4]+"/"+fileext[0]+".pdf"
							print(blob_na,"blob_na")
							block_blob_service.create_blob_from_path(
							container_name,
							blob_na,
							"/home/cdms_backend/cdms2/server_DocFile/"+fileext[0]+".pdf",content_settings=ContentSettings(content_type='application/pdf'))
							print("file uploaded in server")
							print("/home/cdms_backend/cdms2/server_DocFile/"+fileext[0]+".pdf")
							os.remove('/home/cdms_backend/cdms2/DocFile_server/'+pathsplit[-1])
							os.remove('/home/cdms_backend/cdms2/server_DocFile/'+fileext[0]+".pdf")
							print("file removed in folder")
							# import PyPDF2

							# # fn='./'+str(dataroom_folder_serializer.data.get('path')).replace('%20',' ')
							# # fn = str(fn).split('.')[0]+'.pdf'
							# fn = './'+str(dataroom_folder_serializer.data.get('name')).replace(' ','_').split('.')[0]+'.pdf'
							# # #print(fn,3434)
							# # fn = str(fn).split('.')[0]+'.pdf'
							# #print("value of fn_docx ====>",fn)
							# import PyPDF2
							# filepath=dataroom_folder_serializer.data['path']
							# print(filepath,"filepath")
							# pathsplit=filepath.split('/')
							# print(pathsplit,"pathsplit")
							# from azure.storage.blob import BlockBlobService, PublicAccess
							# pathu=pathsplit[-2]+"/"+pathsplit[-1]
							# block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
							# container_name='docullycontainer'
							# block_blob_service.get_blob_to_path(container_name, pathu, pathsplit[-1])
							# dataroom_file.pages = int()
							# dataroom_file.save()
							# unzip -p pathsplit[-1] docProps/app.xml | grep -oP '(?<=\<Pages\>).*(?=\</Pages\>)'
						# 	with open(pathsplit[-1], 'rb') as fileee:
						# 		pdfReader = PyPDF2.PdfFileReader(fileee)
						# 		pn = pdfReader.numPages
						# dataroom_file.pages = pn
						# dataroom_file.save()
							# os.system('rm -r '+fn)

							# wordFileObj = open(fn, 'rb')
							# document = Document(wordFileObj)
							
							# import re
							# for p in document.paragraphs:
							#     r = re.match('Chapter \d+',p.text)
							#     if r:
							#         #print(r.group(),pn)
							#     for run in p.runs:
							#         if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
							#             pn+=1
							#             #print(pn)
						if str(dataroom_folder_serializer.data.get('name')).split('.')[1]=='pdf' or str(dataroom_folder_serializer.data.get('name')).split('.')[1]=='PDF':
							print("if2")
							import PyPDF2
							filepath=dataroom_folder_serializer.data['path']
							print(filepath,"filepath")
							pathsplit=filepath.split('/')
							print(pathsplit,"pathsplit")
							from azure.storage.blob import BlockBlobService, PublicAccess
							pathu=pathsplit[-2]+"/"+pathsplit[-1]
							block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
							container_name='docullycontainer'
							block_blob_service.get_blob_to_path(container_name, pathu, pathsplit[-1])
							with open(pathsplit[-1], 'rb') as fileee:
								pdfReader = PyPDF2.PdfFileReader(fileee)
								pn = pdfReader.numPages
						dataroom_file.pages = pn
						dataroom_file.save()

						# dataroom_file.path = dataroom_files
						# #print("printtttttt", dataroom_file.path)
						# dataroom_file.save()
						# file is successfully saved
						if is_notify == "Yes":
							from . import utils
							dr_overview = DataroomOverview.objects.filter(dataroom_id=dataroomId).first()
							if dr_overview.send_daily_email_updates:
								utils.send_notify_to_all_members_regarding_uploaded_file(dataroom_file)
				else:
					print("else")
					data['path'] = dataroom_files
					data['version'] = folder_obj.version + 1
					data['user'] = folder_obj.user.id
					data['dataroom'] = dataroomId
					data['last_updated_user'] = user.id
					serializer = DataroomFolderSerializer(folder_obj, data = data)
					# #print("serializer", serializer.is_valid())
					# #print("serializer", serializer.errors)
					if serializer.is_valid():
						serializer.save()
		else:
			print("else2")
			for path in datas.getlist('path'):
				# #print("pathhhhhh", path)
				pa = path.split("/")
				del pa[-1]
				# #print("paaaaaaaa", pa, path)
				parent_id = pk
				# #print("parent_id", parent_id)
				for p in pa:
					if int(parent_id) == 0:
						da = DataroomFolder.objects.filter(name = p, dataroom_id =dataroomId, is_deleted=False, is_root_folder=True)
					else:
						da = DataroomFolder.objects.filter(name = p,dataroom_id = dataroomId, parent_folder_id=parent_id, is_deleted= False, is_folder=True)
					# #print("daaaaa", len(da))
					if len(da) > 0:
						parent_id = da.first().id
					else:
						folder = self.create_folder(p, parent_id, dataroomId, user)
						parent_id = folder.id


		data = { 'message':"File successfully uploaded !"}
		return Response(data = data, status=status.HTTP_201_CREATED)
		# return Response(status=status.HTTP_400_BAD_REQUEST)


class UpdateFolderOrFileName(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def post(self, request, format=None):
		data = request.data
		user = request.user
		folder = data.get('folder')
		try:
			dataroom_folder = DataroomFolder.objects.get(id=folder.get('id'))
			dataroom_folder.name = folder.get('name')
			dataroom_folder.save()
			return Response({'msg':'Folder/File successfully updated !'}, status=status.HTTP_201_CREATED)
		except:
			# print ("model dows not exist")
			return Response({'msg':'Folder/File does not exist !'}, status=status.HTTP_400_BAD_REQUEST)


class DeleteAllSelectedFiles(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def post(self, request, format=None):
		data = request.data
		user = request.user
		folders = data.get('folder')
		folders_id = [folder.get('id') for folder in folders if folder.get('selected') == True]
		# print ("selected folders_id", folders_id)
		current_date = timezone.now()
		try:
			new_data = []
			for _id in folders_id:
				dataroom_folder = DataroomFolder.objects.get(id=_id)
				dataroom_folder.is_deleted =  True
				dataroom_folder.deleted_by = user
				dataroom_folder.deleted_by_date = current_date#add current date
				dataroom_folder.save()
				new_data = change_all_indexes(dataroom_folder)
			dataroom_folder_serializer = DataroomFolderSerializer(new_data, many=True)    
			return Response(dataroom_folder_serializer.data, status=status.HTTP_201_CREATED)
		except:
			# print ("model does not exist")
			return Response({'msg':'Folder/File successfully deleted !'}, status=status.HTTP_400_BAD_REQUEST)

class GetAllTrashedFilesAndFolders(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		data = request.data
		user = request.user
		count = 0
		deleted_bys = request.GET.get('deleted_by')
		q_list = [Q(deleted_by__last_name__icontains=deleted_bys), Q(deleted_by__first_name__icontains=deleted_bys)]
		dataroom_folder = DataroomFolder.objects.filter(dataroom_id=pk, is_deleted=True, user_id = user.id).filter(reduce(operator.or_, q_list))
		count = DataroomFolder.objects.filter(dataroom_id=pk, is_deleted=True, user_id = user.id).filter(reduce(operator.or_, q_list)).count()
		page = paginator.paginate_queryset(dataroom_folder, request)
		dataroom_folder_serializer = DataroomFolderSerializer(page, many=True)
		data = dataroom_folder_serializer.data
		for da in data:
			deleted_by = User.objects.get(id=da['deleted_by'])
			da['deleted_by'] = (UserSerializer(deleted_by, many=False)).data
		data.sort(key=lambda item:item['deleted_by_date'], reverse=True)
		return Response({'data':data, 'size':count})

class GetAllRecentDocument(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		# print ("get all recent document data")        
		data = request.data
		user = request.user
		from_date = request.GET.get('from_date')
		to_date = request.GET.get('to_date')
		if from_date == '' and to_date == '':
			dataroom_folder = DataroomFolder.objects.filter(dataroom_id=pk, is_deleted=False, is_folder=False).order_by('-created_date')
		else:   
			import datetime
			todays_date = datetime.datetime.strftime(datetime.datetime.strptime( to_date, '%Y-%m-%d'),'%Y-%m-%d 23:59:59+05:30')
			first_date = datetime.datetime.strftime(datetime.datetime.strptime( from_date, '%Y-%m-%d'),'%Y-%m-%d 00:00:00+05:30')

			#date_from = datetime.now() - datetime.timedelta(days=30)
			dataroom_folder = DataroomFolder.objects.filter(dataroom_id=pk, is_deleted=False, is_folder=False, created_date__gte=first_date, created_date__lte=todays_date).order_by('-created_date')
		dataroom_folder_serializer = DataroomFolderSerializer(dataroom_folder, many=True)
		data = dataroom_folder_serializer.data
		from . import utils
		for da in data:
			da['index'] = utils.getIndexes(da)
			da['parent_folders'] = DataroomFolder.objects.get(id = da['parent_folder']).name
			try:
			   perm = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=pk)
			   perm_serializer = DataroomGroupFolderSpecificPermissionsSerializer(perm, many=False)
			   da['perm'] = perm_serializer.data
			except:
			   da['perm'] = {}
			if user.is_superadmin == True or user.is_admin == True:
				da['perm']['is_view_only'] = True
				da['perm']['is_no_access'] = False
				da['perm']['is_access'] = True
				da['perm']['is_view_and_print'] = True
				da['perm']['is_view_and_print_and_download'] = True
				da['perm']['is_upload'] = True
				da['perm']['is_watermarking'] = True
				da['perm']['is_drm'] = True
				da['perm']['is_editor'] = True
		return Response(data)

class EditIndexOfRootFolder(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def update_indexing_logic(self, objects, old_index, index_to_be_updated):
		# print ("update indexing logic")
		for obj in objects:
			# print ("obj id", obj.index)
			if obj.index == old_index:
				try:
					obj.index = int(index_to_be_updated)
					obj.save()
					# print ("index is successfully updated")
				except:
					print ("error in saving index")
			else:
					obj.index += 1
					obj.save()
					print ("reindexed all indexes")

	def post(self, request, format=None):
		#pdb.set_trace()
		print ("inside edit index of root folder")
		data  = request.data
		user = request.user
		dataroomId = data.get('dataroomId')
		folder_data = data.get('folder')
		print ("folder data is", folder_data)
		index_to_be_updated = folder_data.get('index')
		id = folder_data.get('id')
		old_index = DataroomFolder.objects.get(id=folder_data.get('id'))
		old_index = old_index.index
		# print ("new index is", index_to_be_updated)
		# print ("old index is", old_index)

		
		if index_to_be_updated>0:
			if folder_data.get('is_folder') :
				if folder_data.get('is_root_folder'):
					#print(90)
					min_data_index = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=True, is_folder=True, is_deleted=False).aggregate(Min('index'))
					max_data_index = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=True, is_folder=True, is_deleted=False).aggregate(Max('index'))
					#print(max_data_index)
					min_data_index = min_data_index.get('index__min')
					max_data_index = max_data_index.get('index__max')

					dataroom_folder = DataroomFolder.objects.get(id=id, is_deleted=False, is_root_folder=True, is_folder=True)
					try:
						dataroom_folder_new = DataroomFolder.objects.get(dataroom_id = dataroomId, index=index_to_be_updated, is_deleted=False, is_root_folder=True, is_folder=True)
					except:
						get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_deleted=False, is_root_folder=True).order_by('index')
						dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)

						return Response({'msg':'Max '+''+str(max_data_index)+''+' is Allowed!', 'data': dataroom_folder_serializer.data}, status=status.HTTP_400_BAD_REQUEST)

					if (index_to_be_updated >=  min_data_index) and (index_to_be_updated <=max_data_index):
						temp_idx = dataroom_folder.index
						dataroom_folder.index = index_to_be_updated
						dataroom_folder_new.index = temp_idx
						dataroom_folder.save()
						dataroom_folder_new.save()

						get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_deleted=False, is_root_folder=True).order_by('index')
						dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)

						return Response({'msg':'Index is successfully updated !', 'data': dataroom_folder_serializer.data}, status=status.HTTP_201_CREATED)
					else:
						get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_deleted=False, is_root_folder=True).order_by('index')
						dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)

						return Response({'msg':'Index is successfully updated !', 'data': dataroom_folder_serializer.data}, status=status.HTTP_400_BAD_REQUEST)
				else:

					min_data_index = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=False, is_folder=True, is_deleted=False, parent_folder_id= folder_data.get('parent_folder')).aggregate(Min('index'))
					max_data_index = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=False, is_folder=True, is_deleted=False, parent_folder_id= folder_data.get('parent_folder')).aggregate(Max('index'))
					#print(max_data_index)
					min_data_index = min_data_index.get('index__min')
					max_data_index = max_data_index.get('index__max')

					dataroom_folder = DataroomFolder.objects.get(id=id, is_deleted=False, is_root_folder=False, is_folder=True, parent_folder_id= folder_data.get('parent_folder'))
					try:
						dataroom_folder_new = DataroomFolder.objects.get(dataroom_id = dataroomId, index=index_to_be_updated, is_deleted=False, is_root_folder=False, is_folder=True, parent_folder_id= folder_data.get('parent_folder'))
					except:
						get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, parent_folder_id=folder_data.get('parent_folder'), is_deleted=False, is_root_folder=False).order_by('index')
						dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)

						return Response({'msg':'Max '+''+str(max_data_index)+''+' is Allowed!', 'data': dataroom_folder_serializer.data}, status=status.HTTP_400_BAD_REQUEST)

					if (index_to_be_updated >=  min_data_index) and (index_to_be_updated <=max_data_index):
						temp_idx = dataroom_folder.index
						dataroom_folder.index = index_to_be_updated
						dataroom_folder_new.index = temp_idx
						dataroom_folder.save()
						dataroom_folder_new.save()

						get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, parent_folder_id=folder_data.get('parent_folder'), is_deleted=False, is_root_folder=False).order_by('index')
						dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)

						return Response({'msg':'Index is successfully updated !', 'data': dataroom_folder_serializer.data}, status=status.HTTP_201_CREATED)
					else:
						get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, parent_folder_id=folder_data.get('parent_folder'), is_deleted=False, is_root_folder=False).order_by('index') 
						dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)

						return Response({'msg':'Index is Not updated!', 'data': dataroom_folder_serializer.data}, status=status.HTTP_400_BAD_REQUEST)
			else:

				if folder_data.get('parent_folder'):
					min_data_index = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=False, is_folder=False, is_deleted=False, parent_folder_id= folder_data.get('parent_folder')).aggregate(Min('index'))
					max_data_index = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=False, is_folder=False, is_deleted=False, parent_folder_id= folder_data.get('parent_folder')).aggregate(Max('index'))
					#print(max_data_index)
					min_data_index = min_data_index.get('index__min')
					max_data_index = max_data_index.get('index__max')

					dataroom_file = DataroomFolder.objects.get(id=id, is_deleted=False, is_root_folder=False, is_folder=False, parent_folder_id= folder_data.get('parent_folder'))
					try:
						dataroom_file_new = DataroomFolder.objects.get(dataroom_id = dataroomId, index=index_to_be_updated, is_deleted=False, is_root_folder=False, is_folder=False, parent_folder_id= folder_data.get('parent_folder'))
					except:
						get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, parent_folder_id=folder_data.get('parent_folder'), is_deleted=False, is_root_folder=False).order_by('index')
						dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)

						return Response({'msg':'Max '+''+str(max_data_index)+''+' is Allowed!', 'data': dataroom_folder_serializer.data}, status=status.HTTP_400_BAD_REQUEST)

					if (index_to_be_updated >=  min_data_index) and (index_to_be_updated <=max_data_index):
						temp_idx = dataroom_file.index
						dataroom_file.index = index_to_be_updated
						dataroom_file_new.index = temp_idx
						dataroom_file.save()
						dataroom_file_new.save()

						get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, parent_folder_id=folder_data.get('parent_folder'), is_deleted=False, is_root_folder=False).order_by('index')
						dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)
						return Response({'msg':'Index is successfully updated !', 'data': dataroom_folder_serializer.data}, status=status.HTTP_201_CREATED)
					else:
						get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, parent_folder_id=folder_data.get('parent_folder'), is_deleted=False, is_root_folder=False).order_by('index')
						dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)

						return Response({'msg':'Index is Not updated!', 'data': dataroom_folder_serializer.data}, status=status.HTTP_400_BAD_REQUEST)
		else:
			if folder_data.get('parent_folder'):

				get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, parent_folder_id=folder_data.get('parent_folder'), is_deleted=False, is_root_folder=False).order_by('index')
				dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)
			else:
				get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_deleted=False, is_root_folder=True).order_by('index')
				dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)

			return Response({'msg':'Index should be greater than Zero!', 'data': dataroom_folder_serializer.data}, status=status.HTTP_400_BAD_REQUEST)
					# for i in dataroom_folder:
						# #print(i)
						# if i.path=='/media/f383d620-6f05-4f61-b6a6-8d6d8ebc074a':
						#     #print(1)
							# object.save()
					# dataroom_folder.path='/media/f383d620-6f05-4f61-b6a6-8d6d8ebc074a'
					# dataroom_folder.save()

					# #print(dataroom_folder.get('index'))
					# data = json.loads(dataroom_folder_serializer.data)
					# data_1 = json.loads(dataroom_folder_serializer_new.data)

					# try:
					#     obj.index = int(index_to_be_updated)
					#     obj.save()
					#     #print("Sucess")
					# except:
					#     print ("error in saving index")
					# #print(dataroom_folder_serializer.data)
					# #print(data)
					# if data:
					#     for obj in data:
					#         #print(obj)


		# if index_to_be_updated > 0:
		#     if folder_data.get('is_folder') :
		#         if folder_data.get('is_root_folder'):
		#             # print ("folder is root folder")
		#             min_data_index = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=True, is_folder=True, is_deleted=False).aggregate(Min('index'))
		#             max_data_index = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=True, is_folder=True, is_deleted=False).aggregate(Max('index'))
		#             min_data_index = min_data_index.get('index__min')
		#             max_data_index = max_data_index.get('index__max')
		#             # print ("min_data_index", min_data_index)
		#             # print ("max_data_index", max_data_index)
		#             # index should be greater than minimum amount and less than maximum amount.
		#             if (index_to_be_updated >=  min_data_index) and (index_to_be_updated <=max_data_index):
		#                 """ update index here """
		#                 get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=True, is_folder=True, is_deleted=False, index__lte=old_index, index__gte=index_to_be_updated)
		#                 """ iterate all objects here """
						
		#                 # update all indexes here
		#                 self.update_indexing_logic(get_all_objects, old_index, index_to_be_updated)
		#                 # edit expected index 
		#                 get_all_objects = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=True, is_folder=True, is_deleted=False).order_by('index')
			
		#                 dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)
		#                # print ("dataroom folder serializer data", dataroom_folder_serializer.data)
		#                 return Response({'msg':'Index is successfully updated !', 'data': dataroom_folder_serializer.data}, status=status.HTTP_201_CREATED)
		#             else:
		#                 return Response({'msg':'Index should be less than maximum index. '}, status=status.HTTP_400_BAD_REQUEST)
		#         else:
		#                 # check expected index should be  greater than min index and less than max index.  
		#                 min_data_index = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=False, is_folder=True, is_deleted=False, parent_folder_id= folder_data.get('parent_folder')).aggregate(Min('index'))
		#                 max_data_index = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=False, is_folder=True, is_deleted=False, parent_folder_id= folder_data.get('parent_folder')).aggregate(Max('index'))
		#                 min_data_index = min_data_index.get('index__min')
		#                 max_data_index = max_data_index.get('index__max')
		#                 # print ("min_data_index", min_data_index)
		#                 # print ("max_data_index", max_data_index)
		#                 # index should be greater than minimum amount and less than maximum amount. 
		#                 if (index_to_be_updated >=  min_data_index) and (index_to_be_updated <= max_data_index):
		#                     # update index here
		#                     get_all_objects = DataroomFolder.objects.filter(parent_folder_id=folder_data.get('parent_folder'), is_deleted=False, is_folder=True, is_root_folder=False, index__lte=old_index, index__gte=index_to_be_updated)#.update(index = F('index')-1)  
		#                     # update indexing logic
		#                     self.update_indexing_logic(get_all_objects, old_index, index_to_be_updated)
		#                     # get all objects
		#                     get_all_objects = DataroomFolder.objects.filter(parent_folder_id=folder_data.get('parent_folder'), is_deleted=False, is_folder=True, is_root_folder=False)  
		#                     dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)
		#                     return Response({'msg':'Index is successfully updated !', 'data': dataroom_folder_serializer.data}, status=status.HTTP_201_CREATED)
		#                 else:
		#                     return Response({'msg':'Index should be less than maximum index. '}, status=status.HTTP_400_BAD_REQUEST)
		#     else :
		#         # print ("folder data is file")
		#         min_data_index = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=False, is_folder=False, is_deleted=False, parent_folder_id= folder_data.get('parent_folder')).aggregate(Min('index'))
		#         max_data_index = DataroomFolder.objects.filter(dataroom_id = dataroomId, is_root_folder=False, is_folder=False, is_deleted=False, parent_folder_id= folder_data.get('parent_folder')).aggregate(Max('index'))
		#         min_data_index = min_data_index.get('index__min')
		#         max_data_index = max_data_index.get('index__max')
		#         # print ("min_data_index", min_data_index)
		#         # print ("max_data_index", max_data_index)
		#         # index should be greater than minimum amount and less than maximum amount. 
		#         if (index_to_be_updated >=  min_data_index) and (index_to_be_updated <= max_data_index):
		#             # update index here
		#             get_all_objects = DataroomFolder.objects.filter(parent_folder_id=folder_data.get('parent_folder'), is_deleted=False, is_folder=False, is_root_folder=False, index__lte=old_index, index__gte=index_to_be_updated)#.update(index = F('index')-1)  
		#             self.update_indexing_logic(get_all_objects, old_index, index_to_be_updated)
		#             get_all_objects = DataroomFolder.objects.filter(parent_folder_id=folder_data.get('parent_folder'), is_deleted=False, is_folder=False, is_root_folder=False)  
		#             dataroom_folder_serializer = DataroomFolderSerializer(get_all_objects, many=True)
		#             return Response({'msg':'Index is successfully updated !', 'data': dataroom_folder_serializer.data}, status=status.HTTP_201_CREATED)
		#         else:
		#             return Response({'msg':'Index should be less than maximum index. '}, status=status.HTTP_400_BAD_REQUEST)
		# else:   
		#     return Response({'msg':'Index should be less than maximum index. '}, status=status.HTTP_400_BAD_REQUEST)

class PermanantDeleteFolder(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def delete(self, request, pk, format=None):
		# print ("dataroom folder pk is ", pk)
		delete_folder = DataroomFolder.objects.get(id=pk)
		delete_folder.delete()
		return Response({'msg':'Folder / File successfully deleted !'}, status=status.HTTP_201_CREATED)

class RestoreFolderFiles(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		# print ("dataroom folder pk is ", pk)
		delete_folder = DataroomFolder.objects.get(id=pk)
		delete_folder.is_deleted = False
		if delete_folder.is_root_folder == True:
			# #print("It is root folder")
			max_index = DataroomFolder.objects.filter(dataroom_id=delete_folder.dataroom_id, is_root_folder=True, is_folder=True, is_deleted=False).aggregate(Max('index'))
		else:
			# #print("It is not root folder")
			if delete_folder.is_folder == True:
				# #print("But it is a folder")
				max_index = DataroomFolder.objects.filter(parent_folder_id=delete_folder.parent_folder_id, dataroom_id=delete_folder.dataroom_id, is_root_folder=False, is_folder=True, is_deleted=False).aggregate(Max('index'))
			else:
				# #print("But it is a file")
				max_index = DataroomFolder.objects.filter(parent_folder_id=delete_folder.parent_folder_id, dataroom_id=delete_folder.dataroom_id, is_root_folder=False, is_folder=False, is_deleted=False).aggregate(Max('index'))
		# print ("max indx is", max_index)
		index = max_index.get('index__max')
		# print ("index is", index)
		if index is None:
			index = 0
		delete_folder.index = index + 1
		delete_folder.save()
		return Response({'msg':'Folder / File successfully restored !'}, status=status.HTTP_201_CREATED)


	def post(self, request, pk, format=None):
		# print ("dataroom folder pk is ", pk)
		# data = json.dumps(request.data)
		# #print(request.data)
		# data = json.loads(request.data)
		# #print()
		if request.data:
			for i in request.data:
				# #print(i)
				delete_folder = DataroomFolder.objects.get(id=i['id'])
				delete_folder.is_deleted = False
				if delete_folder.is_root_folder == True:
					# #print("It is root folder")
					max_index = DataroomFolder.objects.filter(dataroom_id=delete_folder.dataroom_id, is_root_folder=True, is_folder=True, is_deleted=False).aggregate(Max('index'))
				else:
					# #print("It is not root folder")
					if delete_folder.is_folder == True:
						# #print("But it is a folder")
						max_index = DataroomFolder.objects.filter(parent_folder_id=delete_folder.parent_folder_id, dataroom_id=delete_folder.dataroom_id, is_root_folder=False, is_folder=True, is_deleted=False).aggregate(Max('index'))
					else:
						# #print("But it is a file")
						max_index = DataroomFolder.objects.filter(parent_folder_id=delete_folder.parent_folder_id, dataroom_id=delete_folder.dataroom_id, is_root_folder=False, is_folder=False, is_deleted=False).aggregate(Max('index'))
				# print ("max indx is", max_index)
				index = max_index.get('index__max')
				# #print(pk, delete_folder.parent_folder_id, index)
				# print ("index is", index)
				if index is None:
					index = 0
				delete_folder.index = index + 1
				delete_folder.save()
			return Response({'msg':'Folder / File successfully restored !'}, status=status.HTTP_201_CREATED)
		else:
			return Response({'msg':'Something Went Wrong!'}, status=status.HTTP_400_BAD_REQUEST)

class GetAllFilesandFolders(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		document = DataroomFolder.objects.filter(dataroom_id = pk, is_root_folder=True, is_deleted=False)
		group_id = request.GET.get('group_id')
		# #print("Group Id", group_id)
		data = []
		from . import utils
		for doc in document:
			docu = DataroomFolder.objects.get(id = doc.id)
			docu_serializer = DataroomFolderSerializer(docu)
			datas = utils.getIndexofFolder(docu_serializer.data)
			data.append(datas)
			docu1 = DataroomFolder.objects.filter(parent_folder = doc.id, is_deleted=False)
			if len(docu1) > 0:
				datas = []
				data.extend(utils.get_under_file(docu1,datas))
		for da in data:
			da['perm'] = {}
			# if user.is_superadmin == False or user.is_admin or False:
			try:
				if group_id == '':
					member = DataroomMembers.objects.get(member_id=user.id, dataroom_id=pk)
					# #print("member--------", member.end_user_group.first().id, user.id, pk)
					perm_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=da['id'],dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id)
				else:
					perm_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=da['id'],dataroom_id=pk, dataroom_groups_id=int(group_id))
				da['perm']['is_view_only'] = perm_obj.is_view_only
				da['perm']['is_no_access'] = perm_obj.is_no_access
				da['perm']['is_view_and_print'] = perm_obj.is_view_and_print
				da['perm']['is_view_and_print_and_download'] = perm_obj.is_view_and_print_and_download
				da['perm']['is_upload'] = perm_obj.is_upload
				da['perm']['is_watermarking'] = perm_obj.is_watermarking
				da['perm']['is_drm'] = perm_obj.is_drm
				da['perm']['is_editor'] = perm_obj.is_editor
				da['perm']['is_access'] = perm_obj.is_access
				da['perm']['is_shortcut'] = perm_obj.is_shortcut
			except:
				da['perm']['is_view_only'] = False
				da['perm']['is_no_access'] = True
				da['perm']['is_access'] = False
				da['perm']['is_view_and_print'] = False
				da['perm']['is_view_and_print_and_download'] = False
				da['perm']['is_upload'] = False
				da['perm']['is_watermarking'] = False
				da['perm']['is_drm'] = False
				da['perm']['is_editor'] = False
				da['perm']['is_shortcut'] = False
			# else:
			#     da['perm']['is_view_only'] = True
			#     da['perm']['is_no_access'] = False
			#     da['perm']['is_access'] = True
			#     da['perm']['is_view_and_print'] = True
			#     da['perm']['is_view_and_print_and_download'] = True
			#     da['perm']['is_upload'] = True
			#     da['perm']['is_watermarking'] = True
			#     da['perm']['is_drm'] = True
			#     da['perm']['is_editor'] = True
		return Response(data,status=status.HTTP_201_CREATED)


class GetAllFoldersAPI(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		name = request.GET.get("name")

		document = DataroomFolder.objects.filter(dataroom_id = pk, is_root_folder=True, is_deleted=False)
		group_id = request.GET.get('group_id')
		# #print("Group Id", group_id)

		data = []
		from . import utils
		for doc in document:
			docu = DataroomFolder.objects.get(id = doc.id)
			docu_serializer = DataroomFolderSerializer(docu)
			datas = utils.getIndexofFolder(docu_serializer.data)
			data.append(datas)
			docu1 = DataroomFolder.objects.filter(parent_folder = doc.id, is_deleted=False)
			if len(docu1) > 0:
				datas = []
				data.extend(utils.get_under_file(docu1,datas))
		for da in data:
			da['perm'] = {}
			if user.is_superadmin == False or user.is_admin or False:
				try:
					if group_id == '' :
						member = DataroomMembers.objects.get(member_id=user.id, dataroom_id=pk)
						# #print("member--------", member.end_user_group.first().id, user.id, pk)
						perm_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=da['id'],dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id)
					else:
						perm_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=da['id'],dataroom_id=pk, dataroom_groups_id=int(group_id))
					da['perm']['is_view_only'] = perm_obj.is_view_only
					da['perm']['is_no_access'] = perm_obj.is_no_access
					da['perm']['is_view_and_print'] = perm_obj.is_view_and_print
					da['perm']['is_view_and_print_and_download'] = perm_obj.is_view_and_print_and_download
					da['perm']['is_upload'] = perm_obj.is_upload
					da['perm']['is_watermarking'] = perm_obj.is_watermarking
					da['perm']['is_drm'] = perm_obj.is_drm
					da['perm']['is_editor'] = perm_obj.is_editor
					da['perm']['is_access'] = perm_obj.is_access
					da['perm']['is_shortcut'] = perm_obj.is_shortcut
				except:
					da['perm']['is_view_only'] = False
					da['perm']['is_no_access'] = True
					da['perm']['is_access'] = False
					da['perm']['is_view_and_print'] = False
					da['perm']['is_view_and_print_and_download'] = False
					da['perm']['is_upload'] = False
					da['perm']['is_watermarking'] = False
					da['perm']['is_drm'] = False
					da['perm']['is_editor'] = False
					da['perm']['is_shortcut'] = False
			else:
				da['perm']['is_view_only'] = True
				da['perm']['is_no_access'] = False
				da['perm']['is_access'] = True
				da['perm']['is_view_and_print'] = True
				da['perm']['is_view_and_print_and_download'] = True
				da['perm']['is_upload'] = True
				da['perm']['is_watermarking'] = True
				da['perm']['is_drm'] = True
				da['perm']['is_editor'] = True
				da['perm']['is_shortcut'] = True
		return Response(data,status=status.HTTP_201_CREATED)


class GetGroupAllFilesandFolders(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		da = []
		data = []
		perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(dataroom_groups_id=int(pk))
		#print("pk", pk, "perm_obj", perm_obj)
		if perm_obj:
			for each in perm_obj:
				docu = DataroomFolder.objects.get(id = each.folder.id)
				docu_serializer = DataroomFolderSerializer(docu)
				docu_serializer = docu_serializer.data
				docu_serializer['perm'] = {}
				docu_serializer['perm']['is_view_only'] = each.is_view_only
				docu_serializer['perm']['is_no_access'] = each.is_no_access
				docu_serializer['perm']['is_view_and_print'] = each.is_view_and_print
				docu_serializer['perm']['is_view_and_print_and_download'] = each.is_view_and_print_and_download
				docu_serializer['perm']['is_upload'] = each.is_upload
				docu_serializer['perm']['is_watermarking'] = each.is_watermarking
				docu_serializer['perm']['is_drm'] = each.is_drm
				docu_serializer['perm']['is_editor'] = each.is_editor
				docu_serializer['perm']['is_access'] = each.is_access
				docu_serializer['perm']['is_shortcut'] = each.is_shortcut
				da.append(docu_serializer)
			return Response(da, status=status.HTTP_201_CREATED)
		else:
			return Response(da, status=status.HTTP_201_CREATED)


class GetUpdateGroupAllFilesandFolders(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def post(self, request, pk, format=None):
		user = request.user
		data = request.data
		files = data['file']
		for da in files:
			folder_permission = DataroomGroupFolderSpecificPermissions.objects.filter(folder_id=da['id'],dataroom_id=int(pk),dataroom_groups_id=data['group_id']).first()
			if not folder_permission:
				serializer = DataroomGroupFolderSpecificPermissionsSerializer(data=da['perm'], context={'folder': da['id'], 'permission_given_by': user, 'dataroom': int(pk), 'dataroom_groups': request.data['group_id']})
				if serializer.is_valid():
					serializer.save()
			else:
				serializer = DataroomGroupFolderSpecificPermissionsSerializer(folder_permission,data = da['perm'])
				if serializer.is_valid():
					serializer.save()
			folder_obj = DataroomFolder.objects.get(id=da['id'])
			flag = True
			while flag:
				perm_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=da['id'],dataroom_id=int(pk),dataroom_groups_id=data['group_id'])
				if perm_obj.is_no_access == False:
					perm_obj.is_access = True
					perm_obj.save()
					if folder_obj.parent_folder_id == None or folder_obj.is_root_folder == True:
						flag = False
					else:
						folder1_obj = DataroomFolder.objects.get(id=folder_obj.parent_folder_id)
						try:
							perm1_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=folder1_obj.id,dataroom_id=int(pk),dataroom_groups_id=data['group_id'])
						except:
							perm1_obj = DataroomGroupFolderSpecificPermissions.objects.create(folder_id=folder1_obj.id,dataroom_id=int(pk),dataroom_groups_id=data['group_id'])
						# perm1_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=folder1_obj.id,dataroom_id=int(pk),dataroom_groups_id=data['group_id'])
						if perm_obj.is_no_access == False:
							perm1_obj.is_access = True
							perm1_obj.is_no_access = False
						else:
							perm1_obj.is_access = False
						perm1_obj.save()
						folder_obj = folder1_obj
						flag = False
				else:
					flag = False
		return Response(status=status.HTTP_204_NO_CONTENT)


class GetAllFolders(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		# user = request.user
		#print("All Folder")
		folder = DataroomFolder.objects.filter(dataroom_id = int(pk), is_deleted=False, is_folder=True)
		#print( "All Folder Queryset", folder)
		docu_serializer = DataroomFolderSerializer(folder, many=True)
		# #print("All FOlderss",docu_serializer.errors)
		data = docu_serializer.data
		#print( "All Folder Data", data)
		return Response(data,status=status.HTTP_201_CREATED)

class CountPdfPages(APIView):
	def get(self, request, pk, format=None):
		data = request.data
		user = request.user
		page = 0
		document = 0
		mydata = []
		new_page_count = []
		#date_from = datetime.now() - datetime.timedelta(days=30)
		# #print("pk=====>",pk)
		dataroom_folder_list = DataroomFolder.objects.filter(dataroom_id=pk, is_deleted=False, is_folder=False).aggregate(Sum('pages'))
		Document_list = DataroomFolder.objects.filter(dataroom_id=pk).values_list('path')
		from azure.storage.blob import BlockBlobService, PublicAccess
		block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
		container_name ='docullycontainer'
		# block_blob_service.get_blob_to_path(container_name, filename, file_name[-1])
		for i in Document_list:
			matchobj = re.match( r'\(\'(.*)\'\,\)', str(i), re.M|re.I)
			if matchobj.group(1).endswith('xlsx'):
				match_blob = matchobj.group(1).split("/")
				# #print(match_blob[-2]+'/'+match_blob[-1])
				isExist = block_blob_service.exists(container_name, match_blob[-2]+'/'+ match_blob[-1])
				if isExist:
					block_blob_service.get_blob_to_path(container_name, match_blob[-2]+'/'+match_blob[-1], match_blob[-1])
					if os.path.exists(match_blob[-1]):
						xl = pd.ExcelFile(match_blob[-1])
						res = len(xl.sheet_names)
						new_page_count.append(res)
						os.remove(match_blob[-1])
			elif matchobj.group(1).endswith('pdf'):
				# #print("pdf method is called ===>")
				match_blob = matchobj.group(1).split("/")
				# #print(match_blob[-2]+'/'+match_blob[-1])
				isExist = block_blob_service.exists(container_name, match_blob[-2]+'/'+ match_blob[-1])
				if isExist:
					block_blob_service.get_blob_to_path(container_name, match_blob[-2]+'/'+match_blob[-1], match_blob[-1])
					if os.path.exists(match_blob[-1]):
						pdf = PdfFileReader(open(match_blob[-1],'rb'))
						a = pdf.getNumPages()
						# #print("value of pdf_pages ==>",a)
						new_page_count.append(a)
						os.remove(match_blob[-1])
			elif matchobj.group(1).endswith('pptx'):
				# #print("pdf method is called ===>")
				match_blob = matchobj.group(1).split("/")
				# #print(match_blob[-2]+'/'+match_blob[-1])
				isExist = block_blob_service.exists(container_name, match_blob[-2]+'/'+ match_blob[-1])
				if isExist:
					block_blob_service.get_blob_to_path(container_name, match_blob[-2]+'/'+match_blob[-1], match_blob[-1])
					if os.path.exists(match_blob[-1]):
						os.remove(match_blob[-1])
			elif matchobj.group(1).endswith('docx'):
				# #print("pdf method is called ===>")
				match_blob = matchobj.group(1).split("/")
				# #print(match_blob[-2]+'/'+match_blob[-1])
				isExist = block_blob_service.exists(container_name, match_blob[-2]+'/'+ match_blob[-1])
				if isExist:
					block_blob_service.get_blob_to_path(container_name, match_blob[-2]+'/'+match_blob[-1], match_blob[-1])
					if os.path.exists(match_blob[-1]):
						os.remove(match_blob[-1])
			elif matchobj.group(1).endswith('csv'):
				# #print("csv method is called ===>")
				match_blob = matchobj.group(1).split("/")
				# #print(match_blob[-2]+'/'+match_blob[-1])
				isExist = block_blob_service.exists(container_name, match_blob[-2]+'/'+ match_blob[-1])
				if isExist:
					block_blob_service.get_blob_to_path(container_name, match_blob[-2]+'/'+match_blob[-1], match_blob[-1])
					if os.path.exists(match_blob[-1]):
						# #print("len ====>",match_blob[-1])
						# xl = pd.read_csv(match_blob[-1])
						# res = len(xl.sheet_names)
						# #print("len of csv =================>",xl)
						new_xlsx_file = match_blob[-1].split(".")
						filepath_out =  new_xlsx_file[0] +".xlsx"
						pd.read_csv(match_blob[-1], delimiter=",").to_excel(filepath_out)
						xl = pd.ExcelFile(filepath_out)
						res_csv = len(xl.sheet_names)
						# #print("res_csv_file =====>",res_csv)
						new_page_count.append(res_csv)
						os.remove(match_blob[-1])
						os.remove(filepath_out)

		# #print("sum==>",sum(new_page_count))
		document = DataroomFolder.objects.filter(dataroom_id=pk, is_deleted=False, is_folder=False).count()
		return Response({'page_count': sum(new_page_count), 'document_count': document})

class GetAllNotViewedDocument(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		# print ("get all recent document data")        
		data = request.data
		user = request.user
		count = 0
		is_type =  int(request.GET.get('type'))
		view_list = FolderView.objects.filter(dataroom_id=pk, user_id=user.id).values('folder_id')
		member = DataroomMembers.objects.filter(dataroom_id=pk,member_id=user.id).first()
		# perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id).values('folder_id')
		try:
			perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id).values('folder_id')
		except:
			perm_obj = None
		if is_type == 2:
			dataroom_folder = DataroomFolder.objects.filter(id__in=view_list, dataroom_id=pk,is_deleted=False, is_folder=False)
			count = DataroomFolder.objects.filter(id__in=view_list, dataroom_id=pk,is_deleted=False, is_folder=False).count()
		else:
			if member.is_dataroom_admin or member.is_la_user:
				dataroom_folder = DataroomFolder.objects.exclude(id__in=view_list).filter(dataroom_id=pk, is_deleted=False, is_folder=False)                
				count = DataroomFolder.objects.exclude(id__in=view_list).filter(dataroom_id=pk, is_deleted=False, is_folder=False).count()                
			else:
				dataroom_folder = DataroomFolder.objects.exclude(id__in=view_list).filter(id__in=perm_obj,dataroom_id=pk, is_deleted=False, is_folder=False)
				count = DataroomFolder.objects.exclude(id__in=view_list).filter(id__in=perm_obj,dataroom_id=pk, is_deleted=False, is_folder=False).count()

		# count = dataroom_folder
		# #print(count)
		page = paginator.paginate_queryset(dataroom_folder, request)

			
			# dataroom_folder = DataroomFolder.objects.exclude(id__in=view_list).filter(dataroom_id=pk, id__in=perm_obj,is_deleted=False, is_folder=False)
		dataroom_folder_serializer = DataroomFolderSerializer(page, many=True)
		data = dataroom_folder_serializer.data
		from . import utils
		for da in data:
			da['index'] = utils.getIndexes(da)
			da['parent_folders'] = DataroomFolder.objects.get(id = da['parent_folder']).name
			if member.is_end_user:
				perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id).first()
				perm_serializer = DataroomGroupFolderSpecificPermissionsSerializer(perm_obj, many=False)
				da['perm'] = perm_serializer.data
			else:
				da['perm'] = {}
				da['perm']['is_view_only'] = True
				da['perm']['is_no_access'] = False
				da['perm']['is_access'] = True
				da['perm']['is_view_and_print'] = True
				da['perm']['is_view_and_print_and_download'] = True
				da['perm']['is_upload'] = True
				da['perm']['is_watermarking'] = True
				da['perm']['is_drm'] = True
				da['perm']['is_editor'] = True
				da['perm']['is_shortcut'] = True
		data = sorted(data,key=lambda x : x['created_date'], reverse=True)
		return Response({'data':data, 'size':count})

class GetAllNotPrintedDocument(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		# print ("get all recent document data")        
		data = request.data
		user = request.user
		is_type =  int(request.GET.get('type'))
		member = DataroomMembers.objects.filter(dataroom_id=pk,member_id=user.id).first()
		# perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id).values('folder_id')
		try:
			perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id).values('folder_id')
		except:
			perm_obj = None
		print_list = FolderPrint.objects.filter(dataroom_id=pk, user_id=user.id).values('folder_id')
		if is_type == 2:
			dataroom_folder = DataroomFolder.objects.filter(id__in=print_list,dataroom_id=pk, is_deleted=False, is_folder=False)
			count = DataroomFolder.objects.filter(id__in=print_list,dataroom_id=pk, is_deleted=False, is_folder=False).count()
		else:
			if member.is_dataroom_admin or member.is_la_user:
				dataroom_folder = DataroomFolder.objects.exclude(id__in=print_list).filter(dataroom_id=pk, is_deleted=False, is_folder=False)                
				count = DataroomFolder.objects.exclude(id__in=print_list).filter(dataroom_id=pk, is_deleted=False, is_folder=False).count()                
			else:
				dataroom_folder = DataroomFolder.objects.exclude(id__in=print_list).filter(id__in=perm_obj,dataroom_id=pk, is_deleted=False, is_folder=False)
				count = DataroomFolder.objects.exclude(id__in=print_list).filter(id__in=perm_obj,dataroom_id=pk, is_deleted=False, is_folder=False).count()
			# dataroom_folder = DataroomFolder.objects.exclude(id__in=print_list).filter(id__in=perm_obj,dataroom_id=pk, is_deleted=False, is_folder=False)
		page = paginator.paginate_queryset(dataroom_folder, request)
		dataroom_folder_serializer = DataroomFolderSerializer(page, many=True)
		data = dataroom_folder_serializer.data
		from . import utils
		for da in data:
			da['index'] = utils.getIndexes(da)
			da['parent_folders'] = DataroomFolder.objects.get(id = da['parent_folder']).name
			if member.is_end_user:
				perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id).first()
				perm_serializer = DataroomGroupFolderSpecificPermissionsSerializer(perm_obj, many=False)
				da['perm'] = perm_serializer.data
			else:
				da['perm'] = {}
				da['perm']['is_view_only'] = True
				da['perm']['is_no_access'] = False
				da['perm']['is_access'] = True
				da['perm']['is_view_and_print'] = True
				da['perm']['is_view_and_print_and_download'] = True
				da['perm']['is_upload'] = True
				da['perm']['is_watermarking'] = True
				da['perm']['is_drm'] = True
				da['perm']['is_editor'] = True
				da['perm']['is_shortcut'] = True
		data = sorted(data,key=lambda x : x['created_date'], reverse=True)
		return Response({'data':data, 'size':count})

class GetAllNotDownloadDocument(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		# print ("get all recent document data")        
		data = request.data
		user = request.user
		count = 0
		is_type =  int(request.GET.get('type'))
		member = DataroomMembers.objects.filter(dataroom_id=pk,member_id=user.id).first()
		# if member.is_dataroom_admin or member.is_la_user:
		try:
			perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id).values('folder_id')
		except:
			perm_obj = None
		download_list = FolderDownload.objects.filter(dataroom_id=pk, user_id=user.id).values('folder_id')
		if is_type == 2:
			dataroom_folder = DataroomFolder.objects.filter(id__in=download_list, dataroom_id = pk, is_deleted=False, is_folder=False)
			count = DataroomFolder.objects.filter(id__in=download_list, dataroom_id = pk, is_deleted=False, is_folder=False).count()
		else:
			if member.is_dataroom_admin or member.is_la_user:
				dataroom_folder = DataroomFolder.objects.exclude(id__in=download_list).filter(dataroom_id=pk, is_deleted=False, is_folder=False)                
				count = DataroomFolder.objects.exclude(id__in=download_list).filter(dataroom_id=pk, is_deleted=False, is_folder=False).count()                
			else:
				dataroom_folder = DataroomFolder.objects.exclude(id__in=download_list).filter(id__in=perm_obj,dataroom_id=pk, is_deleted=False, is_folder=False)
				count = DataroomFolder.objects.exclude(id__in=download_list).filter(id__in=perm_obj,dataroom_id=pk, is_deleted=False, is_folder=False).count()

		page = paginator.paginate_queryset(dataroom_folder, request)
		dataroom_folder_serializer = DataroomFolderSerializer(page, many=True)
		data = dataroom_folder_serializer.data
		from . import utils
		for da in data:
			da['index'] = utils.getIndexes(da)
			da['parent_folders'] = DataroomFolder.objects.get(id = da['parent_folder']).name
			# #print("Permmmmmmm", perm_obj)
			if member.is_end_user:
				perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id).first()
				perm_serializer = DataroomGroupFolderSpecificPermissionsSerializer(perm_obj, many=False)
				da['perm'] = perm_serializer.data
			else:
				da['perm'] = {}
				da['perm']['is_view_only'] = True
				da['perm']['is_no_access'] = False
				da['perm']['is_access'] = True
				da['perm']['is_view_and_print'] = True
				da['perm']['is_view_and_print_and_download'] = True
				da['perm']['is_upload'] = True
				da['perm']['is_watermarking'] = True
				da['perm']['is_drm'] = True
				da['perm']['is_editor'] = True
				da['perm']['is_shortcut'] = True
		data = sorted(data,key=lambda x : x['created_date'], reverse=True)
		return Response({'data':data, 'size':count})


class GetAllUploadStatusDocument(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		print ("get all recent document data")        
		data = request.data
		user = request.user
		#date_from = datetime.now() - datetime.timedelta(days=30)
		dataroom_folder = DataroomFolder.objects.filter(dataroom_id=pk, is_deleted=False, is_folder=False, user_id=user.id)
		count = DataroomFolder.objects.filter(dataroom_id=pk, is_deleted=False, is_folder=False, user_id=user.id).count()
		#print(count)
		page = paginator.paginate_queryset(dataroom_folder, request)
		dataroom_folder_serializer = DataroomFolderSerializer(page, many=True)
		data = dataroom_folder_serializer.data
		from . import utils
		for da in data:
			# da['index'] = utils.getIndexes(da)
			user_obj = User.objects.get(id=da['user'])
			da['user'] = (UserSerializer(user_obj, many=False)).data
			da['parent_folders'] = DataroomFolder.objects.get(id = da['parent_folder']).name
			try:
			   perm = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=pk)
			   perm_serializer = DataroomGroupFolderSpecificPermissionsSerializer(perm, many=False)
			   da['perm'] = perm_serializer.data
			except:
			   da['perm'] = {}
			# if user.is_superadmin == True or user.is_admin == True:
			#     da['perm']['is_view_only'] = True
			#     da['perm']['is_no_access'] = False
			#     da['perm']['is_access'] = True
			#     da['perm']['is_view_and_print'] = True
			#     da['perm']['is_view_and_print_and_download'] = True
			#     da['perm']['is_upload'] = True
			#     da['perm']['is_watermarking'] = True
			#     da['perm']['is_drm'] = True
			#     da['perm']['is_editor'] = True
		data = sorted(data,key=lambda x :x['created_date'], reverse=True)
		return Response({'data':data,'size':count})


class GetFilesUpdated(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		# print ("get all recent document data")    
		update_view = request.GET.get("type")
		user = request.user
		folder = DataroomFolder.objects.get(id=pk)
		#print("folder pathhhh", folder.path)
		if update_view == 'view':
			#print("m in view")
			folder.is_view = True
			folder.save()
			view = FolderView()
			#print(view,"finding folder view")
			view.folder_id = folder.id
			view.user_id = user.id
			view.dataroom_id = folder.dataroom_id
			view.save()
		elif update_view == 'print':
			folder.is_print = True
			folder.save()
			print_obj = Folder#print()
			print_obj.folder_id = folder.id
			print_obj.user_id = user.id
			print_obj.dataroom_id = folder.dataroom_id
			print_obj.save()
		elif update_view == 'download':
			folder.is_download = True
			folder.save()
			download = FolderDownload()
			download.folder_id = folder.id
			download.user_id = user.id
			download.dataroom_id = folder.dataroom_id
			download.save()
		else:
			pass
		#response = Response(status=status.HTTP_204_NO_CONTENT)
		#response["Access-Control-Allow-Origin"] = "*"
		#response["Access-Control-Allow-Methods"] = "*"
		#response["Access-Control-Allow-Headers"] = "*"
		#return response
		return Response(status=status.HTTP_204_NO_CONTENT)



class DownloadFiles(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )
	def get(self, request, pk, format=None):
		user = request.user
		print(user,"user of download")
		print(pk,"pkkkkk")
		update_view = request.GET.get("type")
		print("========>",pk)
		print("update_view=======>",update_view)
		from dataroom.models import Watermarking
		# #print("watermark_data ========>",Watermarking.objects.filter().values())
		# #print(DataroomFolder.objects.filter(id=pk).values())
		# #print("new_data==>",DataroomFolder.objects.filter(id=pk).values_list('dataroom_id'))
		# new_dataroom_id = list(DataroomFolder.objects.filter(id=pk).values_list('dataroom_id'))
		# #print(new_dataroom_id[0],len(new_dataroom_id),"new_dataroom_id ====>",type(new_dataroom_id))
		# #print("new_data==>",DataroomFolder.objects.get(id=pk))
		# data_new = re.match(r'', new_dataroom_id, re.M|re.I)
		# new_dataroom_id = str(new_dataroom_id).replace('\(\)\,','')
		# #print("data=>",new_dataroom_id)

		# #print("watermark_data ========>",Watermarking.objects.filter(dataroom_id=new_dataroom_id[0]).values())
		# #print("datroom_id===>",DataroomFolder.objects.filter(id=pk).only('dataroom_id').values())
		# #print('###########################################')

		# #print(request.data)
		# #print('###########################################')
		# #print(update_view)
		# #print('###########################################')

		from wsgiref.util import FileWrapper
		from django.http import FileResponse
		folder = DataroomFolder.objects.get(id=pk)
		print("id===>",folder.dataroom_id)
		room_id=folder.dataroom_id
		print(room_id,"room_id")
		print("folder.id",folder.id)
		folder_path = str(folder.path)
		print("folder_path",folder_path)
		filename=str(folder.path)
		# #print(folder_path,"<==========Filenameeeeeee=========>", filename)
		# try: 
		#print("watermark_data===>",Watermarking.objects.filter(dataroom_id=folder.dataroom_id).values())
		member_idd = DataroomMembers.objects.filter(dataroom_id=room_id,is_dataroom_admin=True).values('member_id')
		print(member_idd,"member_idd")
		member = DataroomMembers.objects.filter(member_id=user.id, dataroom_id=folder.dataroom_id).first()
		print(member.is_la_user,"<====== memeber ====>",member.is_dataroom_admin)
		print("member is up ",member)
		if member.is_dataroom_admin or member.is_la_user:
			is_download = True
			is_print = True
		else:
			perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(folder_id=folder.id,dataroom_id=folder.dataroom_id, dataroom_groups_id=member.end_user_group.first().id).first()
			is_download = perm_obj.is_view_and_print_and_download
			is_print = perm_obj.is_view_and_print
			group_perm = DataroomGroupPermission.objects.filter(dataroom_groups_id=member.end_user_group.first().id,dataroom=folder.dataroom.id).first()
			# #print('################%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%###########################')

			#print("Permm", group_perm.id, group_perm.is_watermarking, group_perm.is_doc_as_pdf, group_perm.is_excel_as_pdf)
			
			if group_perm.is_watermarking == True:
				if group_perm.is_doc_as_pdf == True or group_perm.is_excel_as_pdf==True:
					if perm_obj.watermarking_file:
						folder_path = str(perm_obj.watermarking_file)
						filename = str(perm_obj.watermarking_file)
						print(filename,"filename in perm")
		# except:
		#     is_download = False
		#     is_print = False
		if folder:  
			if update_view == 'view':
				folder.is_view = True
				folder.save()
				view = FolderView()
				view.folder_id = folder.id
				view.user_id = user.id
				view.dataroom_id = folder.dataroom_id
				view.save()
			elif update_view == 'print':
				folder.is_print = True
				folder.save()
				print_obj = Folder#print()
				print_obj.folder_id = folder.id
				print_obj.user_id = user.id
				print_obj.dataroom_id = folder.dataroom_id
				print_obj.save()
			elif update_view == 'download' and (is_download==True or user.is_superadmin==True or user.is_admin==True):
				print('########download details###############')

				folder.is_download = True
				folder.save()
				download = FolderDownload()
				download.folder_id = folder.id
				download.user_id = user.id
				download.dataroom_id = folder.dataroom_id
				download.save()
			else:
				pass 
		#print(filename,"---print ----", is_download, user.is_admin, user.is_superadmin) 
		# group_perm = DataroomGroupPermission.objects.filter(dataroom_id=data.dataroom_id, dataroom_groups_id=member.end_user_group.first().id).first()
		# #print("Download_function ======>",group_perm.is_watermarking)
		from dataroom import watermarking
		#print("mail_31_march===>",watermarking.getmail(user))
		# test_dataroom_17_March&yaberig541@mailimail.com.png
		import os
		new_pdf_list = []
		if (is_download == True or is_print==True) or user.is_superadmin == True or user.is_admin == True:
			#print("Filenameeeeeee", filename)
			from constants import constants
			extensions = constants.extensions
		#     print ("download method called")        
			for key , value in extensions.items():
		#         #print("key", key, "value", value)
		#         #print('--------------------------------------------')
		#         path_list = filename.split("/")
		#         #print("key", filename)
		#         #print('////////////////////////////////////////////////')
		#         path = settings.MEDIA_ROOT+"/"
		#         #print()
		#         print ("inside for loop")
				# if filename.endswith('.pdf'):
				#     # #print("watermarking_user",watermarking.getmail(user))
				#     from azure.storage.blob import BlockBlobService, PublicAccess
				#     block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
				#     container_name ='docullycontainer'
				#     global b
				#     b = watermarking.getmail(user)
				#     file_name = filename.split("/")
				#     #print("Downloadf---",type(file_name[-1]),"<======= file_name==>",filename)
				#     block_blob_service.get_blob_to_path(container_name, filename, file_name[-1])
				#     #print("watermark_data===>",Watermarking.objects.filter(dataroom_id=folder.dataroom_id).values())
				#     from pdf2image import convert_from_path

					# #print("set ====>",set(new_pdf_list))
				if filename.endswith(key):
		#             #print('<<<<<<<<<<<<<<<<<<<<<<<<<<found')
		#             path += filename
		#             name = str(path_list[-1])
					# #print("filename===>",filename.split("/"))

					import os, uuid, sys
					from azure.storage.blob import BlockBlobService, PublicAccess
					block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
					container_name ='docullycontainer'
					file_name = filename.split("/")
					print(filename,"filename check")
					print(file_name[-1],"file_name[-1] check")
					block_blob_service.get_blob_to_path(container_name, filename, file_name[-1])
					path = settings.MEDIA_ROOT
					print(path,"settings.path")
					print(file_name[-1],"filename")
					extension=os.path.splitext(file_name[-1])[-1]
					import PyPDF2
					print(extension,"extension")
					data_of_per=DataroomGroupPermission.objects.filter(dataroom=room_id).values("is_doc_as_pdf","is_watermarking","is_excel_as_pdf")
					print(data_of_per,"data_of_per")
					print(data_of_per[0],"data_of_per[0]")
					print(type(data_of_per))
					if(extension=='.pdf'):
						import os, uuid, sys
						from azure.storage.blob import BlockBlobService, PublicAccess
						block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
						container_name ='docullycontainer'
						file_name = filename.split("/")
						print(filename,"filename check")
						print(file_name[-1],"file_name[-1] check")
						block_blob_service.get_blob_to_path(container_name, filename, file_name[-1])
						path = settings.MEDIA_ROOT
						print(path,"settings.path")
						print(file_name[-1],"filename")
						extension=os.path.splitext(file_name[-1])[-1]
						import PyPDF2
						print(extension,"extension")
						data_of_per=DataroomGroupPermission.objects.filter(dataroom=room_id).values("is_doc_as_pdf","is_watermarking","is_excel_as_pdf")
						print(data_of_per,"data_of_per")
						print(data_of_per[0],"data_of_per[0]")
						print(type(data_of_per))
						# admin_name=User.objects.filter()
						# '/home/cdms_backend/cdms2/media/'+filename
						from dataroom.serializers import WatermarkingSerializer
						from dataroom.pdf_watermarking import GeneratePDF
						userid=user.id
						watermarking = Watermarking.objects.filter(dataroom_id=int(room_id)).order_by('id')
						for i in watermarking:
							i.user_id=userid
						# print("value of watermarking =====>",Watermarking.objects.all().values())
						print("watermarking =====>",watermarking)
						serializer = WatermarkingSerializer(watermarking,many=True)
						data = serializer.data
						print(data,"in download")
						from userauth import utils
						ip = utils.get_client_ip(request)
						print(pk,"pkk")
						GeneratePDF(data,ip,user,pk)
						watermarkfile="/home/cdms_backend/cdms2/Admin_Watermark/"+str(room_id)+".pdf"
						outputfile="/home/cdms_backend/cdms2/dataroom/success.pdf"
						pdf_writer=PyPDF2.PdfFileWriter()
						if (os.path.exists(file_name[-1])):
							with open(file_name[-1], 'rb') as fh:
								pdf=PyPDF2.PdfFileReader(fh)
								with open(watermarkfile,'rb') as watermarkfile:
									watermarkfile_pdf=PyPDF2.PdfFileReader(watermarkfile)
									for i in range(pdf.getNumPages()):
										p=pdf.getPage(i)
										p.mergePage(watermarkfile_pdf.getPage(0))
										pdf_writer.addPage(p)
									with open(outputfile,'wb') as outputfileeee:
										pdf_writer.write(outputfileeee)
									with open(outputfile, 'rb') as output:
										response = HttpResponse(output.read(), content_type=value)
										response['Content-Disposition'] = 'inline; filename=' + filename
										os.remove(file_name[-1])
									return response
								raise Http404   
					elif (extension!='.pdf') and (data_of_per[0]["is_watermarking"]==False) and (data_of_per[0]["is_excel_as_pdf"]==False):
						print("second elif")
						import os, uuid, sys
						from azure.storage.blob import BlockBlobService, PublicAccess
						block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
						container_name ='docullycontainer'
						file_name = filename.split("/")
						print(filename,"filename check")
						print(file_name[-1],"file_name[-1] check")
						block_blob_service.get_blob_to_path(container_name, filename, file_name[-1])
						path = settings.MEDIA_ROOT
						print(path,"settings.path")
						print(file_name[-1],"filename")
						with open(file_name[-1], 'rb') as output:
							response = HttpResponse(output.read(), content_type=value)
							response['Content-Disposition'] = 'inline; filename=' + filename
							os.remove(file_name[-1])
							return response
						raise Http404
						# ,data_of_per["is_excel_as_pdf"]==False or True
					elif (extension=='.docx' or ".doc" or ".ppt" or ".pptx" or '.xlsx' or '.xls' or '.csv') and (data_of_per[0]["is_doc_as_pdf"]==True) and (data_of_per[0]["is_watermarking"]==True):
						print("in last elif")
						import os, uuid, sys
						from azure.storage.blob import BlockBlobService, PublicAccess
						block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
						container_name ='docullycontainer'
						from pathlib import Path
						filename1 = Path(file_name[-1])
						filename_wo_ext = filename1.with_suffix('')
						print(filename_wo_ext,"filename_wo_ext")
						pdf_filename=str(filename_wo_ext)+".pdf"
						print(pdf_filename,"pdf_filename")
						blobname=file_name[0]+"/"+pdf_filename
						print(blobname,"blobname")
						print(container_name, blobname,pdf_filename,"details download")
						block_blob_service.get_blob_to_path(container_name, blobname,pdf_filename)
						with open(pdf_filename, 'rb') as output:
							response = HttpResponse(output.read(), content_type="application/pdf")
							response['Content-Disposition'] = 'inline; filename=' + pdf_filename
							# os.remove(file_name[-1])
							return response
						# raise Http404
					# #print("path =====>",path)
					# storage_path = "https://docullystorage.blob.core.windows.net/docullycontainer/"
					# file_path = os.path.join(storage_path, filename)
					# r = requests.get(file_path)
					# response = HttpResponse(r.content,content_type=value)
					# response['Content-Disposition'] = 'attachment; filename='+ filename
					# return response

					# download_file_path = os.path.join(local_path, str.replace(file_path ,'.txt', 'DOWNLOAD.txt'))
					# with open(download_file_path, "wb") as download_file:
					#     download_file.write(blob_client.download_blob().readall())

					# #print("file_path ====>",file_path)
					# r = requests.get(file_path)
					# #print("value of r=====>",r)
					
					# #print(r.content,"value of r ====>",len(r.text),r.text)
					# with open(r.text,"rb") as fh:
					# 	#print("readlines==========>",fh.readlines())
					# text_file = r.text
					# s#print("value of response =====>",response)
					# #print("content =====>",r.content)
					# with open(file_name[-1],'rb') as fh:
					# 	#print("===========>",fh.read())
					# response = HttpResponse(content_type=value)
					# response['Content-Disposition'] = 'attachment; filename='+ file_name[-1]
					# return response
					# with open(file_path + filename,"rb") as fh:
					# 	response = HttpResponse(fh.read(),content_type=value)
					# 	response['Content-Disposition'] = 'inline; filename=' + file_name[-1]
					# 	#print("response=====>",response)
					# 	return response

					# return HttpResponse(file_path + filename,content_type=value)
					# #print("file_path =======>",file_path)
					# #print("file_path =======>",os.path.exists(file_path))
					# #print(os.path.basename(file_path),"value==>",value,"content-type==>",folder.file_content_type,"file_path =======>",type(os.path.basename(file_path)))
					# if os.path.exists(file_path.rstrip()):
					#     #print("entered in the if condition")
					#     with open(file_path.rstrip(), 'rb') as fh:
					#         response = HttpResponse(fh.read(), content_type=folder.file_content_type)
					#         response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
					#         # response['Content-Type'] = value
					#         response['name'] = os.path.basename(file_path)
					#         return response
					# raise Http404

				else:
					pass
		#             print ("no extension found")
		#             path += filename
		#             name = str(path_list[-1])
		#             if os.path.exists(path):
		#                 with open(path, 'rb') as fh:
		#                     response = HttpResponse(fh.read(), content_type=folder.file_content_type)
		#                     response['Content-Disposition'] = 'inline; filename=' + str(name) 
		#                     response['name'] = str(name) 
		#                     return response
		# else:
		return Response(None, status=status.HTTP_201_CREATED)
		# return Http404

def build_tree_recursive(tree, parent, nodes):
	children  = [n for n in nodes if n['parent_folder'] == parent]
	for child in children:
		# #print("Childdddddddd-------", child)
		data = {}
		data['name'] = child['names']
		data['id'] = child['id']
		data['children'] = []
		tree.append(data)
		build_tree_recursive(data['children'], child['id'], nodes)

class FoldersHierarchy(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		perm_obj = None
		# document = DataroomFolder.objects.filter(id__in = perm_obj,dataroom_id = pk, is_root_folder=True, is_deleted=False).order_by('index')
		# if user.is_superadmin == True:
		#     document = DataroomFolder.objects.filter(dataroom_id = pk, is_root_folder=True, is_deleted=False).order_by('index')
		#     perm_obj = None
		# else:
		member = DataroomMembers.objects.filter(member_id=user.id, dataroom_id=pk, is_deleted=False).first()
		# #print("memberrrrrr", member.is_la_user, user.id, pk)
		if member.is_la_user == True or member.is_dataroom_admin == True:
			document = DataroomFolder.objects.filter(dataroom_id = pk, is_root_folder=True, is_deleted=False).order_by('index')
			perm_obj = None
		else:
			perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id, is_view_only=True).values('folder_id')
			document = DataroomFolder.objects.filter(id__in = perm_obj,dataroom_id = pk, is_root_folder=True, is_deleted=False).order_by('index')

		data = []
		from . import utils
		for doc in document:
			docu = DataroomFolder.objects.get(id = doc.id)
			docu_serializer = DataroomFolderSerializer(docu)
			datas = docu_serializer.data
			utils.getIndexofFolder(datas)
			data.append(datas)

			# if user.is_superadmin == True:
			#     docu1 = DataroomFolder.objects.filter(parent_folder = doc.id, is_folder=True, is_deleted=False).order_by('index')
			# else:
			member = DataroomMembers.objects.filter(member_id=user.id, dataroom_id=pk, is_deleted=False).first()
			if member.is_la_user == True or member.is_dataroom_admin == True:
				docu1 = DataroomFolder.objects.filter(parent_folder = doc.id, is_folder=True, is_deleted=False).order_by('index')
			else:
				docu1 = DataroomFolder.objects.filter(parent_folder = doc.id, is_folder=True, is_deleted=False, id__in = perm_obj).order_by('index')
			if len(docu1) > 0:
				datas = []
				data.extend(utils.get_under_folder(docu1, datas, user, perm_obj, pk))
				# #print("dataaa", data)
		# #print("dataaaaa", data)
		tree = []
		# fill in tree starting with roots (those with no parent)
		build_tree_recursive(tree, None, data)
		# #print("treee", tree)
		return Response(tree, status=status.HTTP_201_CREATED)


class FoldersTree(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, parent_id, format=None):
		user = request.user
		perm_obj = None
		# document = DataroomFolder.objects.filter(id__in = perm_obj,dataroom_id = pk, is_root_folder=True, is_deleted=False).order_by('index')
		# if user.is_superadmin == True:
		#     document = DataroomFolder.objects.filter(dataroom_id = pk, is_root_folder=True, is_deleted=False).order_by('index')
		#     perm_obj = None
		# else:
		print(user.id,"user id")
		print(pk,"pk")
		member = DataroomMembers.objects.filter(member_id=user.id, dataroom_id=pk, is_deleted=False).first()
		print("memberrrrrr", member)

		if member.is_la_user == True or member.is_dataroom_admin == True:
			if int(parent_id)==0:
				document = DataroomFolder.objects.filter(dataroom_id = pk, is_root_folder=True, is_deleted=False).order_by('index')
				#print(parent_id, pk,1958)
			else:
				document = DataroomFolder.objects.filter(dataroom_id = pk, parent_folder=parent_id, is_root_folder=False, is_deleted=False, is_folder=True).order_by('index')
				#print(parent_id, pk,1961)
			perm_obj = None
		else:
			perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id, is_view_only=True).values('folder_id')
			if int(parent_id)==0:
				document = DataroomFolder.objects.filter(id__in = perm_obj,dataroom_id = pk, is_root_folder=True, is_deleted=False).order_by('index')
				#print(parent_id, pk,1967)
			else:
				document = DataroomFolder.objects.filter(dataroom_id = pk, parent_folder=parent_id, is_root_folder=False, is_deleted=False, is_folder=True).order_by('index')
				#print(parent_id, pk,1970)

		data = []
		from . import utils
		for doc in document:
			docu = DataroomFolder.objects.get(id = doc.id)
			docu_serializer = DataroomFolderSerializer(docu)
			datas = docu_serializer.data
			utils.getIndexofFolder(datas)

			# if user.is_superadmin == True:
			#     docu1 = DataroomFolder.objects.filter(parent_folder = doc.id, is_folder=True, is_deleted=False).order_by('index')
			# else:
			member = DataroomMembers.objects.filter(member_id=user.id, dataroom_id=pk, is_deleted=False).first()
			if member.is_la_user == True or member.is_dataroom_admin == True:
				docu1 = DataroomFolder.objects.filter(parent_folder = doc.id, is_folder=True, is_deleted=False).order_by('index')
			else:
				docu1 = DataroomFolder.objects.filter(parent_folder = doc.id, is_folder=True, is_deleted=False, id__in = perm_obj).order_by('index')
			if len(docu1) > 0:
				# if(utils.check_subfolder(docu1, datas, user, perm_obj, pk)):
				#     datas['hasChildren']=True
				# else:
				datas['hasChildren']=True
			else:
				datas['hasChildren']=False
			data.append(datas)
				# #print("dataaa", data)
		# #print("dataaaaa", data)
		# tree = []
		# fill in tree starting with roots (those with no parent)
		# build_tree_recursive(tree, None, data)
		# #print("treee", tree)
		return Response(data, status=status.HTTP_201_CREATED)


class FileActivitybyDateReport(APIView):
	"""docstring for ActivitybyDateReport"""
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		import datetime
		from_date = request.GET.get("from_date")
		to_date = request.GET.get("to_date")
		deleted = False if request.GET.get("deleted") == 'false' else True
		# #print("deleted",deleted)
		todays_date = datetime.datetime.strftime(datetime.datetime.strptime( to_date, '%Y-%m-%d'),'%Y-%m-%d 23:59:59+05:30')
		first_date = datetime.datetime.strftime(datetime.datetime.strptime( from_date, '%Y-%m-%d'),'%Y-%m-%d 00:00:00+05:30')
		#print('<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<')
		#print("request get", first_date, todays_date)
		document = DataroomFolder.objects.filter(dataroom_id = pk, is_folder=False, is_deleted=deleted, created_date__gte=first_date, created_date__lte=todays_date)
		count = DataroomFolder.objects.filter(dataroom_id = pk, is_folder=False, is_deleted=deleted, created_date__gte=first_date, created_date__lte=todays_date).count()
		page = paginator.paginate_queryset(document, request)
		serializer = DataroomFolderSerializer(page, many=True)
		data = serializer.data
		data.sort(key=lambda item:item['created_date'], reverse=True)
		for da in data:
			da['parent'] = DataroomFolder.objects.get(id = da['parent_folder']).name
			da['user'] = UserSerializer(User.objects.get(id=da['user']),many=False).data
		return Response({'data':data,'size':count}, status=status.HTTP_201_CREATED)

class ActivitybyDateReport(APIView):
	"""docstring for ActivitybyDateReport"""
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )


	def get(self, request, pk, format=None):
		user = request.user
		import datetime
		from_date = request.GET.get("from_date")
		to_date = request.GET.get("to_date")
		todays_date = datetime.datetime.strptime( to_date, '%Y-%m-%d')
		first_date = datetime.datetime.strptime( from_date, '%Y-%m-%d')
		days = ((todays_date-first_date).days)+1
		#lines addded by harish 
		datarooms_date = Dataroom.objects.get(id=pk).created_date
		dateobject=datetime.datetime.strptime(str(datarooms_date),'%Y-%m-%d %H:%M:%S.%f')
		d_date =dateobject.strftime("%Y-%m-%d")
		#print(d_date)
		days2 = ((todays_date-datarooms_date).days)+2
		# #print("date-------------2",days2)
		# #print("this method is called ")
		# #print(str(todays_date)+"   date "+str(first_date) +"days "+str(days))
		if days >=days2:
			days3 = days2
		else:
			days3= days
		# end lines 

		data = []
		for i in reversed(range(0,days3)):
			da = {}
			dates = todays_date - timedelta(days=i)
			da['dataroom_views'] = DataroomView.objects.filter(dataroom_id=pk,created_date__day=dates.day, created_date__month=dates.month, created_date__year = dates.year).count()
			da['documents_view'] = FolderView.objects.filter(dataroom_id=pk,created_date__day=dates.day, created_date__month=dates.month, created_date__year = dates.year).count()
			da['documents_print'] = FolderPrint.objects.filter(dataroom_id=pk,created_date__day=dates.day, created_date__month=dates.month, created_date__year = dates.year).count()
			da['documents_download'] = FolderDownload.objects.filter(dataroom_id=pk,created_date__day=dates.day, created_date__month=dates.month, created_date__year = dates.year).count()
			# #print("End-----", todays_date, dates, dates.day, dates.month, dates.year)
			da['date'] = dates.strftime('%d/%m/%Y')
			data.append(da)
		# #print("data", data) 
		data.reverse()
		return Response(data, status=status.HTTP_201_CREATED)



class ExportActivityByFile(APIView):
	"""docstring for ActivitybyDateReport"""
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		import datetime
		from_date = request.GET.get("from_date")
		to_date = request.GET.get("to_date")
		deleted = False if request.GET.get("deleted") == 'false' else True
		# #print("deleted",deleted)
		todays_date = datetime.datetime.strftime(datetime.datetime.strptime( to_date, '%Y-%m-%d'),'%Y-%m-%d 23:59:59+05:30')
		first_date = datetime.datetime.strftime(datetime.datetime.strptime( from_date, '%Y-%m-%d'),'%Y-%m-%d 00:00:00+05:30')
		# #print("request get", first_date, todays_date)
		document = DataroomFolder.objects.filter(dataroom_id = pk, is_folder=False, is_deleted=deleted, created_date__gte=first_date, created_date__lte=todays_date)
		serializer = DataroomFolderSerializer(document, many=True)
		data = serializer.data
		data.sort(key=lambda item:item['created_date'], reverse=True)
		from . import utils
		import csv

		for da in data:
			da['index'] = utils.getIndexes(da)
		response = HttpResponse(content_type='text/csv')
		response['Content-Disposition'] = 'attachment; filename="Activity.csv"'
		writer = csv.writer(response)
		header_data, datas = utils.getExcelDataActivityByDate(data)
		
		writer.writerow(header_data)
		writer.writerows(datas)
		return response
		# return Response(data, status=status.HTTP_201_CREATED)

class ExportActivitybyDateReport(APIView):
	"""docstring for ActivitybyDateReport"""
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		import datetime
		from_date = request.GET.get("from_date")
		to_date = request.GET.get("to_date")
		todays_date = datetime.datetime.strptime( to_date, '%Y-%m-%d')
		first_date = datetime.datetime.strptime( from_date, '%Y-%m-%d')
		days = ((todays_date-first_date).days)+1
		 #lines addded by harish 
		datarooms_date = Dataroom.objects.get(id=pk).created_date
		dateobject=datetime.datetime.strptime(str(datarooms_date),'%Y-%m-%d %H:%M:%S.%f')
		d_date =dateobject.strftime("%Y-%m-%d")
		#print(d_date)
		days2 = ((todays_date-datarooms_date).days)+2
		# #print("date-------------2",days2)
		# #print("this method is called ")
		# #print(str(todays_date)+"   date "+str(first_date) +"days "+str(days))
		if days >=days2:
			days3 = days2
		else:
			days3= days
		# end lines 
		data = []
		for i in reversed(range(0,days3)):
			da = {}
			dates = todays_date - timedelta(days=i)
			da['dataroom_views'] = DataroomView.objects.filter(dataroom_id=pk,created_date__day=dates.day, created_date__month=dates.month, created_date__year = dates.year).count()
			da['documents_view'] = FolderView.objects.filter(dataroom_id=pk,created_date__day=dates.day, created_date__month=dates.month, created_date__year = dates.year).count()
			da['documents_print'] = FolderPrint.objects.filter(dataroom_id=pk,created_date__day=dates.day, created_date__month=dates.month, created_date__year = dates.year).count()
			da['documents_download'] = FolderDownload.objects.filter(dataroom_id=pk,created_date__day=dates.day, created_date__month=dates.month, created_date__year = dates.year).count()
			# #print("End-----", todays_date, dates, dates.day, dates.month, dates.year)
			da['date'] = dates.strftime('%d/%m/%Y')
			data.append(da)
		# #print("data", data)
		data.reverse
		from . import utils
		import csv

		response = HttpResponse(content_type='text/csv')
		response['Content-Disposition'] = 'attachment; filename="Activity.csv"'
		writer = csv.writer(response)
		
		header_data, datas = utils.getActivityByDateExport(data)
		
		writer.writerow(header_data)
		writer.writerows(datas)
		return response
		
class GetIndexReports(APIView):
	authentication_classes = ( TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		index_list = IndexDownload.objects.filter(dataroom_id=pk)
		serializer = IndexDownloadSerializer(index_list, many=True)
		data = serializer.data
		data.sort(key=lambda item:item['created_date'], reverse=True)
		return Response(data, status=status.HTTP_201_CREATED)

	def delete(self, request, pk, format=None):
		index = IndexDownload.objects.get(id=pk)
		index.delete()
		return Response(status=status.HTTP_204_NO_CONTENT)


class RecentDocument(APIView):
	authentication_classes = (TokenAuthentication,)
	permission_classes = (IsAuthenticated,)

	def get(self, request, pk, format=None):
		folder_data = DataroomFolder.objects.filter(dataroom_id=pk).order_by('-created_date') #('updated_date')
		folder_data = folder_data.order_by('-updated_date')
		journal = User.objects.all().order_by('-date_joined')
		# Rajendra code, Combine two queryset from diff model
		report = chain(folder_data, journal)
		# report = report.order_by('-updated_date', '-date_joined', '-created_date')
		qs_json = sez.serialize('json', report)  
		# #print("Tango:----------------------", qs_json)
		d = json.loads(qs_json)
		for each in d:
			# #print("Tango:----------------------", each, "\n")
			try:
				if each['fields']['name']:
					user = User.objects.get(id=each['fields']['user'])
					each['fields']['username'] = user.username
			except:
				each['fields']['name'] = ''
		cat_list = Categories.objects.all()
		return Response({"folder_data": d})

class GetAllFilesViewed(APIView):
	authentication_classes = ( TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		file_list = FolderView.objects.filter(dataroom_id=pk, folder__is_folder=False)
		serializer = FolderViewSerializer(file_list, many=True)
		data = serializer.data
		data.sort(key=lambda item:item['created_date'], reverse=True)
		return Response(data, status=status.HTTP_201_CREATED)

class GetAllFilesPrinted(APIView):
	authentication_classes = ( TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		print_list = FolderPrint.objects.filter(dataroom_id=pk, folder__is_folder=False)
		serializer = FolderPrintSerializer(print_list, many=True)
		data = serializer.data
		data.sort(key=lambda item:item['created_date'], reverse=True)
		return Response(data, status=status.HTTP_201_CREATED)

class GetAllFilesDownloaded(APIView):
	authentication_classes = ( TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		import datetime
		user = request.user
		from_date = request.GET.get("from_date")
		to_date = request.GET.get("to_date")
		todays_date = datetime.datetime.strftime(datetime.datetime.strptime( to_date, '%Y-%m-%d'),'%Y-%m-%d 23:59:59+05:30')
		first_date = datetime.datetime.strftime(datetime.datetime.strptime( from_date, '%Y-%m-%d'),'%Y-%m-%d 00:00:00+05:30')
		download_list = FolderDownload.objects.filter(dataroom_id=pk, folder__is_folder=False, created_date__gte=first_date, created_date__lte=todays_date)
		serializer = FolderDownloadSerializer(download_list, many=True)
		data = serializer.data
		data.sort(key=lambda item:item['created_date'], reverse=True)
		return Response(data, status=status.HTTP_201_CREATED)

class ExportFilesDownloaded(APIView):
	authentication_classes = ( TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		import datetime
		user = request.user
		from_date = request.GET.get("from_date")
		to_date = request.GET.get("to_date")
		todays_date = datetime.datetime.strftime(datetime.datetime.strptime( to_date, '%Y-%m-%d'),'%Y-%m-%d 23:59:59+05:30')
		first_date = datetime.datetime.strftime(datetime.datetime.strptime( from_date, '%Y-%m-%d'),'%Y-%m-%d 00:00:00+05:30')
		download_list = FolderDownload.objects.filter(dataroom_id=pk, folder__is_folder=False, created_date__gte=first_date, created_date__lte=todays_date)
		serializer = FolderDownloadSerializer(download_list, many=True)
		data = serializer.data
		data.sort(key=lambda item:item['created_date'], reverse=True)
		from . import utils
		import csv

		response = HttpResponse(content_type='text/csv')
		response['Content-Disposition'] = 'attachment; filename="Activity.csv"'
		writer = csv.writer(response)
		
		header_data, datas = utils.getExcelDataDownloadedFiles(data)
		
		writer.writerow(header_data)
		writer.writerows(datas)
		return response


class RecentDocument(APIView):
	"""
	Retrieve, update or delete a snippet instance.
	"""
	authentication_classes = (TokenAuthentication,)
	permission_classes = (IsAuthenticated,)

	def get(self, request, pk, format=None):
		is_uploaded = 0
		is_updated = 0
		is_deleted = 0
		folder_data = DataroomFolder.objects.filter(dataroom_id=pk).order_by('-created_date') #('updated_date')
		folder_data = folder_data.order_by('-updated_date')
		journal = User.objects.all().order_by('-date_joined')
		from itertools import chain
		report = chain(folder_data, journal)
		qs_json = sez.serialize('json', report)        
		d = json.loads(qs_json)
		# #print("dataa", d)
		for each in d:
			# #print("Tango:----------------------", each, "\n")
			try:
				if each['fields']['name']:
					user = User.objects.get(id=each['fields']['user'])
					each['fields']['username'] = user.username
			except:
				each['fields']['name'] = ''
			try:
				if each['fields']['is_deleted']:
					is_deleted += 1 
			except:
				pass
			try:
				if each['fields']['is_uploaded']:
					is_uploaded += 1 
			except:
				pass
			try:
				if not each['fields']['is_uploaded']:
					is_updated += 1 
			except:
				pass

		recent_update = RcentUpdate.objects.filter(dataroom_id=pk).order_by('-created_at')
		recent_update = recent_update.order_by('-modified_at')
		recent = chain(recent_update,)
		# #print("dataa", recent)
		recent_json = sez.serialize('json', recent)
		# #print("dataa", recent_json)        
		recent_data = json.loads(recent_json)
		
		# #print("Else part here.", d)
		# print ("is_uploaded", is_uploaded)
		# print ("is_updated", is_updated)
		# print ("is_deleted", is_deleted)
		return Response({"folder_data": d, "is_deleted": is_deleted, 'is_uploaded': is_uploaded, 'is_updated': is_updated, 'recent_data': recent_data})

class FoldersList(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		folder = DataroomFolder.objects.filter(dataroom_id=pk, is_deleted=False, is_folder=False)
		serializer = DataroomFolderSerializer(folder, many=True)
		return Response(serializer.data, status=status.HTTP_201_CREATED)


class UpdateList(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		if user.is_superadmin == True or user.is_admin == True:
			update = RcentUpdate.objects.filter(dataroom_id=pk).order_by('-id')

		else:
			update = RcentUpdate.objects.filter(dataroom_id=pk, member__in=[user.id]).order_by('-id')
		serializer = RecentUpdateSerializer(update, many=True)
		data = serializer.data
		for da in data:
			da['member_details'] = []
			if da.get('file'):
				folder = DataroomFolder.objects.get(id = da.get('file'))
				da['file'] = DataroomFolderSerializer(folder, many=False).data
				if user.is_superadmin == True or user.is_admin == True:
					da['file']['file_view'] = True
				else:
					try:
						member = DataroomMembers.objects.filter(member_id=user.id, dataroom_id=pk).first()
						perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(folder_id=da.get('file').get('id'),dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id).first()
						da['file']['file_view'] = file_list.is_view_only
					except:
						da['file']['file_view'] = False
			if da.get('member'):
				# #print(da.get('member'))
				member_array = []
				for member_id in da.get('member'):
					# #print(member_id)
					members = DataroomMembers.objects.filter(member_id=member_id, dataroom_id=pk).first()
					serializer = DataroomMembersSerializer(members, many=False)
					# #print(serializer.data)
					if serializer.data and 'member' in serializer.data:
						# #print(1)
						member_array.append(str(serializer.data.get('member').get('first_name')))
				# #print(member_array)
				da['member_details'] = member_array

		return Response(data, status=status.HTTP_201_CREATED)


class NewUpdate(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		update = RcentUpdate.objects.get(id=pk)
		serializer = RcentUpdateSerializer(update, many=False)
		return Response(serializer.data, status=status.HTTP_201_CREATED)

	def post(self, request, pk, format=None):
		user = request.user
		# #print('#####################################################################################')
		# print ("user:-----", user)
		# #print('#####################################################################################')

		data = request.data

		# #print('#####################################################################################')
		# print ("data:-----", data)
		# #print('#####################################################################################')

		myfile = request.FILES.getlist('file')
		# print ("data:------", data)
		member_list = []
		group_list = []
		categorie_list = []
		try:
			for cat in data['categories']:
				categories = Categories.objects.filter(categories_name=cat).first()
				categorie_list.append(categories.id)
		except:
			for cat in data['categories']:
				categories = Categories.objects.filter(id=cat).first()
				categorie_list.append(categories.id)
		try:
			for each in data['member']:
				member = User.objects.filter(email=each).first()
				member_list.append(member.id)
		except:
			for each in data['member']:
				member = DataroomGroups.objects.filter(id=each).first()
				group_list.append(member.id)
	   
		# import ast
		# data = ast.literal_eval(json.dumps(data))
		# #print("member_list", member_list, "group_list", group_list)
		data['categories'] = []
		# data['member'] = []
		data['user'] = user.id
		data['dataroom'] = int(pk)
		# print ("data is ", data)
		# message = get_template('data_documents/new_updates.html').render(ctx)
		# send_email_to_members(data['groups',])
		if data['send_update_email'] == True:
			# #print('&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&')
			subject =data['subject']
			from_email = settings.EMAIL_HOST_USER
			to =data['member']
			# #print(to[0])
			text="@" in str(to[0])
			# #print(text)
			# if to[0]__icontains < 5 :
				# #print('9876')
			if text:
				to = to
			else :
				userss = DataroomMembers.objects.filter(end_user_group__in=to,dataroom_id =pk, is_deleted=False)
				# #print("@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@")
				# #print(group_list)
				# #print(userss)
				to = []
				for user in userss:
					to.append(user.member.email)
					# #print("to------------",to)
			# #print(to)
			dataroomname = Dataroom.objects.get(id=pk)
			message=data['message']
			# ctx = {
			#     'email': to,
			#     'subject': subject,
			#     'messages':message,

			#     # 'link': link,

			#     # 'update':update
			#         }

			# message = get_template('data_documents/new_updates.html').render(ctx)

			# msg = EmailMessage(subject, message, to=to, from_email=from_email)
			# msg.content_subtype = 'html'
			# msg.send()

		for i in to:
			name =User.objects.get(email=i)

			ctx = {
					'email': i,
					'subject': subject,
					'messages':message,
					'name':name.first_name,
					'dataroom':dataroomname.dataroom_name,
					# 'update':update
						}
			# #print("******************************************************")
			# #print(i)
			# #print('name -----------------------')
			# #print('666666666666666666666666666666666666666666666666666')
			ne = [i]
			# #print(ne)


			message = get_template('data_documents/new_updates.html').render(ctx)

			msg = EmailMessage(subject, message, to=ne, from_email=from_email)
			msg.content_subtype = 'html'
			msg.send()

		data.pop('member', None)
		data.pop('categories', None)
		serializer = RecentUpdateSerializer(data=data, context={'member': member_list, 'groups': group_list, 'categories': categorie_list})
		if serializer.is_valid():
		   serializer.save()

		   return Response(serializer.data, status=status.HTTP_201_CREATED)        
		# print ("error:---", serializer.errors)
		return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)    


class CategoriesListApi(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		cate_list = Categories.objects.filter(dataroom_id=pk)
		cat_serializer = CategoriesSerializer(cate_list, many=True)
		data = cat_serializer.data
		for da in data:
			try:
				da['user'] = ManageDataroomCategories.objects.filter(category_id=da['id'], dataroom_id=da['dataroom']).values_list('user_id', flat=True)
			except:
				da['user'] = []
		return Response(data, status=status.HTTP_201_CREATED)


class DataroomUserList(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user_id = []
		dataroom_list = Dataroom.objects.filter(id=pk).first()
		dataroom_user_list = DataroomMembers.objects.filter(dataroom_id=pk)
		for dataroom_user in dataroom_user_list:
			if (dataroom_user.is_la_user or dataroom_user.is_dataroom_admin) and dataroom_user.is_q_a_user:
				user_id.append(dataroom_user.member.id)
		user_list = User.objects.filter(id__in=user_id)
		# #print("user_id_10_april ===>",user_list)
		serializer = UserSerializer(user_list, many=True)
		# cat_serializer = CategoriesSerializer(cate_list, many=True)
		return Response(serializer.data, status=status.HTTP_201_CREATED)


class DataroomMemList(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	# def get(self, request, pk, format=None):
	#     import pdb;pdb.set_trace();
	#     user_id = []
	#     dataroom_user_list = DataroomMembers.objects.filter(dataroom_id=pk)
	#     serializer = DataroomMembersSerializer(dataroom_user_list, many=True)
	#     return Response(serializer.data, status=status.HTTP_201_CREATED)

	def post(self, request, pk, format=None):
		# #print("Request Data", request.data, request.user.id)
		user_id = []
		dataroom_user_list = DataroomMembers.objects.filter(dataroom_id=pk, member_id=request.user.id)
		serializer = DataroomMembersSerializer(dataroom_user_list, many=True)
		return Response(serializer.data, status=status.HTTP_201_CREATED)


class GetRecentUpdates(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )
	
	def post(self, request, pk, format=None):
		cat_id = []
		categories_obj = request.data['cat_list']
		for cat in categories_obj:
			cat_id.append(cat['id'])
		recent_update = RcentUpdate.objects.filter(dataroom_id=pk, categories__in=cat_id).order_by('-created_at').distinct()
		recent_update = recent_update.order_by('-modified_at')
		recent = chain(recent_update,)
		recent_json = sez.serialize('json', recent)        
		recent_data = json.loads(recent_json)
		return Response({'recent_data': recent_data})


class CopyFiles(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def assign_index(self, data , user, pk):
		index = 0
		try :
			# print ("data under assign_index is:", data)
			# print ("priamry key is", pk)
			max_index = DataroomFolder.objects.filter(parent_folder_id=pk, dataroom_id=data.get('dataroomId'), is_root_folder=False, is_folder=False, is_deleted=False).aggregate(Max('index'))
			# print ("max indx is", max_index)
			index = max_index.get('index__max')
			# print ("index is", index)
			if index is None:
				index = 0
		except :
			pass
		# print ("new index is", index+1)
		return index+1

	def assign_index_folder(self, data , user, pk):
		index = 0
		try :
			# print ("inside create new folder", data)
			max_index = DataroomFolder.objects.filter(dataroom_id=data.get('dataroomId'), parent_folder_id=pk, is_root_folder=False, is_folder=True, is_deleted=False).aggregate(Max('index'))
			index = max_index.get('index__max')
			# print ('max index now is', index)
			if index is None:
				index = 0
		except :
			pass
		# print ("index after +1 is", index+1)
		return index+1

	def copy_folder(self, files, filed, user):
		data = {}
		for file in files:
			copy_file = DataroomFolder()
			copy_file.user_id = user.id
			copy_file.dataroom_id = file.dataroom_id
			copy_file.name = file.name
			copy_file.last_updated_user_id = user.id
			copy_file.parent_path = file.parent_path
			copy_file.file_content_type = file.file_content_type
			copy_file.file_size = file.file_size
			copy_file.parent_folder_id = filed.id
			copy_file.pages = file.pages
			copy_file.version = file.version
			copy_file.path = file.path
			copy_file.is_folder = file.is_folder
			data['dataroomId'] = file.dataroom_id 
			if file.is_folder == True:
				copy_file.index = self.assign_index_folder(data, user, filed.id)

			else:
				copy_file.index = self.assign_index(data, user, filed.id)
			copy_file.save()
			if file.is_folder == True:
				files = DataroomFolder.objects.filter(parent_folder_id=file.id)
				self.copy_folder(files, copy_file, user)
		return True

	def post(self, request, format=None):
		data = {}
		user = request.user
		# #print("requesttttttt", request.data)
		file_obj = DataroomFolder.objects.filter(id=request.data['file_id']).first()
		folder_obj = DataroomFolder.objects.filter(id=request.data['folder_id']).first()
		
		copy_file = DataroomFolder()
		copy_file.user_id = user.id
		copy_file.dataroom_id = file_obj.dataroom_id
		copy_file.name = file_obj.name
		copy_file.last_updated_user_id = request.user.id
		copy_file.parent_path = file_obj.parent_path
		copy_file.file_content_type = file_obj.file_content_type
		copy_file.file_size = file_obj.file_size
		copy_file.parent_folder_id = folder_obj.id
		copy_file.pages = file_obj.pages
		copy_file.version = file_obj.version
		copy_file.is_folder = file_obj.is_folder
		if file_obj.is_root_folder == True:
			copy_file.is_root_folder = False 
		copy_file.path = file_obj.path
		data['dataroomId'] = file_obj.dataroom_id 
		if file_obj.is_folder == True:
			copy_file.index = self.assign_index_folder(data, user, folder_obj.id)
		else:
			copy_file.index = self.assign_index(data, user, folder_obj.id)
		copy_file.save()
		if file_obj.is_folder == True:
			files = DataroomFolder.objects.filter(parent_folder_id=file_obj.id)
			self.copy_folder(files, copy_file, user)
		data = request.data
		return Response({'recent_data': True}, status=status.HTTP_201_CREATED)


class MoveFiles(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def assign_index(self, data , user, pk):
		index = 0
		try :
			# print ("data under assign_index is:", data)
			# print ("priamry key is", pk)
			max_index = DataroomFolder.objects.filter(parent_folder_id=pk, dataroom_id=data.get('dataroomId'), is_root_folder=False, is_folder=False, is_deleted=False).aggregate(Max('index'))
			# print ("max indx is", max_index)
			index = max_index.get('index__max')
			# print ("index is", index)
			if index is None:
				index = 0
		except :
			pass
		# print ("new index is", index+1)
		return index+1

	def assign_index_folder(self, data , user, pk):
		index = 0
		try :
			# print ("inside create new folder", data)
			max_index = DataroomFolder.objects.filter(dataroom_id=data.get('dataroomId'), parent_folder_id=pk, is_root_folder=False, is_folder=True, is_deleted=False).aggregate(Max('index'))
			index = max_index.get('index__max')
			# print ('max index now is', index)
			if index is None:
				index = 0
		except :
			pass
		# print ("index after +1 is", index+1)
		return index+1

	def post(self, request, format=None):
		data = {}
		user = request.user
		file_obj = DataroomFolder.objects.filter(id=request.data['file_id']).first()
		folder_obj = DataroomFolder.objects.filter(id=request.data['folder_id']).first()
		
		copy_file = DataroomFolder()
		copy_file.user_id = user.id
		copy_file.dataroom_id = file_obj.dataroom_id
		copy_file.name = file_obj.name
		copy_file.last_updated_user_id = request.user.id
		copy_file.parent_path = file_obj.parent_path
		copy_file.file_content_type = file_obj.file_content_type
		copy_file.file_size = file_obj.file_size
		copy_file.parent_folder_id = folder_obj.id
		copy_file.pages = file_obj.pages
		copy_file.version = file_obj.version
		copy_file.path = file_obj.path
		data['dataroomId'] = file_obj.dataroom_id 
		copy_file.is_folder = file_obj.is_folder
		if file_obj.is_root_folder == True:
			copy_file.is_root_folder = False 
		if file_obj.is_folder == True:
			copy_file.index = self.assign_index_folder(data, user, folder_obj.id)
		else:
			copy_file.index = self.assign_index(data, user, folder_obj.id)
		copy_file.dataroom_folder_uuid = file_obj.dataroom_folder_uuid
		change_all_indexes(file_obj)
		copy_file.save()
		if file_obj.is_folder == True:
			DataroomFolder.objects.filter(parent_folder_id=file_obj.id).update(parent_folder_id=copy_file.id)
		file_obj.delete()
		return Response({'recent_data': True}, status=status.HTTP_201_CREATED)

class ChangeRootIndex(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):

		temp = 0
		data = []
		dataroomModel = Dataroom.objects.filter()
		datroomseralizer = DataroomSerializer(dataroomModel, many=True)

		if datroomseralizer.data:

			for k in datroomseralizer.data:

				#print(k.get('id'))
				pk = k.get('id')
				temp +=1

				folder_obj = DataroomFolder.objects.filter(dataroom_id=pk, is_root_folder=True, is_folder=True, is_deleted=False).order_by('index')
				serializer = DataroomFolderSerializer(folder_obj, many=True)
				# #print(serializer.data)
				j=0
				if serializer.data:
					for i in serializer.data:
						# #print(i.id)
						j+=1
						#print(i.get('id'))
						DataroomFolder.objects.filter(id=i.get('id')).update(index=j)

		return Response({'recent_data': temp, 'data':datroomseralizer.data}, status=status.HTTP_201_CREATED)
	 
class NotifyotherMembers(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, format=None):
		return Response({'recent_data': True}, status=status.HTTP_201_CREATED)

class ManageCategory(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def post(self, request, pk, format=None):
		data = request.data['update']
		# #print("data", data)
		for each in data:
			if each['category']:
				data = {}
				user = request.user
				categ_list = ManageDataroomCategories.objects.filter(dataroom_id=pk, category_id=each['categories']).values_list('user_id', flat=True)
				for usr in each['user']:
					if usr not in categ_list:
						ManageDataroomCategories.objects.create(user_id=usr, dataroom_id=pk, category_id=each['categories'])
					else:
						pass
				for categ in categ_list:
					if categ not in each['user']:
						ManageDataroomCategories.objects.filter(dataroom_id=pk, category_id=each['categories'], user_id=categ).delete()
					
			else:
				cat = Categories.objects.create(categories_name=each['category_name'], dataroom_id=pk)
				if each['user']:
					for usr in each['user']:
						ManageDataroomCategories.objects.create(user_id=usr, dataroom_id=pk, category_id=cat.id)
					
		return Response({"result": "Added category successfully!"}, status=status.HTTP_201_CREATED)


@api_view(('GET',))
@renderer_classes((TemplateHTMLRenderer, ))
# @authentication_classes((TokenAuthentication,))
@permission_classes((AllowAny, ))
def file_view(request, pk, pk1):
	# import pdb;pdb.set_trace()
	# #print("dataa --------", pk, pk1)
	print("this api is calling")
	activity_list = []
	context_data = {}
	from constants import constants
	base_url = "http://52.172.204.103:8000"
	frontend_url = "http://52.172.204.103"
	# frontend_url = "https://services.docullyvdr.com/projectName"
	data = DataroomFolder.objects.filter(id=pk).first()
	
	folder =  data
	conversion_path=str(data.path.url)
	print(conversion_path,"con path")
	path = str(data.path.url)
	#print('path',path)
	dpath = str(data.path.url)
	#print('dpath', dpath)
	docs = False
	odf = False
	csv = False
	pdf = False
	excel = False
	psd = False
	pptx = False
	png = False
	message = ''
	from . import utils
	user = User.objects.get(id=pk1)
	# #print("Group Folder", group_folder)
	try:
		member = DataroomMembers.objects.filter(member_id=user.id, dataroom_id=data.dataroom.id).first()
	except:
		member = None
	if request.accepted_renderer.format == 'html':
		# #print("dataa --------", data.path)
		if folder.name.endswith('.odt') or folder.name.endswith('.pdf') or folder.name.endswith('.ods') or folder.name.endswith('.odp'):
			odf = True
			path = str(folder.path.url)
			# #print("path_file_view ======>",path)
			dpath = path 
			if folder.name.endswith('.pdf'):
				pdf = True
		elif folder.name.endswith('.csv'):
			csv = True
			path = str(folder.path.url)
			dpath = path
		elif folder.name.endswith('.xlsx') or folder.name.endswith('.xls'):
			excel = True
			path = str(folder.path.url)
			# #print('path_xlsx ====>',path)
			dpath = path
			# if folder.conversion_path == '':
			#     conversion_path = utils.saveConversionPath(folder)
			#     folder.conversion_path = conversion_path
			#     folder.save()
			# conversion_path =  str(folder.conversion_path.url)
			# #print(conversion_path,99999)
			# if str(conversion_path)=='/media/This%20file%20is%20not%20supported':
			#     message ='This file is not supported.'
			# path = str(folder.path.url)
			# dpath = path
			# #print("conversion_path", conversion_path, member.is_dataroom_admin, member.is_la_user)

			# if member.is_end_user:
			#     if conversion_path.endswith('.pdf'):
			#         odf = True
			#         excel = False

		# 
		elif folder.name.endswith('.docx') or folder.name.endswith('.doc'):
			print("in docs condition")
			docs = True
			path = str(folder.path.url)
			dpath = path
		elif folder.name.endswith('.ppt') or folder.name.endswith('.pptx'):
			pptx = True
			path = str(folder.path.url)
			dpath = path
			# docx_text_list = []
			# file = urlopen(dpath).read()
			# file = BytesIO(file)
			# document = ZipFile(file)
			# content = document.read('word/document.xml')
			# word_obj = BeautifulSoup(content.decode('utf-8'),features="html.parser")
			# text_document = word_obj.findAll('w:t')
			# for t in text_document:
			#     # docx_text = t.text
			#     docx_text_list.append(t.text)
			# listToStr = ' '.join([str(elem) for elem in docx_text_list])
			# context_data['docx_text'] = listToStr
			# docx = BytesIO(requests.get(dpath).content)
			# text = docx2txt.process(docx)
			# #print("text======>",text)
			# from docx import Document
			# document = Document(dpath)
			# for p in document.paragraphs:
			#     #print("ptext ====>",p.text)

			# r = requests.get(folder.path.url)
			# #print(type(r.text),"value of r ====>",len(r.text),r.content)
			# response = HttpResponse(r.content,content_type="application/ms-word")
			# response['Content-Disposition'] = 'attachment; filename='+ filename
			# #print("value of response =====>",response)
			# return response
			# if folder.conversion_path == '':
			#     conversion_path = utils.saveConversionPath(folder)
			#     folder.conversion_path = conversion_path
			#     folder.save()
			# path = str(folder.conversion_path.url)
			# dpath = str(folder.path.url)
			# conversion_path = str(folder.conversion_path.url
		elif folder.name.endswith('.psd'):
			psd = True
		elif folder.name.endswith('.png'):
			# #print("entered in png file today")
			png = True
			path = str(folder.path.url)
			dpath = path
			# #print(png,"============",dpath)
		else:
			path = str(folder.path.url)
			dpath = path
		# #print("docssssssssss",docs, "odf", odf, "csvvvvvvv", csv)
		parent_list = []
		data.parent_folders = []
		get_root_folder = False
		folder_id =folder.parent_folder_id
		while (get_root_folder == False):
		   if folder.is_root_folder == False:
			   folder_obj = DataroomFolder.objects.get(id = folder_id)
			   folder_serializer = DataroomFolderSerializer(folder_obj, many=False)
			   parent_list.append(folder_serializer.data)
			   folder_id = folder_obj.parent_folder_id
			   folder = folder_obj
		   else:
			   get_root_folder = True
		data.parent_folders = parent_list
		data.parent_folders.reverse()
		folders_list = []
		all_folders = DataroomFolder.objects.filter(dataroom_id=data.dataroom.id).values('id','name', 'parent_folder_id')
		for allf in all_folders:
			if allf['parent_folder_id'] == None:
				allf['parent_folder_id'] = ""
			folders_list.append(allf)
		

		if member.is_dataroom_admin == True or member.is_la_user==True:
			user.is_la_user = True
			is_upload = True
			is_download = True
			is_shortcut = True
			is_print = True
			view_list = FolderView.objects.filter(folder_id=data.id, dataroom_id= data.dataroom.id)

			### FolderActivityViewSerializer
			view_serializer = FolderActivityViewSerializer(view_list, many=True).data
			activity_list.extend(view_serializer)
			download_list = FolderDownload.objects.filter(folder_id=data.id, dataroom_id=data.dataroom.id)

			#### FolderActivityDownloadSerializer

			download_serializer = FolderDownloadSerializer(download_list, many=True).data
			activity_list.extend(download_serializer)
			print_list = FolderPrint.objects.filter(folder_id=data.id, dataroom_id=data.dataroom.id)

			#### FolderActivityPrintSerializer
			print_serializer = FolderActivityPrintSerializer(print_list, many=True).data
			activity_list.extend(print_serializer)
			# #print("file_view_if")
		else:
			# #print("file_view_else")
			group_folder = DataroomGroupFolderSpecificPermissions.objects.filter(folder_id=pk,dataroom_id=data.dataroom.id, dataroom_groups_id=member.end_user_group.first().id).first()    
			# #print("------", group_folder.folder_id, group_folder.dataroom_groups_id, group_folder.is_view_and_print_and_download, group_folder.is_view_and_print)
			is_upload = group_folder.is_upload
			is_download = group_folder.is_view_and_print_and_download
			is_shortcut = group_folder.is_shortcut
			is_print = group_folder.is_view_and_print
			# #print("Member----",member.end_user_group.first().id, pk)
			group_perm = DataroomGroupPermission.objects.filter(dataroom_id=data.dataroom_id, dataroom_groups_id=member.end_user_group.first().id).first()
			# #print("group_perm==>", group_perm.is_watermarking)
			if group_perm.is_watermarking == True:
				# #print("watermarking permission is True")
				group_folder = DataroomGroupFolderSpecificPermissions.objects.filter(folder_id=pk,dataroom_id=data.dataroom.id, dataroom_groups_id=member.end_user_group.first().id).first()
				watermark = Watermarking.objects.filter(dataroom_id=data.dataroom_id)
				# #print("type ========>",Watermarking.objects.filter(dataroom_id=data.dataroom_id).values())
				# #print("watermark ==>",Watermarking.objects.filter(dataroom_id=data.dataroom_id).values())
				# watermark_data = WatermarkingSerializer(watermark,many=True).data
				# #print("========>",watermark_data)
				from dataroom import pdf_watermarking
				# #print("path", "dpath")
				from userauth import utils
				ip = utils.get_client_ip(request)
				# #print("ip====>",ip)
				if conversion_path.endswith('.pdf'):
					# #print("conversion_path ====>",conversion_path)
					from azure.storage.blob import BlockBlobService, PublicAccess
					block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
					container_name ='docullycontainer'
					path_name = conversion_path.split("/")
					print(path_name,"ravi1")
					print( path_name[-2]+'/'+path_name[-1],"ravi2")
					print(path_name[-1],"ravi3")
					block_blob_service.get_blob_to_path(container_name, path_name[-2]+'/'+path_name[-1], path_name[-1])
					new_path = "/home/cdms_backend/cdms2/"+ path_name[-1]
					# #print("new_path======>",new_path)
					# pdf_watermarking.GeneratePDF(watermark, group_folder, new_path, ip, user)
					# path = str(group_folder.watermarking_file.url)
					# #print("value of path =====>",path)
					# #print("group_perm.is_doc_as_pdf", group_perm.is_doc_as_pdf)
					# if dpath.endswith('.pdf'):
					#     dpath = str(group_folder.watermarking_file.url)
					#     #print("Grouppp", group_perm.is_excel_as_pdf)
					# if group_perm.is_doc_as_pdf == True or group_perm.is_excel_as_pdf == True:
						# dpath =str(group_folder.watermarking_file.url)
						# #print("Grouppp", group_perm.is_excel_as_pdf)
				elif conversion_path.endswith('.png') or conversion_path.endswith('.jpg'):
					ip = utils.get_client_ip(request)

			# except:
			#     is_upload = False
			#     is_download = False
			#     is_print = False
		# #print("is_download", is_download, is_upload, is_print)

		qna_list = []
		if member.is_dataroom_admin == True or member.is_la_user==True:
			# #print(data.dataroom.id,"<==============>",pk)
			qna_list_obj = QuestionAnswer.objects.filter(dataroom_id=data.dataroom.id, folder_id=pk)
			# qna_list_obj = QuestionAnswer.objects.filter()
			# #print("qna_list_obj=====>",qna_list_obj)
		else:
			qna_list_obj = QuestionAnswer.objects.filter(dataroom_id=data.dataroom.id,folder_id=pk, user_id=pk1)
		for each in qna_list_obj:
			# if not each.answer or each.final:
			if not each.answer:
				# print ("answer:-", each.answer, "/", "final:-", each.final)
				qna_list.append(each)

		category_list = Categories.objects.filter(dataroom_id=data.dataroom.id).exclude(category_details__user=None)
		cat_serializer = CategoriesSerializer(category_list, many=True)
		qna_serializer = QuestionFileViewerAnswerSerializer(qna_list, many=True)
		for qna in qna_serializer.data:
			quest = QuestionAnswer.objects.filter(id=qna['id']).first()
			quest_categ = ManageDataroomCategories.objects.filter(category_id = quest.category_id, dataroom_id=quest.dataroom_id).values_list('user_id', flat=True)
			categ = ManageDataroomCategories.objects.filter(dataroom_id=quest.dataroom_id).values_list('user_id', flat=True)
			# #print("========>",quest.dataroom_id,"pk1==>",pk1)
			member = DataroomMembers.objects.filter(dataroom_id=quest.dataroom_id,member_id=pk1).first()
			# member = DataroomMembers.objects.filter(dataroom_id=quest.dataroom_id).first()
			# #print(member.is_q_a_user, categ, user.id)
			# #print("memeber ====>",member)
			if member.is_q_a_user == True and  user.id not in categ:
				# #print("not in")
				reply = True
			elif user.id == quest.user.id or user.id in quest_categ:
				reply = True
				# #print("in")
			else:
				reply = False
			qna['reply'] = reply
		# is_end_user
		# if docs == True:
		#     is_download == False
		# import os, uuid, sys
		# from azure.storage.blob import BlockBlobService, PublicAccess

		# try:
		# Create the BlockBlockService that is used to call the Blob service for the storage account
			# block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
			# container_name ='docullycontainer'
			# #print("value of path =====>",path)
			# path_split = path.split("/")
			# #print("path_split ===>",path_split[-2],path_split[-1])
			# # new_name = '/'+ path_split[-2]+'/'+path_list[-1]
			# # #print("value of blob name ===>",new_name)
			# file_size = BlockBlobService.get_blob_properties(block_blob_service,container_name,path_split[-2]+'/'+path_split[-1]).properties.content_length
			# #print("value of file_size ====>",file_size)
			# if file_size > 0:
			#     context_data['path'] = path+'#toolbar=0'
			# else:
			#     context_data['path'] = ''

			# #print("value of file_size===>",file_size)
			# if block_blob_service.exists(container_name, path_list[-1]):
			#     #print("true")
			# else:
			#     pass
			# https://docullystorage.blob.core.windows.net/docullycontainer/test_dataroom_17_March_33da3675-0289-4480-9350-386ef8f42644/test_1.pdf
		# except:
			# #print("errors")
		filnn=extension11[0].split("/")
		print(filnn,"filnn")
		# pdf_filename=str(filnn[-1])+".pdf"
		# print(pdf_filename,"pdf_filename")
		# blobname=filnn[-2]+"/"+pdf_filename
		# print(blobname,"blobname")
		dataroom_user_list = DataroomMembers.objects.filter(dataroom_id=data.dataroom_id, member_id=pk1).first()
		context_data['category_list'] = cat_serializer.data
		context_data['qna_list'] = qna_serializer.data#sorted(qna_serializer.data, key=lambda k: k['created_date'], reverse=True)
		context_data['path'] = path+'#toolbar=0'
		context_data['dpath'] = dpath
		context_data['all_folders'] = all_folders
		context_data['folder'] = data
		context_data['base_url'] = base_url
		context_data['frontend_url'] = frontend_url
		context_data['docs'] = docs
		context_data['odf'] = odf
		context_data['pdf'] = pdf
		context_data['csv'] = csv
		context_data['psd'] = psd
		context_data['pptx'] = pptx
		context_data['png'] = png
		context_data['url'] = "https://services.docullyvdr.com"+'/projectName/file-view/'+str(pk)+'/'+str(pk1)+'/'
		context_data['excel'] = excel
		context_data['all_folders'] = folders_list
		context_data['is_upload']  = is_upload
		context_data['is_download']  = is_download
		context_data['is_shortcut']  = is_shortcut
		context_data['is_print']  = is_print
		context_data['login_user'] = pk1
		context_data['message'] = message
		context_data['dataroom_user_list'] = dataroom_user_list
		context_data['user'] = user
		context_data['refpath'] = 

		# #print("user_file_viewer ===>",user)
		getusername(user)
		# #print("List",activity_list)
		# if docx_text is not None:
		#     context_data['docx_text'] = listToStr
		#     #print("docx_text_list ====>",docx_text_list)
		# else:
		#     context_data['docx_text'] = ""



		# #print("docx_text_list ====>",docx_text_list)
		# #print("context_data============>",context_data)
		activity_list.sort(key=lambda r: r['created_date'], reverse=True)
		context_data['activity_tracker'] = activity_list
	   # print("List11",activity_list)
		# print(context_data,"context_data")
		if member.is_dataroom_admin == True or member.is_la_user==True:
			return render(request, 'file-view.html', context_data )
		else:
			print(pk1,"m in else file view")
			view_list = FolderView.objects.filter(folder_id=data.id, dataroom_id= data.dataroom
				)
			view_serializer = FolderActivityViewSerializer(view_list, many=True).data
			activity_list.extend(view_serializer)
			context_data['activity_tracker'] = activity_list
			print(context_data['activity_tracker'],"enddddd user")
			return render(request,'file-view.html', context_data )

def getusername(self):
	# #print("getusername======>",self.user)
	return

from django.views.generic import TemplateView

class IframeView(TemplateView):
	template_name = 'iframe.html'

	def get_context_data(self, **kwargs):
		context = super(IframeView, self).get_context_data(**kwargs)
		# import pdb;pdb.set_trace()
		context['path'] = kwargs.get('path')
		return context


@api_view(('POST',))
# @renderer_classes((TemplateHTMLRenderer, ))
# @authentication_classes((TokenAuthentication,))
@permission_classes((AllowAny, ))
def upload_file_version(request, pk):
	data = request.data
	# #print("data--------8_april==>",data)
	file = request.FILES.getlist('file0')[0]
	# #print("file--------",file.name)
	folder = DataroomFolder.objects.filter(id=pk).first()
	folder.name = file.name
	# #print("value of file_size ===>",file.size)
	folder.file_size_mb = file.size
	folder.file_content_type = file.content_type
	folder.path = file
	folder.save()
	from . import utils
	if file.name.endswith('.docx') or file.name.endswith('.doc') or file.name.endswith('.ppt') or file.name.endswith('.pptx') or file.name.endswith('.pdf') or folder.name.endswith('.xlsx') or folder.name.endswith('.xls') or folder.name.endswith('.csv') or folder.name.endswith('.odt') or folder.name.endswith('.ods') or folder.name.endswith('.odp'):
		conversation_path = utils.saveConversionPath(folder)
		folder.conversion_path = conversation_path
		folder.save()
	data = {'msg':'successfully updated'}
	return HttpResponse(data,content_type='application/javascript')
   

@permission_classes((AllowAny, ))
def get_print_download_logs(request, pk):
	# print ("get all recent document data", request.GET)    
	update_view = request.GET.get("type")
	user = request.GET.get("user")
	folder = DataroomFolder.objects.get(id=pk)
	print("folder", folder.path, user)
	if update_view == 'view':
		folder.is_view = True
		folder.save()
		view = FolderView()
		view.folder_id = folder.id
		view.user_id = int(user)
		view.dataroom_id = folder.dataroom_id
		view.save()
	elif update_view == 'print':
		folder.is_print = True
		folder.save()
		print_obj = folder#print()
		print_obj.folder_id = folder.id
		print_obj.user_id = int(user)
		print_obj.dataroom_id = folder.dataroom_id
		print_obj.save()
	elif update_view == 'download':
		folder.is_download = True
		folder.save()
		download = FolderDownload()
		download.folder_id = folder.id
		download.user_id = int(user)
		download.dataroom_id = folder.dataroom_id
		download.save()
	else:
		pass
	data = {'msg':'successfully updated'}
	return HttpResponse(data,content_type='application/javascript')



@renderer_classes((TemplateHTMLRenderer, ))
@permission_classes((AllowAny, ))
def file_comment(request):
	if request.method == "POST" and request.is_ajax():
		# import pdb;pdb.set_trace();
		context_data = {}
		category_id = request.POST['category_id']
		que_title = request.POST['que_title']
		file_id = request.POST['file_id']
		dataroom_id = request.POST['dataroom_id']
		login_user_id = request.POST['login_user_id']
		qna_obj = QuestionAnswer.objects.create(user_id=login_user_id, dataroom_id=dataroom_id, question=que_title, folder_id=file_id, category_id=category_id)
		qna_list = QuestionAnswer.objects.filter(dataroom_id=dataroom_id, folder_id=file_id)
		category_list = Categories.objects.all()
		cat_serializer = CategoriesSerializer(category_list, many=True)
		qna_serializer = QuestionAnswerSerializer(qna_list.reverse(), many=True)
		s_qna_serializer = QuestionAnswerSerializer(qna_list.reverse().first(),)
		context_data['category_list'] = cat_serializer.data
		context_data['single_qna'] = s_qna_serializer.data
		context_data['qna_list'] = sorted(qna_serializer.data, key=lambda k: k['created_date'], reverse=True)
		category_obj = ManageDataroomCategories.objects.filter(category_id=category_id, dataroom_id=dataroom_id).first()
		from . import utils
		utils.send_mail_to_coordinator(qna_obj, category_obj)
		return HttpResponse(json.dumps(context_data), content_type='application/json')
		# return JsonResponse(context_data)


def file_comment_reply(request):
	if request.method == "POST" and request.is_ajax():
		context_data = {}
		qna_answer = request.POST['qna_answer']
		qna_answer_id = request.POST['qna_answer_id']
		login_user_id = request.POST['login_user_id']
		qna_obj = QuestionAnswer.objects.filter(id=qna_answer_id).first()
		qna_created = QuestionAnswer.objects.create(user_id=login_user_id, answer=qna_answer, question=qna_obj.question, dataroom=qna_obj.dataroom, qna_id=qna_answer_id, folder=qna_obj.folder, category=qna_obj.category, final=True)
		# qna_serializer = QuestionAnswerSerializer(qna_list.reverse(), many=True)
		qna_list = []
		qna_list_obj = QuestionAnswer.objects.filter(dataroom_id=qna_obj.dataroom.id)
		for each in qna_list_obj:
			if not each.answer or each.final:
				# print ("answer:-", each.answer, "/", "final:-", each.final)
				qna_list.append(each)

		category_list = Categories.objects.filter(dataroom_id=qna_obj.dataroom.id)
		cat_serializer = CategoriesSerializer(category_list, many=True)
		qna_serializer = QuestionAnswerSerializer(qna_list, many=True)
		# is_end_user
		dataroom_user_list = DataroomMembers.objects.filter(dataroom_id=qna_obj.dataroom_id, member_id=login_user_id).first()
		context_data['category_list'] = cat_serializer.data
		context_data['qna_list'] = qna_serializer.data#sorted(qna_serializer.data, key=lambda k: k['created_date'], reverse=True)
		return HttpResponse(json.dumps(context_data), content_type='application/json')

class GetAllActivityofDocuments(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		view_list = FolderView.objects.filter(dataroom_id=pk, user_id = user.id)
		view_serializer = FolderViewSerializer(view_list, many=True)
		data = view_serializer.data
		for da in data:
			da['activity'] = 'Viewed'

		print_list = FolderPrint.objects.filter(dataroom_id=pk, user_id = user.id)
		print_serializer = FolderPrintSerializer(print_list, many=True)
		print_data = print_serializer.data
		for prnt in print_data:
			prnt['activity'] = "Printed"
			data.append(prnt)

		download_list = FolderDownload.objects.filter(dataroom_id=pk, user_id = user.id)
		download_serializer = FolderDownloadSerializer(download_list, many=True)
		download_data = download_serializer.data
		for dwnld in download_data:
			dwnld['activity'] = "Downloaded"
			data.append(dwnld)
		data.sort(key=lambda r: r['created_date'], reverse=True)
		return Response(data, status=status.HTTP_201_CREATED)

class CopySelectedFiles(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def assign_index(self, data , user, pk):
		index = 0
		try :
			# print ("data under assign_index is:", data)
			# print ("priamry key is", pk)
			max_index = DataroomFolder.objects.filter(parent_folder_id=pk, dataroom_id=data.get('dataroomId'), is_root_folder=False, is_folder=False, is_deleted=False).aggregate(Max('index'))
			# print ("max indx is", max_index)
			index = max_index.get('index__max')
			# print ("index is", index)
			if index is None:
				index = 0
		except :
			pass
		# print ("new index is", index+1)
		return index+1

	def assign_index_folder(self, data , user, pk):
		index = 0
		try :
			# print ("inside create new folder", data)
			max_index = DataroomFolder.objects.filter(dataroom_id=data.get('dataroomId'), parent_folder_id=pk, is_root_folder=False, is_folder=True, is_deleted=False).aggregate(Max('index'))
			index = max_index.get('index__max')
			# print ('max index now is', index)
			if index is None:
				index = 0
		except :
			pass
		# print ("index after +1 is", index+1)
		return index+1

	def copy_folder(self, files, filed, user):
		data = {}
		for file in files:
			copy_file = DataroomFolder()
			copy_file.user_id = user.id
			copy_file.dataroom_id = file.dataroom_id
			copy_file.name = file.name
			copy_file.last_updated_user_id = user.id
			copy_file.parent_path = file.parent_path
			copy_file.file_content_type = file.file_content_type
			copy_file.file_size = file.file_size
			copy_file.parent_folder_id = filed.id
			copy_file.pages = file.pages
			copy_file.version = file.version
			copy_file.path = file.path
			copy_file.is_folder = file.is_folder
			data['dataroomId'] = file.dataroom_id 
			if file.is_folder == True:
				copy_file.index = self.assign_index_folder(data, user, filed.id)

			else:
				copy_file.index = self.assign_index(data, user, filed.id)
			copy_file.save()
			if file.is_folder == True:
				files = DataroomFolder.objects.filter(parent_folder_id=file.id)
				self.copy_folder(files, copy_file, user)
		return True

	def post(self, request, pk, format=None):
		data = {}
		user = request.user
		# #print("requesttttttt", request.data)
		for da in request.data:
			file_obj = DataroomFolder.objects.filter(id=da['id']).first()
			folder_obj = DataroomFolder.objects.filter(id=pk).first()
			
			copy_file = DataroomFolder()
			copy_file.user_id = user.id
			copy_file.dataroom_id = file_obj.dataroom_id
			copy_file.name = file_obj.name
			copy_file.last_updated_user_id = request.user.id
			copy_file.parent_path = file_obj.parent_path
			copy_file.file_content_type = file_obj.file_content_type
			copy_file.file_size = file_obj.file_size
			copy_file.parent_folder_id = folder_obj.id
			copy_file.pages = file_obj.pages
			copy_file.version = file_obj.version
			copy_file.path = file_obj.path
			if file_obj.is_root_folder == True:
				copy_file.is_root_folder = False
			copy_file.is_folder = file_obj.is_folder
			data['dataroomId'] = file_obj.dataroom_id 
			if file_obj.is_folder == True:
				copy_file.index = self.assign_index_folder(data, user, folder_obj.id)
			else:
				copy_file.index = self.assign_index(data, user, folder_obj.id)

			copy_file.save()
			if file_obj.is_folder == True:
				files = DataroomFolder.objects.filter(parent_folder_id=file_obj.id)
				self.copy_folder(files, copy_file, user)
		data = request.data
		return Response({'recent_data': True}, status=status.HTTP_201_CREATED)

class MoveSelectedFiles(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def assign_index(self, data , user, pk):
		index = 0
		try :
			# print ("data under assign_index is:", data)
			# print ("priamry key is", pk)
			max_index = DataroomFolder.objects.filter(parent_folder_id=pk, dataroom_id=data.get('dataroomId'), is_root_folder=False, is_folder=False, is_deleted=False).aggregate(Max('index'))
			# print ("max indx is", max_index)
			index = max_index.get('index__max')
			# print ("index is", index)
			if index is None:
				index = 0
		except :
			pass
		# print ("new index is", index+1)
		return index+1

	def assign_index_folder(self, data , user, pk):
		index = 0
		try :
			# print ("inside create new folder", data)
			max_index = DataroomFolder.objects.filter(dataroom_id=data.get('dataroomId'), parent_folder_id=pk, is_root_folder=False, is_folder=True, is_deleted=False).aggregate(Max('index'))
			index = max_index.get('index__max')
			# print ('max index now is', index)
			if index is None:
				index = 0
		except :
			pass
		# print ("index after +1 is", index+1)
		return index+1

	def post(self, request, pk, format=None):
		data = {}
		user = request.user
		for da in request.data:
			file_obj = DataroomFolder.objects.filter(id=da['id']).first()
			folder_obj = DataroomFolder.objects.filter(id=pk).first()
			
			copy_file = DataroomFolder()
			copy_file.user_id = user.id
			copy_file.dataroom_id = file_obj.dataroom_id
			copy_file.name = file_obj.name
			copy_file.last_updated_user_id = request.user.id
			copy_file.parent_path = file_obj.parent_path
			copy_file.file_content_type = file_obj.file_content_type
			copy_file.file_size = file_obj.file_size
			copy_file.parent_folder_id = folder_obj.id
			copy_file.pages = file_obj.pages
			copy_file.version = file_obj.version
			copy_file.path = file_obj.path
			data['dataroomId'] = file_obj.dataroom_id
			copy_file.is_folder = file_obj.is_folder 
			if file_obj.is_root_folder == True:
				copy_file.is_root_folder = False
			# #print("file_obj", file_obj.is_folder )
			if file_obj.is_folder == True:
				copy_file.index = self.assign_index_folder(data, user, folder_obj.id)
			else:
				copy_file.index = self.assign_index(data, user, folder_obj.id)
			copy_file.dataroom_folder_uuid = file_obj.dataroom_folder_uuid
			change_all_indexes(file_obj)
			copy_file.save()
			if file_obj.is_folder == True:
				DataroomFolder.objects.filter(parent_folder_id=file_obj.id).update(parent_folder_id=copy_file.id)
			file_obj.delete()
		return Response({'recent_data': True}, status=status.HTTP_201_CREATED)
			
class ExportOverviewReport(APIView):
	"""docstring for ActivitybyDateReport"""
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		import datetime
		data = []
		from_date = request.GET.get("from_date")
		to_date = request.GET.get("to_date")
		#print("from_date",from_date,to_date)
		todays_date = datetime.datetime.strftime(datetime.datetime.strptime( to_date, '%Y-%m-%d'),'%Y-%m-%d 23:59:59+05:30')
		first_date = datetime.datetime.strftime(datetime.datetime.strptime( from_date, '%Y-%m-%d'),'%Y-%m-%d 00:00:00+05:30')
		#print('<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<')
		#print("request get", first_date, todays_date)

		folder_view = FolderView.objects.filter(dataroom_id=pk,created_date__gte=first_date, created_date__lte=todays_date)
		for fold in folder_view:
			data.append({'created_date':fold.created_date,'event':'Viewed', 'title':fold.folder.name, 'user':fold.user.first_name+fold.user.last_name})
		folder_print = FolderPrint.objects.filter(dataroom_id=pk,created_date__gte=first_date, created_date__lte=todays_date)
		for fold in folder_print:
			data.append({'created_date':fold.created_date,'event':'Printed', 'title':fold.folder.name, 'user':fold.user.first_name+fold.user.last_name})
		folder_download = FolderDownload.objects.filter(dataroom_id=pk,created_date__gte=first_date, created_date__lte=todays_date)
		for fold in folder_download:
			data.append({'created_date':fold.created_date,'event':'Downloaded', 'title':fold.folder.name, 'user':fold.user.first_name+fold.user.last_name})
		
		folder = DataroomFolder.objects.filter(dataroom_id = pk, is_deleted=False,created_date__gte=first_date, created_date__lte=todays_date)
		for fold in folder:
			if fold.is_folder == True:
				data.append({'created_date':fold.created_date,'event':'Created', 'title':fold.name, 'user':fold.user.first_name+fold.user.last_name})
			else:
				data.append({'created_date':fold.created_date,'event':'Uploaded', 'title':fold.name, 'user':fold.user.first_name+fold.user.last_name})

		folder_deleted = DataroomFolder.objects.filter(dataroom_id = pk, is_deleted=True, is_folder=False,created_date__gte=first_date, created_date__lte=todays_date)
		for fold in folder_deleted:
			data.append({'created_date':fold.deleted_by_date,'event':'Deleted', 'title':fold.name, 'user':fold.deleted_by.first_name+fold.deleted_by.last_name})
			
		data.sort(key=lambda item:item['created_date'], reverse=True)
		datarooms = Dataroom.objects.filter(id=pk)
		serializer = DataroomSerializer(datarooms, many=True)
		file_name = str(serializer.data[0].get('dataroom_name'))+' - Overview Report - '+str(time.strftime("%d-%m-%Y_%H:%M:%S"))+'.csv'
		from . import utils
		import csv
		response = HttpResponse(content_type='text/csv')
		response['Content-Disposition'] = 'attachment; filename='+str(file_name)+''
		writer = csv.writer(response)
		
		header_data, datas = utils.getExcelOverviewData(data)
		
		writer.writerow(header_data)
		writer.writerows(datas)
		return response

def build_folder_recursive(tree, parent, nodes, user, pk, group_id):
	children  = [n for n in nodes if n['parent_folder'] == parent]
	for child in children:
		# #print("Childdddddddd-------", child)
		data = {'perm':{}}
		data['name'] = child['name']
		data['id'] = child['id']
		data['is_folder'] = child['is_folder']
		data['noaccess_indeterminate'] = False
		data['view_indeterminate'] = False
		data['vp_indeterminate'] = False
		data['vpd_indeterminate'] = False
		data['upload_indeterminate'] = False
		data['drm_indeterminate'] = False
		data['watermarking_indeterminate'] = False
		data['editor_indeterminate'] = False
		data['shortcut_indeterminate'] = False
		try:
			perm_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=data['id'],dataroom_id=pk, dataroom_groups_id=int(group_id))
			data['perm']['is_view_only'] = perm_obj.is_view_only
			data['perm']['is_no_access'] = perm_obj.is_no_access
			data['perm']['is_view_and_print'] = perm_obj.is_view_and_print
			data['perm']['is_view_and_print_and_download'] = perm_obj.is_view_and_print_and_download
			data['perm']['is_upload'] = perm_obj.is_upload
			data['perm']['is_watermarking'] = perm_obj.is_watermarking
			data['perm']['is_drm'] = perm_obj.is_drm
			data['perm']['is_editor'] = perm_obj.is_editor
			data['perm']['is_access'] = perm_obj.is_access
			data['perm']['is_shortcut'] = perm_obj.is_shortcut
		except:
			data['perm']['is_view_only'] = False
			data['perm']['is_no_access'] = True
			data['perm']['is_access'] = False
			data['perm']['is_view_and_print'] = False
			data['perm']['is_view_and_print_and_download'] = False
			data['perm']['is_upload'] = False
			data['perm']['is_watermarking'] = False
			data['perm']['is_drm'] = False
			data['perm']['is_editor'] = False
			data['perm']['is_shortcut'] = False
		data['children'] = []
		tree.append(data)
		build_folder_recursive(data['children'], child['id'], nodes, user, pk, group_id)

class FoldersHierarchyPermission(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		group_id = request.GET.get('group_id')
		# #print("Group Id", group_id)
		document = DataroomFolder.objects.filter(dataroom_id = pk, is_root_folder=True, is_deleted=False).order_by('index')
		data = []
		from . import utils
		for doc in document:
			docu = DataroomFolder.objects.get(id = doc.id)
			docu_serializer = DataroomFolderSerializer(docu)
			datas = docu_serializer.data
			# utils.getIndexofFolder(datas)
			data.append(datas)
			docu1 = DataroomFolder.objects.filter(parent_folder = doc.id, is_deleted=False).order_by('index')
			if len(docu1) > 0:
				datas = []
				data.extend(utils.get_under_file_withoutindex(docu1,datas))
				# #print("dataaa", data)
		# #print("dataaaaa", data)
		tree = []
		# fill in tree starting with roots (those with no parent)
		build_folder_recursive(tree, None, data, user, pk, group_id)
		# #print("treee", tree)
		return Response(tree, status=status.HTTP_201_CREATED)


class TreeViewPermission(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		group_id = request.GET.get('group_id')
		parent_id = request.GET.get('parent')
		#print(parent_id)
		# #print("Group Id", group_id)
		if int(parent_id)==0:
			document = DataroomFolder.objects.filter(dataroom_id = pk, is_root_folder=True, is_deleted=False).order_by('index')
			#print(3434)
		else:
			document = DataroomFolder.objects.filter(dataroom_id = pk, is_root_folder=False, parent_folder=parent_id, is_deleted=False).order_by('index')
			#print(9090)
		data = []
		from . import utils
		for doc in document:
			docu = DataroomFolder.objects.get(id = doc.id)
			docu_serializer = DataroomFolderSerializer(docu)
			datas = docu_serializer.data
			# utils.getIndexofFolder(datas)
			
			docu1 = DataroomFolder.objects.filter(parent_folder = doc.id, is_deleted=False).order_by('index')
			if datas['is_folder']:
				datas['hasChildren'] = True
			else:
				datas['hasChildren'] = False

			datas['perm'] = {}
			datas['noaccess_indeterminate'] = False
			datas['view_indeterminate'] = False
			datas['vp_indeterminate'] = False
			datas['vpd_indeterminate'] = False
			datas['upload_indeterminate'] = False
			datas['drm_indeterminate'] = False
			datas['watermarking_indeterminate'] = False
			datas['editor_indeterminate'] = False
			datas['shortcut_indeterminate'] = False
			# datas['children'] = []
			try:
				
				perm_obj = DataroomGroupFolderSpecificPermissions.objects.get(folder_id=datas['id'],dataroom_id=pk, dataroom_groups_id=int(group_id))
				datas['perm']['is_view_only'] = perm_obj.is_view_only
				datas['perm']['is_no_access'] = perm_obj.is_no_access
				datas['perm']['is_view_and_print'] = perm_obj.is_view_and_print
				datas['perm']['is_view_and_print_and_download'] = perm_obj.is_view_and_print_and_download
				datas['perm']['is_upload'] = perm_obj.is_upload
				datas['perm']['is_watermarking'] = perm_obj.is_watermarking
				datas['perm']['is_drm'] = perm_obj.is_drm
				datas['perm']['is_editor'] = perm_obj.is_editor
				datas['perm']['is_access'] = perm_obj.is_access
				datas['perm']['is_exist'] = True
				datas['perm']['is_shortcut'] = perm_obj.is_shortcut
			except:
				datas['perm']['is_view_only'] = False
				datas['perm']['is_no_access'] = True
				datas['perm']['is_access'] = False
				datas['perm']['is_view_and_print'] = False
				datas['perm']['is_view_and_print_and_download'] = False
				datas['perm']['is_upload'] = False
				datas['perm']['is_watermarking'] = False
				datas['perm']['is_drm'] = False
				datas['perm']['is_editor'] = False
				datas['perm']['is_exist'] = False
				datas['perm']['is_shortcut'] = False
			# if len(docu1) > 0:
				# datas = []
				# datas['hasChildren'] = True
			# else:
				# datas['hasChildren'] = True

			data.append(datas)
				# data.extend(utils.get_under_file_withoutindex(docu1,datas))
				# #print("dataaa", data)
		# #print("dataaaaa", data)
		# tree = []
		# fill in tree starting with roots (those with no parent)
		# build_folder_recursive(tree, None, data, user, pk, group_id)
		# #print("treee", tree)
		return Response(data, status=status.HTTP_201_CREATED)

class UserActivity(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, pk1, format=None):
		data = {}
		datas=[]
		folder_view = FolderView.objects.filter(dataroom_id = pk, user_id = pk1)
		folder_view_data = FolderViewSerializer(folder_view, many=True).data
		datas.extend(folder_view_data)
		folder_download = FolderDownload.objects.filter(dataroom_id=pk, user_id=pk1)
		folder_download_data = FolderDownloadSerializer(folder_download, many=True).data
		datas.extend(folder_download_data)
		workspace_view = DataroomView.objects.filter(dataroom_id=pk, user_id=pk1)
		folder_workspace_data = DataroomViewSerializer(workspace_view, many=True).data
		datas.extend(folder_workspace_data)
		member = DataroomMembers.objects.filter(dataroom_id=pk, member_added_by_id=pk1, is_deleted=False, is_deleted_end=False, is_deleted_la=False)
		member_data = DataroomMembersSerializer(member, many=True).data
		for member in member_data:
			member['created_date'] = member.get('date_joined')
			member['event'] = 'member joined'
			datas.append(member)
		
		data['folder_view_count']  = folder_view.count()
		data['folder_download_count']  = folder_download.count()
		data['warkspace_view_count'] = workspace_view.count()
		# #print("Activityyyyyy",datas)
		data['activity'] = sorted(datas, key=lambda k: k['created_date'], reverse=True)
		# data['activity'] = datas.sort(key=lambda item:item['created_date'], reverse=True)
		# #print("Activityyyyyy",data['activity'])
		return Response(data, status=status.HTTP_201_CREATED)
		
class DataroomActivity(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		data = {}
		datas=[]
		#print("request =====>",request.data)
		#print("pk ======>",pk)

		from_date = request.GET.get("from_date")
		to_date = request.GET.get("to_date")
		#print("dataroom_activity_api_called ====>",from_date,to_date)
		# need to work after the some time
		dateobject1=datetime.strptime(from_date,'%Y-%m-%d')
		first_date=dateobject1.strftime("%Y-%m-%d 00:00:00+05:30")
		dateobject2=datetime.strptime(to_date,'%Y-%m-%d')
		todays_date=dateobject2.strftime("%Y-%m-%d 23:59:59+05:30")
		#print("dataroom_activity_api_called ====>")

		# #print(first_date)
		# #print(todays_date)
		# first_date = datetime.datetime.strftime(datetime.datetime.strptime( from_date, '%Y-%m-%d'),'%Y-%m-%d 00:00:00+05:30')
		# #print('<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<')
		# #print("request get", first_date, todays_date)

		folder_view = FolderView.objects.filter(dataroom_id = pk,created_date__gte=first_date, created_date__lte=todays_date)
		folder_view_data = FolderViewSerializer(folder_view, many=True).data
		datas.extend(folder_view_data)
		folder_download = FolderDownload.objects.filter(dataroom_id=pk,created_date__gte=first_date, created_date__lte=todays_date)
		folder_download_data = FolderDownloadSerializer(folder_download, many=True).data
		datas.extend(folder_download_data)
		workspace_view = DataroomView.objects.filter(dataroom_id=pk,created_date__gte=first_date, created_date__lte=todays_date)
		folder_workspace_data = DataroomViewSerializer(workspace_view, many=True).data
		datas.extend(folder_workspace_data)
		# change made by harish feb 26
		folder = DataroomFolder.objects.filter(dataroom_id = pk, is_deleted=False,created_date__gte=first_date, created_date__lte=todays_date)
		folder_data = DataroomFolderSerializer(folder, many=True).data
		# datas.extend(folder_data)
		#print('data_file_activity_16_april-------------->',folder_data)
		for fold in folder_data:
			if fold['is_folder'] == True:
				fold['event']='Created'
			else:
				fold['event']='Uploaded'
			# fold['first_name'] = User.objects.filter(id=fold.get('user')).values_list('first_name',flat=True)
			name = User.objects.get(id=fold.get('user'))
			if name.first_name != '' or name.last_name !=None and name.last_name != '' or name.last_name != None :
				fold['first_name'] =name.first_name
				fold['last_name'] = name.last_name
			else :
				fold['first_name'] = 'user name is not defined '
			datas.append(fold)


		member = DataroomMembers.objects.filter(dataroom_id=pk, is_deleted=False, is_deleted_end=False, is_deleted_la=False)
		member_data = DataroomMembersSerializer(member, many=True).data
		for member in member_data:
			member['created_date'] = member.get('date_joined')
			member['event'] = 'member joined'
			datas.append(member)
		# #print("dataaaaaaaaa", datas)


		datas = sorted(datas,key=lambda x : x['created_date'], reverse=True)
		return Response(datas, status=status.HTTP_201_CREATED)


def build_tree_recursive_index_report(tree, parent, nodes):
	from . import utils
	children  = [n for n in nodes if n['parent_folder'] == parent]
	for child in children:
		# #print("Childdddddddd-------", child)
		data = {}
		data['name'] = child['name']
		data['id'] = child['id']
		if child['is_folder'] == True:
			data['type'] = 'Folder'
		else:
			data['type']  = 'File'
		data['path'] = str(child['parent_path'])
		data['date'] = child['created_date']
		data['size'] = child['file_size_mb'] if child['file_size_mb'] != None else 0
		data['index'] = utils.getIndexes(child)
		data['children'] = []
		tree.append(data)
		build_tree_recursive_index_report(data['children'], child['id'], nodes)

class IndexReportUserWise(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, pk1, format=None):
		user = request.user
		member = DataroomMembers.objects.filter(member_id=pk1, dataroom_id=pk).first()
		if member.is_dataroom_admin or member.is_la_user:
			document = DataroomFolder.objects.filter(dataroom_id = pk, is_root_folder=True, is_deleted=False).order_by('index')
		else:
			perm_obj = DataroomGroupFolderSpecificPermissions.objects.filter(dataroom_id=pk, dataroom_groups_id=member.end_user_group.first().id, is_view_only=True).values('folder_id')
			document = DataroomFolder.objects.filter(id__in = perm_obj,dataroom_id = pk, is_root_folder=True, is_deleted=False).order_by('index')
		data = []
		from . import utils
		for doc in document:
			docu = DataroomFolder.objects.get(id = doc.id)
			docu_serializer = DataroomFolderSerializer(docu)
			datas = docu_serializer.data
			datas['index'] = utils.getIndexes(datas)
			data.append(datas)
			if member.is_dataroom_admin or member.is_la_user:
				docu1 = DataroomFolder.objects.filter(parent_folder = doc.id, is_deleted=False).order_by('index')
			else:
				docu1 = DataroomFolder.objects.filter(parent_folder = doc.id, is_deleted=False, id__in = perm_obj).order_by('index')
			if len(docu1) > 0:
				datas = []
				data.extend(utils.get_under_file(docu1,datas))
		tree = []
		build_tree_recursive_index_report(tree, None, data)
		import csv
		response = HttpResponse(content_type='text/csv')
		response['Content-Disposition'] = 'attachment; filename="index.csv"'
		writer = csv.writer(response)
		users = User.objects.get(id=pk1)
		dataroom = Dataroom.objects.get(id=pk)
		header_data, datas = utils.getExcelIndexReportUserWise(tree,[])
		writer.writerow(["This index report is of "+users.first_name+" "+users.last_name+" ("+users.email+") from Dataroom "+dataroom.dataroom_name])
		writer.writerow(header_data)
		new_list = []
		for value in datas:
			# #print(type(value),"<===== new_data_for_csv====>",value)
			new_list.append(list(value))
		# #print("new_list==========>",new_list)
		for new_value in new_list:
			# #print("new_value=====>",new_value)
			if new_list[1][-1] is not None:
				new_list[1][-1] = dataroom.dataroom_name +'/'+ new_list[0][1] + '/' + new_list[1][-1]
		# #print("set===>",str(new_list[1][-1]).split('/'))
		new_path = new_list[1][-1].split('/')
		# #print("new_path=====>",new_path[-1],new_path[-2],new_path[-3])
		new_list[1][-1] = new_path[-3] +'/'+ new_path[-2]+'/'+new_path[-1]
		# #print("updated_list===>",new_list)        
		writer.writerows(new_list)
		# #print(datas[0][1],"2323=========>",type(datas),datas[1])
		return response

class DataroomMembersView(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		user = request.user
		members = DataroomMembers.objects.filter(dataroom_id=pk, is_deleted=False)
		serializer = DataroomMembersSerializer(members, many=True)
		return Response(serializer.data, status=status.HTTP_201_CREATED)

class GetSingleMemberView(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )
	def get(self, request, pk, format=None):
		user = request.user
		members = DataroomMembers.objects.filter(id=pk, is_deleted=False)
		serializer = DataroomMembersSerializer(members, many=True)
		return Response(serializer.data, status=status.HTTP_201_CREATED)

class DataroomFilesActivity(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, pk1, format=None):
		user = request.user
		data = []
		#print(pk,"pk====>",pk1)
		view = FolderView.objects.filter(dataroom_id=pk, folder_id=pk1)
		# #print("this_api====>",view)
		# #print("new_data ===>",FolderView.objects.filter(dataroom_id=pk,folder_id=pk1).values())
		view_serializer = FolderViewSerializer(view, many=True).data
		data.extend(view_serializer)
		download = FolderDownload.objects.filter(dataroom_id=pk, folder_id=pk1)
		download_serializer = FolderDownloadSerializer(download, many=True).data
		data.extend(download_serializer)
		prints = FolderPrint.objects.filter(dataroom_id=pk, folder_id=pk1)
		print_serializer = FolderPrintSerializer(prints, many=True).data
		data.extend(print_serializer)
		#print("data_file_activity=====>",data)
		import datetime
		# data.sort(key=lambda r: datetime.datetime.strptime(r['created_date'], '%Y/%m/%d %H:%M:%S'), reverse=True)
		# #print("value of data ====>",data)
		return Response(data, status=status.HTTP_201_CREATED)


class SendDocumentUpdate(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		document = DataroomFolder.objects.get(id=pk)
		# #print("document", document.path)
		utils.send_notify_to_all_members_regarding_uploaded_file(dataroom_file)
		return Response(status=status.HTTP_201_CREATED)


class bulkDownload(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def get(self, request, pk, format=None):
		data = [{'name':'sanket'}]
		return Response(data, status=status.HTTP_201_CREATED)
	def post(self, request, pk, format=None):
		data = json.loads(request.data)
		from azure.storage.blob import BlockBlobService
		import zipfile
		print(data,"data")
		foldername=data['data'][0]['name']
		print(foldername,"foldername")
		datarooms = Dataroom.objects.filter(id=pk).values('dataroom_name','dataroom_uuid')
		print(datarooms,"datarooms")
		folerofazure=datarooms[0]['dataroom_name']+"_"+str(datarooms[0]['dataroom_uuid'])
		print(folerofazure,"folerofazure")
		block_blob_service = BlockBlobService(account_name ='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
		CONTAINER_NAME='docullycontainer'
		generator = block_blob_service.list_blobs(CONTAINER_NAME)
		filelist=[]
		for blob in generator:
			if blob.name.startswith(folerofazure):
				# print("processing right")
				filelist.append(blob.name)
		print(filelist)
		zf = zipfile.ZipFile('/home/cdms_backend/cdms2/media/'+foldername+'.zip', mode='w',compression=zipfile.ZIP_DEFLATED)
		for i in filelist:
			print(i,"iii")
			b = block_blob_service.get_blob_to_bytes(CONTAINER_NAME, i)
			zf.writestr(i, b.content)
		# result = {"size":round(fsize, 2),
		# 	"name":str(file_name)+'.zip',
		# 	"path":'media/'+str(file_name)+'.zip'
		# 	}
		# return Response(result, status=status.HTTP_201_CREATED)
		zip_filess=(open('/home/cdms_backend/cdms2/media/'+foldername+'.zip','rb').read())
		response = HttpResponse(zip_filess,content_type="application/zip")
		filen=foldername+'.zip'
		response["Content-Disposition"] = b'attachment; filename=%s' % filen.encode(encoding="utf-8")
		return response

	# def post(self, request, pk, format=None	# 	from . import utils
	# 	data = json.loads(request.data)
	# 	print(data,"data json")
	# 	datarooms = Dataroom.objects.filter(id=pk)
	# 	serializer = DataroomSerializer(datarooms, many=True)
	# 	#print("time ====>",time.strftime("%d-%m-%Y_%H:%M:%S%p"))

	# 	file_name = str(serializer.data[0].get('dataroom_name'))+'-Multiple'+'-'+str(time.strftime("%d-%m-%Y_%H:%M:%S%p"))
	# 	#print("file_name =======>",file_name)
	# 	os.mkdir('./media/'+str(file_name))
	# 	directory_name = './media/'+''+str(file_name).replace(' ','\ ')+''

	# 	for file in data['data']:
	# 		print(file,"fileeeeeeeeeeeeeeeeeeeeeeeeee")
	# 		if print('is_folder' in file):
	# 			if file['is_folder']:
	# 				print("m in if")
	# 				directory = str(directory_name).replace('\ ',' ')+'/'+str(file['name'])
	# 				#print(directory)
	# 				os.mkdir(''+directory+'')
	# 				# break
	# 				utils.get_subfolder_list(file['id'], directory)
	# 				# break
	# 			if print(not file['is_folder']:
	# 				print("m in if not")
	# 				utils.copy(utils.split_link(str(file['path']).replace('%20','\ ')), str(directory_name)+'/')
	# 				# break
	# 		else:
	# 			if file['folder']['is_folder']:

	# 				try:
	# 					os.mkdir(str(directory_name)+'/'+str(file['folder']['name']).replace(' ','\ '))
	# 					# break
	# 					utils.get_subfolder_list(file['folder']['id'], str(directory_name)+'/'+str(file['folder']['name']))
	# 				except:
	# 					pass
	# 					#print("File Exist")
	# 				# break
	# 			if not file['folder']['is_folder']:
	# 				utils.copy(utils.split_link(str(file['folder']['path']).replace('%20','\ ')), str(directory_name)+'/')

	# 	file_paths_1 = []
	# 	directory_name_array = []
	# 	file_paths = []
	# 	directory_name_array = ['./'+str(file_name)+'']
	# 	os.system("mv -v "+str(directory_name)+" "+str('media/../'))

	# 	for file_name_1 in directory_name_array:
	# 		if os.path.isdir(file_name_1):
	# 			#print(file_name_1)
	# 			file_paths_1 = utils.get_all_file_paths(file_name_1)
	# 			for file in file_paths_1:
	# 				file_paths.append(file)
					
	# 	# zipFileName = utils.randomString2(8)
	# 	# writing files to a zipfile 
	# 	with ZipFile('media/'+str(file_name)+'.zip','w') as zip: 
	# 		# writing each file one by one 
	# 		for file in file_paths:
	# 			zip.write(file)
	# 	#print(file_name, 2323)
	# 	os.system("rm -r ./"+str(file_name).replace(' ','\ ')+"")
	# 	# fsize=os.stat('media'+str(file_name).split('.')[1]+'.zip')
	# # 	fsize=os.path.getsize('media/'+str(file_name)+'.zip')/float(1<<20)
	# 	result = {"size":round(fsize, 2),
	# 		"name":str(file_name)+'.zip',
	# 		"path":'media/'+str(file_name)+'.zip'
	# 		}
	# 	return Response(result, status=status.HTTP_201_CREATED)

# class shortCutDownload(APIView):
#     authentication_classes = (TokenAuthentication, )
#     permission_classes = (IsAuthenticated, )

#     def post(self, request, pk, format=None):

#         from constants import constants
#         base_url = constants.backend_ip

#         docu = DataroomFolder.objects.get(id = pk)
#         docu_serializer = DataroomFolderSerializer(docu)
#         datas = docu_serializer.data
#         file_name = datas.get('name').split('.')
#         iconFile = 'default.ico'
#         #print(datas.get('name').split('.')[0])
#         if datas.get('name'):
#             iconFile = file_name[len(file_name)-1]+'.ico'

#         if not request.user.is_icon_downloaded:
#             User.objects.filter(id=request.user.id).update(is_icon_downloaded=True)
#         file_name = str(request.user.id)+'_'+str(pk)+'.txt'

#         f= open(file_name,"w+")

#         url = 'URL='+str(base_url)+'/file-view/'+str(pk)+'/'+str(request.user.id)+'/'

#         data = ['[InternetShortcut]']
#         url_file = datas.get('name').split('.')[0].replace(' ','_')

#         icon = "IconFile=%USERPROFILE%\\Downloads\\"+str(url_file)+"\\"+str(iconFile)
#         icon = icon.replace(" ", "")
#         data.append(url)
#         data.append('IconIndex=0')
#         data.append(icon)

#         for i in data:
#              f.write("%s\r\n" %(i))
#         f.close()
#         os.system("rm -r ./media/shortcut/"+str(url_file))
#         os.mkdir("./media/shortcut/"+str(url_file))
#         #print(url_file)
#         os.system("mv "+file_name+" ./media/shortcut/"+str(url_file)+'/'+str(url_file)+".url")
#         try:
#             data = os.system("cp -r ./static/file_icon/"+str(iconFile)+" ./media/shortcut/"+str(url_file)+'/')
#             #print(data, 3815)
#         except(OSError, IOError):
#             os.system("cp -r ./static/file_icon/default.ico ./media/shortcut/"+str(url_file)+'/')

#         os.system("rm -r "+str(settings.MEDIA_ROOT)+'../'+str(url_file)+'/')

#         os.system("mv -v "+str(settings.MEDIA_ROOT)+"shortcut/"+str(url_file)+'/ '+str(settings.MEDIA_ROOT)+'../'+str(url_file)+'/')

#         path = str(settings.MEDIA_ROOT)+'../'+str(url_file)+'/'+str(url_file)+".url"
#         file_paths = []
#         file_paths.append('./media/../'+str(url_file)+'/'+str(url_file)+".url")
#         file_paths.append('./media/../'+str(url_file)+'/'+str(iconFile))
#         zipFile = url_file
#         os.system("rm -r "+str(settings.MEDIA_ROOT)+str(zipFile)+'.zip')
#         result = {}
#         if os.path.exists(path):

#             with ZipFile('media/'+str(zipFile)+'.zip','w') as zip: 
#                 # writing each file one by one 
#                 for file in file_paths:
#                     zip.write(file)

#             fsize=os.path.getsize('media/'+str(zipFile)+'.zip')/float(1<<20)

#             result = {"size":round(fsize, 2),
#             "name":str(settings.MEDIA_ROOT)+str(zipFile)+'.zip',
#             "path":'media/'+str(zipFile)+'.zip'
#             }
#         return Response(result, status=status.HTTP_201_CREATED)


# def downloadtxt(file_name):
#     #print("hit ==>",file_name)
#     file_name_split = file_name.split("/")
#     if os.path.exists(file_name_split[-1]):
#         #print("existed =====>")
#         with open(file_name_split[-1], 'rb') as fh:
#             #print("fh.read==>",fh.read())
#             response = HttpResponse(fh.read(), content_type='text/plain')
#             #print("response =====>",response)
#             response['Content-Disposition'] = 'inline; filename=' + file_name_split[-1]
#             # os.remove(url_file+'.url')
#             return response
		
class shortCutDownload(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )

	def post(self, request, pk, format=None):

		from constants import constants
		import os
		base_url = constants.backend_ip2

		docu = DataroomFolder.objects.get(id = pk)

		docu_serializer = DataroomFolderSerializer(docu)
		datas = docu_serializer.data
		file_name = datas.get('name').split('.')
		iconFile = 'default.ico'
		#print(datas.get('name').split('.')[0])
		if datas.get('name'):
			iconFile = file_name[len(file_name)-1]+'.ico'

		if not request.user.is_icon_downloaded:
			User.objects.filter(id=request.user.id).update(is_icon_downloaded=True)
		file_name = str(request.user.id)+'_'+str(pk)+'.txt'

		f= open(file_name,"w+")

		url = 'URL='+str(base_url)+'/file-view/'+str(pk)+'/'+str(request.user.id)+'/'

		data = ['[InternetShortcut]']
		url_file = datas.get('name').split('.')[0].replace(' ','_')
		# #print('url_file',url_file)
		# #print('url ------',url)
		url_file_ext = datas.get('name').split('.')
		# #print("url_file_ext =====>",url_file_ext[1])

		# icon_path = '/home/cdms_backend/cdms2/static/Doculink_Icon/'
		# for file_icon in os.listdir(icon_path):
		#     file_icon_type = file_icon.split('.')
		#     if file_icon_type[0] == url_file_ext[1]:
		#         #print("file_icon ==>",file_icon)
		#         icon_type = file_icon


		# icon = "IconFile=%USERPROFILE%\\Downloads\\"+str(url_file)+"\\"+"media"+"\\"+"new_icon_type.ico"
		# icon = icon.replace(" ", "")
		data.append(url)
		# data.append('IconIndex=0')
		# data.append(icon)

		for i in data:  
			f.write("%s\r\n" %(i))
		f.close()
		# #print("file_name ==>",file_name)
		# #print("url_file ===>",url_file)
		# #print("storage_path ===>",settings.MEDIA_ROOT)
		# #print("data =====>",type(data))
		# blob_list = container_client.list_blobs()
		# for blob in blob_list:
		#     #print("\t" + blob.name)
		new_filename = file_name.split('.')
		# icon_path = '/home/cdms_backend/cdms2/static/Doculink_Icon/'
		# for file_icon in os.listdir(icon_path):
		#     file_icon_type = file_icon.split('.')
		#     if file_icon_type[0] == url_file_ext[1]:
		#         #print("file_icon ==>",file_icon)
		#         icon_type = file_icon
		# #print("value of icon_type =====>",icon_type)
		# copy_file = os.system("cp -r /home/cdms_backend/cdms2/"+file_name+" /home/"+url_file+'.url')
		# #print("copy_file =====>",copy_file)
		# data = os.system("cp -r /home/cdms_backend/cdms2/static/Doculink_Icon/"+icon_type+" /home/cdms_backend/cdms2/"+file_name)
		# #print(data,"=========", 3815)
		# #print("directory ===>",os.getcwd())
		# if not os.path.exists("/home/"+url_file+"/"):
		#     os.mkdir("/home/"+url_file+"/")
		#     os.makedirs(directory, 0777)
		#     os.chmod("/home/", 777)
		#     os.chmod("/home/", 0o777)
		#     os.mkdir("/home/"+url_file+"/")
		# new_file_direct = os.mkdir("media"+"/"+url_file+'.url'
		# #print("new_file_direct =====>",new_file_direct)
		#print("filename ==>",file_name)
		print("/home/cdms_backend/cdms2/"+file_name,"before copy")
		print('/home/cdms_backend/cdms2/'+url_file+'.url',"after copy")
		shutil.copy2("/home/cdms_backend/cdms2/"+file_name, '/home/cdms_backend/cdms2/'+url_file+'.url')

		# file_blob_path = 'media'+'/'+url_file+'/'+url_file+'.url'
		# import os, uuid, sys
		# from azure.storage.blob import BlockBlobService, PublicAccess
		# from azure.storage.file import FileService
		# try:
		# Create the BlockBlockService that is used to call the Blob service for the storage account
			# block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
			# container_name ='docullycontainer'
			# blob_name = 'new_media'
			# top_level_container_name = "docullycontainer"
			# block_blob_service.create_container(container_name)
			# block_blob_service.set_container_acl(container_name, public_access=PublicAccess.Container)
			# block_blob_service.create_directory(container_name, 'sampledir')
			# block_blob_service.create_blob_from_path(container_name,blob_name)
			# generator = block_blob_service.list_blobs(top_level_container_name)
			# for blob in generator:
			#     #print("\t Blob name: " + blob.name)
			# new_filename = file_name.split('.')
			# file_blob_path = 'media'+'/'+url_file+'/'+url_file+'.url'
			# a = block_blob_service.create_blob_from_path(container_name,file_blob_path,'/home/cdms_backend/cdms2/'+file_name)
			# block_blob_service.create_blob_from_path(container_name,file_blob_path,'/home/cdms_backend/cdms2/'+url_file+'.url')
			# #print("value of blob =====>",url_file)
			# block_blob_service.create_blob_from_path(container_name,'media'+'/'+url_file+'/'+icon_type,'/home/cdms_backend/cdms2/static/Doculink_Icon/'+icon_type)
			# file_size = BlockBlobService.get_blob_properties(block_blob_service,container_name,'media'+'/'+ url_file+'/'+url_file+'.url').properties.content_length
			# #print("value of file_size ====>",file_size)
			# block_blob_service.get_blob_to_path(container_name,'media'+'/'+ url_file+'/'+url_file+'.url' , url_file+'.url')

			# if not os.path.exists(''):
			#     os.makedirs('my_folder')
			# shutil.copy2("/home/cdms_backend/cdms2/"+url_file+".url", 'media/'+url_file+"_"+".url")
			# #print("value of file_size =====>",file_size)
			# #print("value of filename ===>",file_name)

			# shutil.copy2("/home/cdms_backend/cdms2/static/Doculink_Icon/"+icon_type, 'media/new_icon_type.ico')
			# shutil.copy2("/home/cdms_backend/cdms2/"+file_name, 'media/'+url_file+'.url')
			# zip_file = url_file
			# file_paths = []
			# file_paths.append('media/new_icon_type.ico')
			# file_paths.append('media/'+ url_file+'.url')
			# #print("zip_file ======>",zip_file)
			# if file_size > 0:
			#     with ZipFile("media/"+str(url_file)+'.zip','w') as zip:
			#         for file in file_paths:
			#             zip.write(file)


			# block_blob_service.create_blob_from_path(container_name,'media'+'/'+ url_file + url_file+'.zip','media/'+str(url_file)+'.zip')
			# os.remove('/home/cdms_backend/cdms2/media/'+ url_file+'_'+'.url')
			# os.remove('/home/cdms_backend/cdms2/media/new_icon_type')
			# os.remove("/home/cdms_backend/cdms2/media/"+str(url_file)+'.zip')
			# block_blob_service.create_blob_from_text(container_name,file_name,'hello')
			# generator = block_blob_service.list_blobs(container_name)
			# fp = open(file_name, 'ab')
			# for blob in generator:
			#     #print("blob",blob)
			#     b = service.get_blob_to_bytes(container_name, blob.name)
			#     fp.write(b.content)
			# fp.flush()
			# fp.close()
		# except:
			# #print("errors")

		# os.system("cp  /home/cdms_backend/cdms2/static/Doculink_Icon/"+icon_type+" /media/"+icon_type)
		# os.system("cp" +url_file+'.url'+" /media/"+url_file+'.url')


		# https://docullystorage.blob.core.windows.net/quickstartblobs/219_13663.txt
		
		
		# file_paths.append('./media/../'+str(url_file)+'/'+str(iconFile))

		# zip_file = url_file
		# file_paths = []
		# file_paths.append('/home/cdms_backend/cdms2/static/Doculink_Icon/'+icon_type)
		# file_paths.append('/home/cdms_backend/cdms2/'+ file_name)
		# #print("zip_file ======>",zip_file)
		# file_paths = []
		# file_paths.append('./media/../'+str(url_file)+'/'+str(url_file)+".url")
		# file_paths.append('./media/../'+str(url_file)+'/'+str(iconFile))
		# zipFile = url_file
		# os.system("rm -r "+str(settings.MEDIA_ROOT)+str(zipFile)+'.zip')
		# result = {}
		# #print("============>")
		# if os.path.exists(path):
		#     #print("path exists ====>")
		#     with ZipFile('media/'+str(zipFile)+'.zip','w') as zip: 
		#         # writing each file one by one 
		#         for file in file_paths:
		#             zip.write(file)

		#     fsize=os.path.getsize('media/'+str(zipFile)+'.zip')/float(1<<20)

		#     result = {"size":round(fsize, 2),
		#     "name":str(settings.MEDIA_ROOT)+str(zipFile)+'.zip',
		#     "path":'media/'+str(zipFile)+'.zip'
		#     }

		# if file_size > 0:
		#     #print("true =====>")
		#     with ZipFile(str(url_file)+'.zip','w') as zip:
		#         for file in file_paths:
		#             zip.write(file)




		# r = requests.get("https://docullystorage.blob.core.windows.net/quickstartblobs/" + file_name)
		# response = HttpResponse(r.content,content_type='application/octet-stream')
		# response['Content-Disposition'] = 'inline; filename='+ file_name
		# # with open('new_'+file_name,'wb') as fh:
		# #     fh.write(r.content)
		# #print("response ====>",r.headers)

		# if os.path.exists(file_name):
		#     with open(file_name, 'rb') as fh:
		#         #print("file_found ====>",file_name)
		#         response = HttpResponse()
		#         response["Content-Disposition"]= "attachment; filename="+ url_file +'.url'
		#         #print("response ====>",response)
		#         return response
		#     raise Http404

		# file_path = "/home/cdms_backend/cdms2" +'/'+ url_file + '.url'
		# file_wrapper = FileWrapper(open(file_path,'rb'))
		# file_mimetype = mimetypes.guess_type(file_path)
		# response = HttpResponse(file_wrapper, content_type=file_mimetype )
		# response['X-Sendfile'] = file_path
		# response['Content-Length'] = os.stat(file_path).st_size
		# response['Content-Disposition'] = 'attachment; filename=%s' % url_file + '.url'
		# #print("response ====>",response) 
		# return response
		#print("file ====>",os.getcwd())
		# downloadtxt(file_name)
		file_name_split = file_name.split("/")
		if os.path.exists(file_name_split[-1]):
			#print("existed =====>")
			with open(file_name_split[-1], 'rb') as fh:
				print("fh.read==>",fh.read())
		os.remove(url_file+'.url')
		print(url,"url end in api")
		result = {"path": url_file + '.url',
		"text":'[InternetShortcut]'+'\n'+ url}
		return Response(result, status=status.HTTP_201_CREATED)

		# Set the permission so the blobs are public.
		# storage_path = "https://docullystorage.blob.core.windows.net/docullycontainer/"
		# os.system("rm -r "+ storage_path+"./media/shortcut/"+str(url_file))
		# os.mkdir(storage_path +"media/shortcut/"+str(url_file))

		
		# os.system("mv "+file_name+" ./media/shortcut/"+str(url_file)+'/'+str(url_file)+".url")
		# try:
		#     data = os.system("cp -r ./static/file_icon/"+str(iconFile)+" ./media/shortcut/"+str(url_file)+'/')
		#     #print(data, 3815)
		# except(OSError, IOError):
		#     os.system("cp -r ./static/file_icon/default.ico ./media/shortcut/"+str(url_file)+'/')

		# os.system("rm -r "+str(settings.MEDIA_ROOT)+'../'+str(url_file)+'/')

		# os.system("mv -v "+str(settings.MEDIA_ROOT)+"shortcut/"+str(url_file)+'/ '+str(settings.MEDIA_ROOT)+'../'+str(url_file)+'/')

		# path = str(settings.MEDIA_ROOT)+'../'+str(url_file)+'/'+str(url_file)+".url"
		# file_paths = []
		# file_paths.append('./media/../'+str(url_file)+'/'+str(url_file)+".url")
		# file_paths.append('./media/../'+str(url_file)+'/'+str(iconFile))
		# zipFile = url_file
		# os.system("rm -r "+str(settings.MEDIA_ROOT)+str(zipFile)+'.zip')
		# result = {}
		# #print("============>")
		# if os.path.exists(path):
		#     #print("path exists ====>")
		#     with ZipFile('media/'+str(zipFile)+'.zip','w') as zip: 
		#         # writing each file one by one 
		#         for file in file_paths:
		#             zip.write(file)

		#     fsize=os.path.getsize('media/'+str(zipFile)+'.zip')/float(1<<20)

		#     result = {"size":round(fsize, 2),
		#     "name":str(settings.MEDIA_ROOT)+str(zipFile)+'.zip',
		#     "path":'media/'+str(zipFile)+'.zip'
		#     }
		# #print("result ====>",result)
		# return Response(result, status=status.HTTP_201_CREATED)
import PyPDF2
@permission_classes((AllowAny, ))
def downloadfileskey(request, path, user):
	if request.method == "GET":
		import os
		from constants import constants
		extensions = constants.extensions
		print(path,"pathhhh")
		extension11=os.path.splitext(path)
		print(extension11,"extension11")
		content_type=extension11[-1]
		file_cont_type=extensions[content_type]
		print(file_cont_type,"file_cont_type")
		from azure.storage.blob import BlockBlobService, PublicAccess
		block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
		container_name ='docullycontainer'
		path_name = path.split("/")
		path_name[-2] = str(path_name[-2]).replace('%20',' ')
		print("file_name===>",path_name[-2])
		pathhh=path.split("/")
		pathoffile=pathhh[3]+"/"+pathhh[-1]
		print(pathoffile,"pathoffile")
		dataroomid=DataroomFolder.objects.filter(path=pathoffile).values('dataroom_id','user_id')
		print(dataroomid,"dataroomid")
		room_id_here=dataroomid[0]['dataroom_id']
		print(room_id_here,"room_id_here")
		user_id_here=dataroomid[0]['user_id']
		print(room_id_here,"room_id_here")
		block_blob_service.get_blob_to_path(container_name, pathoffile, pathhh[-1])
		print(block_blob_service.get_blob_to_path(container_name, pathoffile, pathhh[-1]),"blooob")
		extension=os.path.splitext(pathhh[-1])[-1]
		data_of_per=DataroomGroupPermission.objects.filter(dataroom=room_id_here).values("is_doc_as_pdf","is_watermarking","is_excel_as_pdf")
		if (extension=='.pdf'):
			if os.path.exists(pathhh[-1]):
				from dataroom.serializers import WatermarkingSerializer
				from dataroom.pdf_watermarking import GeneratePDF
				print(user,"useruser")
				nameee=User.objects.filter(email=user).values('id')
				print(nameee,"nameee")
				watermarking = Watermarking.objects.filter(dataroom_id=int(room_id_here)).order_by('id')
				for i in watermarking:
					i.user_id=nameee[0]['id']
				print("watermarking =====>",watermarking)
				serializer = WatermarkingSerializer(watermarking,many=True)
				data = serializer.data
				print(data,"in download")
				from userauth import utils
				ip = utils.get_client_ip(request)
				print(room_id_here,"pkk")
				GeneratePDF(data,ip,user,room_id_here)
				watermarkfile="/home/cdms_backend/cdms2/Admin_Watermark/"+str(room_id_here)+".pdf"
				outputfile="/home/cdms_backend/cdms2/dataroom/success.pdf"
				pdf_writer=PyPDF2.PdfFileWriter()
				with open(pathhh[-1], 'rb') as fh:
					pdf=PyPDF2.PdfFileReader(fh)
					with open(watermarkfile,'rb') as watermarkfile:
						watermarkfile_pdf=PyPDF2.PdfFileReader(watermarkfile)
						for i in range(pdf.getNumPages()):
							p=pdf.getPage(i)
							p.mergePage(watermarkfile_pdf.getPage(0))
							pdf_writer.addPage(p)
						with open(outputfile,'wb') as outputfileeee:
							pdf_writer.write(outputfileeee)
						with open(outputfile, 'rb') as output:
							response = HttpResponse(output.read(), content_type=file_cont_type)
							response['Content-Disposition'] = 'inline; filename=' + pathhh[-1]
						return response
					raise Http404
		elif (extension!='.pdf') and (data_of_per[0]["is_watermarking"]==False) and (data_of_per[0]["is_excel_as_pdf"]==False):
			print("in second elif")
			import os, uuid, sys
			from azure.storage.blob import BlockBlobService, PublicAccess
			block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
			container_name ='docullycontainer'
			block_blob_service.get_blob_to_path(container_name, pathoffile, pathhh[-1])
			if os.path.exists(path_name[-1]):
				with open(pathhh[-1], 'rb') as fh:
					response = HttpResponse(fh.read(), content_type=file_cont_type)
					response['Content-Disposition'] = 'inline; filename=' + path_name[-1]
					os.remove(path_name[-1])
					return response
				raise Http404
		elif (extension=='.docx' or ".doc" or ".ppt" or ".pptx" or '.xlsx' or '.xls' or '.csv') and (data_of_per[0]["is_doc_as_pdf"]==True) and (data_of_per[0]["is_watermarking"]==True):
			print("in last elif file-view")
			import os, uuid, sys
			from azure.storage.blob import BlockBlobService, PublicAccess
			block_blob_service = BlockBlobService(account_name='docullystorage', account_key='ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw==')
			container_name ='docullycontainer'
			from pathlib import Path
			filnn=extension11[0].split("/")
			print(filnn,"filnn")
			# pdf_filename=str(filnn[-1])+".pdf"
			# print(pdf_filename,"pdf_filename")
			# blobname=filnn[-2]+"/"+pdf_filename
			# print(blobname,"blobname")
			# print(container_name, blobname,pdf_filename,"details download")
			# block_blob_service.get_blob_to_path(container_name, blobname,pdf_filename)
			# with open(pdf_filename, 'rb') as output:
			# 	response = HttpResponse(output.read(), content_type="application/pdf")
			# 	response['Content-Disposition'] = 'inline; filename=' + pdf_filename
			# 	return response


import codecs

# def page(self,request):
# 	return render(request,'pdffile.html')

# class PrivateView(APIView):
	# authentication_classes = (TokenAuthentication, )
	# permission_classes = (IsAuthenticated, )

	# def post(self,request):
	# 	dd=request.data['dataaa']
	# 	print(dd,"dddd")
	# 	extension=os.path.splitext(dd)[-1]
	# 	print(extension,"extension")
	# 	DATA=({"extension":extension})
	# 	if extension=='.png':
	# 		return Response(DATA,status=200)
	# 	elif extension=='.pdf':
	# 		print("pdf")
	# 		return render(request, 'file-view.html')
from datetime import datetime, timedelta
from azure.storage.blob import (
    BlockBlobService,
    ContainerPermissions,
    BlobPermissions,
    PublicAccess,
)
def PrivateView(request):
	if request.method=='GET':
		dpath=request.GET.get('dataaa')
		print(dpath,"in class")
		print("get")
		extension=os.path.splitext(str(dpath))[-1]
		print(extension,"extension")
		d="https://docullystorage.blob.core.windows.net/docullycontainer/new_test_8_april_later_1c3978ed-c31e-40fa-821b-1d0fcb222730/new_sample.pdf?se=2020-05-24T18%3A28%3A12Z&sp=r&sv=2018-03-28&sr=b&sig=hZeMvUga1/mzOILLOvg08RYl4qs8kPer5lEgB/6fdjk%3D"
		dd="https://docullystorage.blob.core.windows.net/docullycontainer/data_room_new_16_april_dbb18eea-14f1-420a-b303-24fd8a7d7c29/testing35.pdf?sp=r&st=2020-06-09T15:06:13Z&se=2020-06-09T23:06:13Z&spr=https&sv=2019-10-10&sr=b&sig=j8jzz8j8xWfTNPWSlEb99lV%2B%2Bh1LqTR83pWObOX%2BSak%3D"
		AZURE_ACC_NAME = 'docullystorage'
		AZURE_PRIMARY_KEY = 'ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw=='
		AZURE_CONTAINER = 'docullycontainer'
		if extension==".pdf":
			return render(request, 'pdffile.html',{"data":str(dpath)})
		elif extension==".docx" or ".doc":
			from pathlib import Path
			p = Path(str(dpath))
			frbl=str(dpath).split("/")
			print(frbl,"frbl")
			filen=frbl[-1]
			print(filen,"filen")
			from pathlib import Path
			filename1 = Path(filen)
			filename_wo_ext = filename1.with_suffix('')
			print(filename_wo_ext,"filename_wo_ext")
			from datetime import datetime, timedelta
			AZURE_BLOB=frbl[4]+"/"+str(filename_wo_ext)+".pdf"
			block_blob_service = BlockBlobService(account_name=AZURE_ACC_NAME, account_key=AZURE_PRIMARY_KEY)
			sas_url = block_blob_service.generate_blob_shared_access_signature(AZURE_CONTAINER,AZURE_BLOB,permission=BlobPermissions.READ,expiry= datetime.utcnow() + timedelta(hours=1))
			finaldoc='https://'+AZURE_ACC_NAME+'.blob.core.windows.net/'+AZURE_CONTAINER+'/'+AZURE_BLOB+'?'+sas_url
			print(finaldoc,"finaldoc")
			return render(request, 'pdffile.html',{"data":finaldoc})
		elif extension==".pptx" or ".ppt":
			from pathlib import Path
			p = Path(str(dpath))
			frbl=str(dpath).split("/")
			print(frbl,"frbl")
			filen=frbl[-1]
			print(filen,"filen")
			from pathlib import Path
			filename1 = Path(filen)
			filename_wo_ext = filename1.with_suffix('')
			print(filename_wo_ext,"filename_wo_ext")
			from datetime import datetime, timedelta
			AZURE_BLOB=frbl[4]+"/"+str(filename_wo_ext)+".pdf"
			block_blob_service = BlockBlobService(account_name=AZURE_ACC_NAME, account_key=AZURE_PRIMARY_KEY)
			sas_url = block_blob_service.generate_blob_shared_access_signature(AZURE_CONTAINER,AZURE_BLOB,permission=BlobPermissions.READ,expiry= datetime.utcnow() + timedelta(hours=1))
			finaldoc='https://'+AZURE_ACC_NAME+'.blob.core.windows.net/'+AZURE_CONTAINER+'/'+AZURE_BLOB+'?'+sas_url
			print(finaldoc,"finaldoc")
			return render(request, 'pdffile.html',{"data":finaldoc})


# from rest_framework.permissions import AllowAny
# @permission_classes((AllowAny, ))
class PrivateView_image(APIView):
	authentication_classes = (TokenAuthentication, )
	permission_classes = (IsAuthenticated, )
	def post(self,request):
		dpath=request.data['dataaa']
		print(dpath,"in class")
		print("get")
		extension=os.path.splitext(str(dpath))[-1]
		print(extension,"extension")
		frbl=str(dpath).split("/")
		AZURE_BLOB=frbl[4]+"/"+frbl[-1]
		AZURE_ACC_NAME = 'docullystorage'
		AZURE_PRIMARY_KEY = 'ddIGey4fa6zz/FnWjMgPm5zN35BgIEDsaY6K18dpTFpkqUAJRD6efPBpXZGdBG8ICnyWWE8Y/PPGZQ0ajUeZTw=='
		AZURE_CONTAINER = 'docullycontainer'
		block_blob_service = BlockBlobService(account_name=AZURE_ACC_NAME, account_key=AZURE_PRIMARY_KEY)
		sas_url = block_blob_service.generate_blob_shared_access_signature(AZURE_CONTAINER,AZURE_BLOB,permission=BlobPermissions.READ,expiry= datetime.utcnow() + timedelta(hours=1))
		finaldoc='https://'+AZURE_ACC_NAME+'.blob.core.windows.net/'+AZURE_CONTAINER+'/'+AZURE_BLOB+'?'+sas_url
		return Response(finaldoc)
# def PrivateView_image(request):
# 	if request.method=='POST':
		

		#req = urllib.request.urlopen(dd)
	   # charset = req.info().get_content_charset()
	   # content = req.read()
	   # print(type(content))
	   # decoded_data = codecs.encode(content, 'base64')
		
		# print(DATA,"DATA")
		
# def pdffile()